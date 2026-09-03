# -*- coding: utf-8 -*-
"""
DeepSeek 桌面助手 v3 —— 自用微应用
功能：系统托盘常驻 + 全局热键呼出 + 调用 DeepSeek API 聊天
      v3 新增：联网搜索增强 / 流式输出 / 多会话标签 / 历史本地保存
技术：PySide6 (GUI + 原生网络 QNetworkAccessManager)
完全免费，仅依赖公开库。大哥练手可读可改。

v5 Codex UI 改版重设计：
- 三栏布局：左侧栏 220px → 主内容区 flex → 右侧面板 260px
- 侧边栏：Logo + 导航项 + 分隔线 + 分组标签 + 对话列表 + 底部设置
- 主区欢迎页：问候语 + 标题「今天想做什么？」+ 3 张功能卡片 + 最近对话列表
- 输入区：卡片式包裹，附件按钮 + 输入框 + 圆形发送按钮 + 提示文字
- 右侧交付物面板：标题 + 计数 badge + 交付卡片含左侧彩色竖条 + 操作按钮
- 暗色主题：基于参考设计稿的浅色方案映射
"""

import sys
import os
import re
import base64
import ctypes
import html as html_mod
import json
import logging
import urllib.parse
import time
from datetime import datetime
from pathlib import Path

# 长任务断点续跑 / 心跳（纯标准库模块，无 Qt 依赖）
import task_resume

# v4.88：自动化任务（定时提醒 / 定时执行 Agent 任务）
import automation

# v4.109：路由旁路日志（只写不读，任何失败静默吞掉，绝不干扰对话）
try:
    import route_log   # v4.109：旁路埋点（路由/用量/技能），只写不读
except Exception:      # 旁路模块缺失绝不能拖垮主界面（冻结环境 import 失败必炸）
    route_log = None

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QTextEdit, QTextBrowser, QLineEdit, QPushButton, QSystemTrayIcon, QMenu, QLabel,
    QScrollArea, QCheckBox, QComboBox, QFileDialog, QMessageBox, QDialog,
    QStackedWidget, QSplitter, QFrame, QSizePolicy, QListWidget, QGridLayout, QSpinBox, QTimeEdit,
    QGroupBox,
)
from PySide6.QtCore import Qt, QUrl, QTimer, QByteArray, Signal, QPoint, QThread, QRect, QAbstractNativeEventFilter, QMutex, QWaitCondition
from PySide6.QtNetwork import QNetworkAccessManager, QNetworkRequest, QNetworkReply
from PySide6.QtGui import QIcon, QColor, QFont, QAction, QDesktopServices, QTextCursor, QGuiApplication, QPixmap, QPainter

# ===== 无边框窗口：Win32 原生边缘缩放所需常量 =====
import ctypes
_WM_NCHITTEST = 0x0084
_WM_NCLBUTTONDBLCLK = 0x00A3
_HTCLIENT = 1
_HTCAPTION = 2
_HTLEFT = 10
_HTRIGHT = 11
_HTTOP = 12
_HTTOPLEFT = 13
_HTTOPRIGHT = 14
_HT_BOTTOM = 15
_HT_BOTTOMLEFT = 16
_HT_BOTTOMRIGHT = 17
_EDGE_MARGIN = 6  # 边缘命中热区像素

import config
from config import (
    APP_DIR, CONFIG_PATH, load_config, load_skills,
    AGENT_SYS_APPEND, TOOL_RESULT_LIMIT, load_dynamic_skills,
    get_app_icon, APP_VERSION,
)
from session import SessionStore
import search as search_mod
import tools as tools_mod
from permissions import PermissionEngine, MODES
import voice as voice_mod
from agent import AgentWorker
from PySide6.QtMultimedia import QMediaPlayer, QAudioOutput
import chat_web  # v4.104：WebEngine 聊天渲染引擎

log = logging.getLogger("dsdesktop")

# ===== 暗色主题调色板（基于 Codex 参考设计稿浅→暗映射） =====
THEME = {
    # ---- Backgrounds (bento-console-flat-light · Google 蓝) ----
    "bg": "#F7F8FC",          # 主窗口/工作区底色
    "sidebar": "#EEF1F8",      # 侧栏底
    "surface": "#FFFFFF",
    "surface_raised": "#FFFFFF",
    "card": "#FFFFFF",         # 所有浮卡/顶栏/状态条/输入框

    # ---- Panels ----
    "panel": "#FFFFFF",
    "panel2": "#F1F3F4",
    "elev": "#FFFFFF",
    "border": "#E5E7EB",
    "border_highlight": "#D2D5DA",

    # ---- Text ----
    "text": "#202124",
    "dim": "#5F6368",
    "faint": "#9AA0A6",
    "placeholder": "#9AA0A6",
    "weak": "#9AA0A6",

    # ---- Accent (Google 蓝 #1A73E8) ----
    "accent": "#1A73E8",
    "accent_hover": "#1765CC",
    "accent_pressed": "#155BB5",
    "accent_disabled": "#A6C4F0",
    "accent2": "#A142F4",      # 强调紫（审稿类节点）
    "danger": "#EA4335",
    "ok": "#34A853",
    "warn": "#FBBC04",
    "link": "#1A73E8",

    # ---- Delivery ----
    "delivery_purple": "#A142F4",
    "magenta": "#EC4899",

    # ---- Welcome cards ----
    "card_blue_bg": "rgba(26,115,232,0.08)",
    "card_blue_icon": "#1A73E8",
    "card_green_bg": "rgba(52,168,83,0.10)",
    "card_green_icon": "#34A853",
    "card_orange_bg": "rgba(251,188,4,0.12)",
    "card_orange_icon": "#FBBC04",

    # ---- Chat bubbles (浅色，蓝强调) ----
    "user_bg": "#E8F0FE",
    "user_border": "#C6DAFC",
    "user_text": "#202124",
    "asst_bg": "#FFFFFF",
    "asst_border": "#E5E7EB",
    "asst_text": "#202124",

    # ---- Tools ----
    "tool_running": "#FBBC04",
    "tool_done": "#34A853",

    # ---- Sidebar ----
    "sidebar_hover": "#E4E8F2",
    "sidebar_active": "#1A73E8",
    "sidebar_active_bar": "#1A73E8",
    "sidebar_active_text": "#FFFFFF",
    "separator": "#E5E7EB",
    "avatar_asst": "#1A73E8",
    "avatar_user": "#202124",

    # ---- Delivery labels ----
    "delivery_blue": "#1A73E8",
    "delivery_green": "#34A853",
    "delivery_orange": "#FBBC04",

    # ---- 交互状态补充 ----
    "blue_hover": "#E4E8F2",
    "border_hover": "#D2D5DA",
    "focus_glow": "0 0 0 3px rgba(26,115,232,0.15)",
    "row_hover": "#F7F8FC",
    "row_hover_alt": "#E4E8F2",
    "secondary_btn_hover": "#F1F3F4",
    "secondary_btn_press": "#E8EAED",

    # ---- Prism 虹彩渐变 (保留兼容，新版未使用) ----
    "prism": "qlineargradient(x1:0,y1:0,x2:1,y2:0,stop:0 #FF6B9D,stop:0.2 #C44569,stop:0.4 #F8B500,stop:0.6 #00D2FF,stop:0.8 #7B68EE,stop:1 #FF69B4)",
    "prism_soft": "qlineargradient(x1:0,y1:0,x2:1,y2:0,stop:0 rgba(255,107,157,0.14),stop:0.5 rgba(0,210,255,0.10),stop:1 rgba(123,104,238,0.14))",
}

# 交付物类型 → 彩色竖条配色
DELIVERY_COLORS = {
    "image": ("delivery_blue", "accent"),
    "img":   ("delivery_blue", "accent"),
    "video": ("delivery_purple", "magenta"),
    "doc":   ("delivery_green", "ok"),
    "md":    ("delivery_green", "ok"),
    "code":  ("delivery_orange", "tool_running"),
    "py":    ("delivery_orange", "tool_running"),
    "js":    ("delivery_orange", "tool_running"),
    "html":  ("delivery_orange", "tool_running"),
    "css":   ("delivery_orange", "tool_running"),
    "file":  ("delivery_blue", "accent"),
}


class _GenThread(QThread):
    """后台跑阻塞型生成（生图/生视频），结果经信号回主线程。"""
    result = Signal(object)   # (rel, kind, name) 或 错误字符串

    def __init__(self, fn, *args, **kwargs):
        super().__init__()
        self.fn = fn
        self.args = args
        self.kwargs = kwargs

    def run(self):
        try:
            self.result.emit(self.fn(*self.args, **self.kwargs))
        except Exception as e:
            self.result.emit(f"异常：{e}")


class PerfWorker(QThread):
    """v4.78：性能基线后台跑批（不阻塞 UI），结果经 done 信号回主线程。"""
    done = Signal(object)  # 结果 dict 或 {"error": str}

    def __init__(self, save_baseline=False):
        super().__init__()
        self.save_baseline = save_baseline

    def run(self):
        try:
            import perf_baseline as pb
            metrics = pb.run_benchmarks()
            if self.save_baseline:
                pb.save_baseline(metrics)
            cmp, verdict = pb.compare_with_baseline(metrics)
            run_path = pb.save_run(metrics, cmp, verdict)
            self.done.emit({
                "metrics": metrics, "comparison": cmp, "verdict": verdict,
                "run_path": run_path, "saved_baseline": self.save_baseline,
            })
        except Exception as e:
            self.done.emit({"error": str(e)})


class OrchestrateWorker(QThread):
    """小说一条龙（3Phase+2检查）。短篇按目标字数一次性写满；长篇按章，用「续写」出下一章。
    交互：爆款雷达后弹出选项让你选定（模型不自决）；运行中可随时暂停并插入修改意见。"""
    stage = Signal(int, str)
    node_status = Signal(int, str)   # running / done / error
    log = Signal(str)
    done = Signal(str)
    need_choice = Signal(list, str)  # 爆款雷达后：[(label,text)...], raw
    paused = Signal(str)             # 暂停等待插入意见时：当前阶段标签

    STAGES = [
        ("爆款雷达", "blue",   "Phase1 选题：雷达扫描给出 3 个结构化爆款切入点，交给你选定"),
        ("选题验证", "orange", "检查1：用『选题三问』深化你选定的方向，给出核心设定"),
        ("写手成稿", "green",  "Phase2 写作：按选定方向产出成稿（短篇循环凑足目标字数 / 长篇写一章）"),
        ("虚拟编辑 审稿", "purple", "Phase3 审稿：按责编 rubric 挑必须改的问题"),
        ("终稿定稿", "red",    "检查2+定稿：按问题修改并终稿复审，输出可投递版本（含标题）"),
    ]

    def __init__(self, mw, topic, platform, length_type="短篇", target_words=2000,
                 prev_state=None, task_id=None, start_stage=0):
        super().__init__()
        self.mw = mw
        self.topic = topic
        self.platform = platform
        self.length_type = length_type if length_type in ("短篇", "长篇") else "短篇"
        try:
            self.target_words = int(target_words) if target_words else 2000
        except Exception:
            self.target_words = 2000
        self.prev_state = prev_state
        self.task_id = task_id or task_resume.new_task_id()
        self.start_stage = max(0, int(start_stage))
        self._stop_requested = False
        self._cancelled = False  # v4.101：用户取消编排（保留检查点可续跑）
        # 续写时：长篇章节号 +1，短篇仍视为第 1 段（扩写）；断点恢复(start_stage>0)则沿用原章节号
        if prev_state and self.length_type == "长篇":
            self.chapter = prev_state.get("chapter", 0) + (1 if self.start_stage == 0 else 0)
        else:
            self.chapter = 1
        self.full_draft = ""
        # 断点恢复：还原上一阶段已写好的正文，避免终稿拼装时丢失
        self.stage2_draft = (prev_state.get("stage2_draft", "") if prev_state else "")
        self.final_state = None
        self.messages = None  # run() 中赋值，检查点使用
        # 交互状态（选题闸门 + 中途暂停）
        self._pause_requested = False
        self._pending_feedback = None
        self._chosen_direction = None
        self._choice_options = []
        self._pause_mutex = QMutex()
        self._pause_cond = QWaitCondition()
        self._choice_mutex = QMutex()
        self._choice_cond = QWaitCondition()
        # D 项（轨迹记忆）：阶段计时 + 重试计数，供成功时采集轨迹
        self._stage_start = None
        self._stage_durations = {}
        self._retry_count = 0

    def request_stop(self):
        """主线程调用，请求取消（下一节点前生效）。同时释放挂起的等待，避免卡死。"""
        self._stop_requested = True
        self._choice_cond.wakeOne()
        self._pause_cond.wakeOne()

    def release_locks(self):
        """取消时由 UI 调用：唤醒所有挂起等待，让 run() 回到 stop 检查。"""
        self._stop_requested = True
        self._choice_cond.wakeOne()
        self._pause_cond.wakeOne()

    def request_pause(self):
        """主线程调用：请求在下一个阶段边界（或写手下一段）暂停，等待插入意见。"""
        self._pause_requested = True

    def resume_with_feedback(self, feedback):
        """主线程调用：带着作者意见恢复运行（feedback 为空则仅恢复）。"""
        self._pending_feedback = (feedback or "").strip() or None
        self._pause_cond.wakeOne()

    def choose(self, idx):
        """主线程调用：用户从爆款雷达选项中选定第 idx 个。"""
        try:
            idx = int(idx)
        except Exception:
            idx = -1
        if 0 <= idx < len(self._choice_options):
            self._chosen_direction = self._choice_options[idx]["text"]
        self._choice_cond.wakeOne()

    def choose_custom(self, text):
        """主线程调用：用户自填切入点方向。"""
        t = (text or "").strip()
        if t:
            self._chosen_direction = t
        self._choice_cond.wakeOne()

    @staticmethod
    def _build_novelist_system_prompt():
        """专用小说写作系统提示：网文主编/写作教练人设，注入作者方法论，禁止把真实家人写进角色。"""
        return (
            "你是一位资深网络小说主编兼写作教练，负责『小说一条龙』流水线"
            "（爆款雷达→选题验证→写手成稿→虚拟编辑审稿→终稿定稿）。\n"
            "作者的方法论（必须贯穿全流水线，不要违背）：\n"
            "1. 选题三问：①我想写什么（核心设定）②读者为什么看（爽点/情绪钩子）"
            "③凭什么我能写好（差异化卖点）。每个选题都要过这三问。\n"
            "2. 叙事：第一人称「我」；每约 500 字必须有一个钩子"
            "（悬念/反转/信息差/情绪爆发）；番茄小说节奏快、章末必须留悬念。\n"
            "3. 交付：番茄一稿一投（一次性成稿、不反复折腾）；场景化叙事、show-don't-tell。\n"
            "4. 纯虚构：严禁把作者的真实家庭成员或本人写进角色"
            "（例如『大哥』『葱头』『小臭』『雪糕』『xyb』等——这些只是作者的私人信息，与小说无关；"
            "除非作者明确要求，否则不得作为角色名、原型或背景人物出现）。角色、地名、机构一律原创。\n"
            "5. 各司其职：严格按当前流水线环节的要求输出，不越界、不提前做后续环节的事。\n"
            "6. 短篇必须写完整：开头→发展→高潮→结局四段齐备，收尾干净、人物命运有交代，"
            "严禁停在半路或留『待续』悬念。字数是软目标，完结优先于凑字数。\n"
        )

    def _platform_hint(self):
        """平台差异化提示。"""
        hints = {
            "番茄小说": "适配番茄小说：快节奏网文，开头强钩子，每章留悬念，第一人称沉浸。",
            "知乎": "适配知乎：盐选故事风格，文笔细腻，逻辑严密，适度知识性。",
            "公众号": "适配公众号：短段落，金句加粗，情感共鸣，不出现具体城市/行业。",
            "抖音": "适配抖音：极短句，强情绪，画面感强，适合口播节奏。",
            "头条": "适配头条：信息密度高，标题党，争议性切入点。",
        }
        return hints.get(self.platform, "")

    def _prompt(self, i, name, brief):
        base = (f"你正在参与『小说一条龙』流水线，当前环节：{name}（{brief}）。\n"
                f"主题：{self.topic}\n平台：{self.platform}\n{self._platform_hint()}\n")
        if i == 0:
            return base + (
                "你是题材趋势分析师。围绕主题做爆款切入点雷达扫描，输出 3 个结构化切入点，"
                "每个用『选项N：』开头（N=1/2/3），格式如下：\n"
                "选项1：[一句话切入点]\n"
                "·为什么爆：（目标人群 + 情绪钩子 + 市场空白）\n"
                "·风险：（同质化 / 违规 / 难以持续）\n"
                "·与番茄同类爆款的差异点\n"
                "只输出这 3 个选项，不要替作者做决定，也不要写正文。")
        if i == 1:
            cd = self._chosen_direction or "（作者尚未明确，请基于主题给出最稳妥的推荐方向）"
            return base + (
                f"作者已选定切入点：\n{cd}\n"
                "你是选题验证专家，不要另选方向。请用『选题三问』检验该方向的可行性："
                "①我想写什么（核心设定）②读者为什么看（爽点/情绪钩子）③凭什么我能写好（差异化）。"
                "指出 1-2 个风险与强化点，并给出落地核心设定：主角人设、核心冲突、开局强钩子。"
                "输出可直接交给写手的方向书。")
        if i == 2:
            if self.length_type == "长篇":
                return base + (f"本次写第 {self.chapter} 章，本章目标约 {self.target_words} 字。"
                               "你是写作教练：第一人称「我」，场景化叙事、show-don't-tell；"
                               "每约 500 字一个钩子（悬念/反转/信息差）；章末必须留强悬念钩子，本章内不写完结。")
            return base + (f"请按选定方向写小说，总体目标约 {self.target_words} 字。"
                           "你是写作教练：第一人称「我」，场景化叙事、show-don't-tell；"
                           "每约 500 字一个钩子；开头 50 字内强钩子，节奏快、爽点密集。"
                           "本次先写开头 800-1000 字；若已有正文务必接着写、不重复开头。")
        if i == 3:
            return base + ("你是责编。通读当前成稿（短篇=全文 / 长篇=最新一章），"
                           "按审稿 rubric 挑 3 条必须改的问题：①开头钩子是否够强 ②爽点/情绪密度 ③人设一致性 ④节奏与违规词。"
                           "每条给具体位置与可执行的改法。")
        if i == 4:
            if self.length_type == "长篇":
                return base + (f"基于审稿意见把第 {self.chapter} 章定稿："
                               f"输出『标题』一行，然后『完整本章正文（保留全部内容、约 {self.target_words} 字）』。")
            # 短篇：正文已在写手阶段写满，终稿只补标题+定稿说明，避免重写截断
            return base + ("不要重写正文。基于审稿意见只做两件事："
                           "1) 给全文拟一个吸睛标题；2) 用一两句话说明你做了哪些定稿处理。"
                           "（正文已在上一步生成，此处不要重复输出大段正文）")

    def _save_node(self, name, text):
        """落盘：每节点产出存 md 到工作区 orchestrate/ 目录。"""
        try:
            d = os.path.join(APP_DIR, "orchestrate")
            os.makedirs(d, exist_ok=True)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe = name.replace(" ", "_").replace("/", "_")
            if self.length_type == "长篇":
                safe = f"第{self.chapter}章_{safe}"
            fp = os.path.join(d, f"{ts}_{safe}.md")
            with open(fp, "w", encoding="utf-8") as f:
                f.write(f"# {name}\n\n主题：{self.topic}  平台：{self.platform}  "
                        f"类型：{self.length_type}"
                        f"（第{self.chapter}章）\n\n{text}\n")
            self.log.emit(f"💾 已保存：{os.path.basename(fp)}\n")
        except Exception as e:
            self.log.emit(f"⚠️ 落盘失败：{e}\n")

    def _save_checkpoint(self, stage_idx, extra=None):
        """阶段边界存档：把可恢复状态（累积 messages + 阶段号 + 草稿）落到磁盘。"""
        if self.messages is None:
            return
        state = {
            "task_id": self.task_id,
            "task_type": "orchestrate",
            "stage": stage_idx,                    # 已完成的最高阶段索引
            "topic": self.topic,
            "platform": self.platform,
            "length_type": self.length_type,
            "target_words": self.target_words,
            "chapter": self.chapter,
            "messages": self.messages,            # 含已完成各阶段的输入/输出
            "stage2_draft": getattr(self, "stage2_draft", ""),
            "full_draft": getattr(self, "full_draft", ""),
        }
        if extra:
            state.update(extra)
        try:
            task_resume.save_checkpoint(self.mw.cfg, state)
        except Exception:
            pass

    def _recover_chosen_direction(self):
        """断点恢复：从 messages 还原作者已选定的切入点，供选题验证/写手阶段使用。"""
        if not self.messages:
            return
        for m in reversed(self.messages):
            c = m.get("content", "") if isinstance(m, dict) else ""
            if m.get("role") == "user" and "作者选定的切入点" in c:
                try:
                    self._chosen_direction = c.split("作者选定的切入点：\n", 1)[-1]
                except Exception:
                    pass
                return

    def _record_stage_duration(self, name):
        """D 项：记录本阶段耗时（秒），供成功时采集进轨迹。"""
        try:
            if self._stage_start is not None:
                self._stage_durations[name] = round(time.time() - self._stage_start, 1)
        except Exception:
            pass

    def _call(self, messages):
        resp = self.mw._agent_call(messages, [], None)
        return resp.get("content") or ""

    def _call_with_retry(self, messages, retries=3, backoff=3, note=None):
        """C 项（借鉴 Prime-Agent 策略重试）：LLM 调用失败（超时/503/网络）自动换思路重试。

        - 仅重试「抛异常」的失败（网络/API 错误）；空内容不算失败，不重试。
        - 每次重试前在消息副本上追加「换一种写法」提示，改变策略而非单纯重发；
          用副本不改 caller 的 messages，保持检查点/上下文干净。
        - 退避 sleep 用 QThread.msleep（不阻塞主界面，因 worker 在后台线程）。
        - 全部失败则抛出最后一次异常，由调用方干净退出（保留上一节点检查点，可断点续跑）。
        """
        try:
            return self._call(messages)
        except Exception as e:
            last = e
        note = note or ("⚠️ 上一次生成调用失败或超时，请换一种方式重试"
                        "（控制篇幅、避免特殊符号），直接输出正文。")
        for k in range(1, retries + 1):
            self._retry_count += 1  # D 项：累计真实重试次数，供轨迹采集
            self.log.emit(f"🔁 重试（第 {k}/{retries} 次，{backoff*k}s 后）…\n")
            self.msleep(backoff * 1000 * k)
            msgs = list(messages)
            msgs.append({"role": "user", "content": note})
            try:
                return self._call(msgs)
            except Exception as e:
                last = e
        raise last if last is not None else RuntimeError("LLM 调用失败")

    def _parse_options(self, text):
        """从爆款雷达输出中解析出结构化选项，返回 [(label, full_text), ...]（最多 3 个）。"""
        import re
        pat = re.compile(r"选项\s*(\d+)\s*[：:]\s*(.*)")
        blocks = {}
        cur = None
        buf = []
        for line in text.split("\n"):
            m = pat.match(line.strip())
            if m:
                if cur is not None:
                    blocks[cur] = "\n".join(buf).strip()
                cur = int(m.group(1))
                buf = [m.group(2)]
            elif cur is not None:
                buf.append(line)
        if cur is not None:
            blocks[cur] = "\n".join(buf).strip()
        opts = []
        for k in sorted(blocks):
            full = blocks[k]
            first_line = full.split("\n", 1)[0].strip()
            label = (first_line[:30] + "…") if len(first_line) > 30 else first_line
            opts.append({"label": label, "text": full})
        if not opts:
            # 兜底：模型未按格式输出，把整段作为唯一选项
            opts = [{"label": "模型给出的方案", "text": text.strip()}]
        return opts[:3]

    def _maybe_pause(self, messages, label):
        """阶段边界检查暂停请求：若用户点了暂停，则挂起等待插入意见后再继续。"""
        if not self._pause_requested:
            return
        self._pause_requested = False
        self._pending_feedback = None
        self.paused.emit(label)
        self._pause_mutex.lock()
        self._pause_cond.wait(self._pause_mutex)
        self._pause_mutex.unlock()
        if self._pending_feedback:
            self.log.emit(f"\n💬 你插入的修改意见：{self._pending_feedback}\n")
            messages.append({"role": "user", "content":
                f"【作者对上一阶段的修改意见，请在后续环节采纳】{self._pending_feedback}"})
            self._pending_feedback = None

    def _run_writing_stage(self, messages, brief):
        """写手成稿：短篇循环续写直到达到目标字数；长篇单章。返回最终成稿文本（同时存 self.stage2_draft）。"""
        if self.length_type == "长篇":
            p = self._prompt(2, self.STAGES[2][0], brief)
            messages.append({"role": "user", "content": p})
            text = self._call(messages)
            messages.append({"role": "assistant", "content": text})
            self.log.emit(f"（第 {self.chapter} 章，约 {len(text)} 字）\n")
            self.stage2_draft = text
            try:
                task_resume.update_heartbeat(self.mw.cfg, self.task_id)
            except Exception:
                pass
            return text
        # 短篇：循环写主体，接近目标后明确写「完整结局」，确保不半路截断、不烂尾
        target = self.target_words
        draft = ""
        max_iter = min(30, max(3, target // 700 + 3))
        ending_emitted = False
        for it in range(max_iter):
            if self._stop_requested:
                break
            # 写手中也可随时暂停并插入修改意见
            self._maybe_pause(messages, "写手成稿（可在此插入修改意见）")
            if self._stop_requested:
                break
            if it == 0:
                p = self._prompt(2, self.STAGES[2][0], brief)
            else:
                tail = draft[-700:] if len(draft) > 700 else draft
                if len(draft) >= target * 0.6:
                    # 主体已铺垫够，转而写完整结局收尾（只写一次）
                    ending_emitted = True
                    p = (f"当前已写约 {len(draft)} 字（全文目标约 {target} 字）：\n……{tail}\n\n"
                         f"现在写【完整结局】：把前面的铺垫推向高潮并干净收尾，"
                         f"人物命运要有交代、情绪要有落点。约 {max(target - len(draft), 300)} 字左右，"
                         f"必须完结、不要留『待续』悬念。只输出结局正文，不重复前文。")
                else:
                    p = (f"当前已写约 {len(draft)} 字：\n……{tail}\n\n"
                         f"接着上文继续写下一节，约 800-1000 字，推进剧情、保持钩子密度，"
                         f"不要重复开头、此时不要写结局、继续铺垫发展。")
            messages.append({"role": "user", "content": p})
            try:
                chunk = self._call_with_retry(messages, retries=3, backoff=4)
            except Exception as e:
                self.log.emit(f"⚠️ 写手阶段调用失败：{e}（已重试，保留已写约 {len(draft)} 字）\n")
                messages.pop()  # 移除本次未完成的 user 轮，保持上下文干净
                break
            messages.append({"role": "assistant", "content": chunk})
            draft += chunk
            try:
                task_resume.update_heartbeat(self.mw.cfg, self.task_id)
            except Exception:
                pass
            # 一旦写完结局立即收尾；或已到目标且结局已写，停止（避免双结尾）
            if ending_emitted or (len(draft) >= target and ending_emitted):
                break
        self.log.emit(f"（已累计约 {len(draft)} 字）\n")
        self.stage2_draft = draft
        return draft

    def run(self):
        # 续写：复用上一次的底稿上下文（选题已完成，跳过前两个节点）
        if self.prev_state:
            messages = list(self.prev_state["messages"])
            for idx in range(0, max(self.start_stage, 2) if self.start_stage > 0 else 2):
                # 断点恢复：把已完成阶段节点标为 done（含 0/1；start_stage>0 时覆盖到 k-1）
                if idx < self.start_stage:
                    self.node_status.emit(idx, "done")
            if self.start_stage >= 1:
                self._recover_chosen_direction()
        else:
            sys_prompt = self._build_novelist_system_prompt()
            # D 项（轨迹记忆）：新鲜启动检索相似成功轨迹，拼成 few-shot 参考注入系统提示
            try:
                import trace_log
                if self.mw.cfg.get("orch_trace_enabled", True):
                    few = trace_log.build_fewshot_for(
                        self.mw.cfg, self.topic, self.platform, self.length_type)
                    if few:
                        sys_prompt += "\n\n" + few
            except Exception:
                pass
            messages = [{"role": "system", "content": sys_prompt}]
        self.messages = messages
        for i, (name, color, brief) in enumerate(self.STAGES):
            if i < self.start_stage:
                continue
            if i in (0, 1) and self.prev_state and self.start_stage == 0:
                continue
            # 中途暂停 + 插入意见（阶段边界，选题外的每个节点前都允许）
            if i >= 1:
                self._maybe_pause(messages, f"准备进入【{name}】前")
            if self._stop_requested:
                self.node_status.emit(i, "error")
                self.log.emit(f"⏹ 已取消（{name} 前）\n")
                self.done.emit("已取消")
                return
            self._stage_start = time.time()  # D 项：阶段计时起点
            self.stage.emit(i, name)
            self.node_status.emit(i, "running")
            self.log.emit(f"\n【{name}】{brief}\n")
            if i == 0:
                # 爆款雷达：产出 3 个结构化切入点，交作者选定（模型不自决）
                messages.append({"role": "user", "content": self._prompt(0, name, brief)})
                try:
                    text = self._call_with_retry(messages, retries=3)
                except Exception as e:
                    self.node_status.emit(i, "error")
                    self.log.emit(f"❌ {name} 多次重试仍失败：{e}\n")
                    self.done.emit(f"生成失败：{name}")
                    return
                messages.append({"role": "assistant", "content": text})
                self.log.emit(text + "\n")
                self._save_node(name, text)
                self._record_stage_duration(name)  # D 项：记录阶段耗时
                self.node_status.emit(i, "done")
                self._save_checkpoint(i)
                # —— 选题闸门：暂停，等作者从 3 个切入点中选定 ——
                self._choice_options = self._parse_options(text)
                self.need_choice.emit(self._choice_options, text)
                self._choice_mutex.lock()
                self._choice_cond.wait(self._choice_mutex)
                self._choice_mutex.unlock()
                if self._stop_requested:
                    self.log.emit("⏹ 已取消（选题阶段）\n")
                    self.done.emit("已取消")
                    return
                chosen = self._chosen_direction or self._choice_options[0]["text"]
                self.log.emit(f"\n✅ 你选定的切入点：\n{chosen}\n")
                messages.append({"role": "user", "content": f"作者选定的切入点：\n{chosen}"})
                messages.append({"role": "assistant", "content": "明白，将以该方向进入选题验证与写作。"})
                continue
            if i == 2:
                try:
                    text = self._run_writing_stage(messages, brief)
                except Exception as e:
                    self.node_status.emit(i, "error")
                    self.log.emit(f"❌ {name} 多次重试仍失败：{e}\n")
                    self.done.emit(f"生成失败：{name}")
                    return
            else:
                messages.append({"role": "user", "content": self._prompt(i, name, brief)})
                try:
                    text = self._call_with_retry(messages, retries=3)
                except Exception as e:
                    self.node_status.emit(i, "error")
                    self.log.emit(f"❌ {name} 多次重试仍失败：{e}\n")
                    self.done.emit(f"生成失败：{name}")
                    return
                messages.append({"role": "assistant", "content": text})
            if i == 4:  # 终稿定稿
                # 短篇：正文已在写手阶段写满，终稿只补标题+说明，拼到正文前，避免重写截断
                if self.length_type == "短篇" and getattr(self, "stage2_draft", ""):
                    self.full_draft = f"{text}\n\n{self.stage2_draft}"
                else:
                    self.full_draft = text
                # 终稿节点必须展示并保存完整可投递文本（标题+全部正文），而非仅标题
                self.log.emit("\n📜 终稿（可投递完整版，含全部正文）\n" + "-" * 24 + "\n" + self.full_draft + "\n")
                self._save_node(name, self.full_draft)
                self._record_stage_duration(name)  # D 项：记录阶段耗时
                self.node_status.emit(i, "done")
                self._save_checkpoint(i)
            else:
                self.log.emit(text + "\n")
                self._save_node(name, text)
                self._record_stage_duration(name)  # D 项：记录阶段耗时
                self.node_status.emit(i, "done")
                self._save_checkpoint(i)
        # 兜底：极少数情况下终稿被跳过，短篇优先用写手正文，否则取最后一段 assistant
        if not self.full_draft:
            if self.length_type == "短篇" and self.stage2_draft:
                self.full_draft = self.stage2_draft
            else:
                for m in reversed(messages):
                    if m.get("role") == "assistant":
                        self.full_draft = m.get("content") or ""
                        break
        self.final_state = {
            "topic": self.topic,
            "platform": self.platform,
            "length_type": self.length_type,
            "target_words": self.target_words,
            "messages": messages,
            "chapter": self.chapter,
            "draft": self.full_draft,
        }
        self.done.emit("小说一条龙生成完成")


class _ASRWorker(QThread):
    """后台语音识别：wav -> 文本（硅基流动 ASR）。"""
    sig_text = Signal(str)
    sig_error = Signal(str)

    def __init__(self, wav_path, sf):
        super().__init__()
        self.wav_path = wav_path
        self.sf = sf

    def run(self):
        try:
            txt = voice_mod.transcribe(self.wav_path, self.sf)
            self.sig_text.emit(txt)
        except Exception as e:
            self.sig_error.emit(str(e))


class _TTSWorker(QThread):
    """后台语音合成：文本 -> mp3（edge-tts）。"""
    sig_mp3 = Signal(str)

    def __init__(self, text, tts):
        super().__init__()
        self.text = text
        self.tts = tts

    def run(self):
        try:
            mp3 = voice_mod.synthesize(self.text, self.tts)
            self.sig_mp3.emit(mp3)
        except Exception as e:
            log.error("TTS 失败: %s", e)


def resource_path(rel):
    """定位资源：兼容 PyInstaller onefile/onedir 与开发时。"""
    rel = rel.replace("/", os.sep)
    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        p = os.path.join(meipass, rel)
        if os.path.exists(p):
            return p
    exe_dir = os.path.dirname(sys.executable)
    candidates = [
        os.path.join(exe_dir, rel),
        os.path.join(exe_dir, "_internal", rel),
    ]
    dev_dir = os.path.dirname(os.path.abspath(__file__))
    candidates.append(os.path.join(dev_dir, rel))
    for p in candidates:
        if os.path.exists(p):
            return p
    return candidates[0]


class ThemedDialog(QDialog):
    """深色、无边框、置顶、自动居中的模态对话框基类。"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowFlags(Qt.Dialog | Qt.WindowStaysOnTopHint | Qt.FramelessWindowHint)
        self.setStyleSheet(
            f"QDialog{{background:{THEME['card']};border:1px solid {THEME['border_highlight']};"
            f"border-radius:10px;}}")
        self._build()

    def _build(self):
        pass

    def showEvent(self, e):
        super().showEvent(e)
        if self.parent():
            p = self.parent().geometry()
            self.move(p.x() + (p.width() - self.width()) // 2,
                      p.y() + (p.height() - self.height()) // 2)

    @staticmethod
    def _btn_style(bg, fg):
        hover = THEME["accent_hover"] if bg == THEME["accent"] else THEME["panel2"]
        pressed = THEME["accent"] if bg == THEME["accent"] else THEME["elev"]
        return (f"QPushButton{{background:{bg};color:{fg};border:none;border-radius:10px;"
                f"padding:8px 20px;font-size:13px;font-weight:600;}}"
                f"QPushButton:hover{{background:{hover};}}"
                f"QPushButton:pressed{{background:{pressed};}}")


class ConfirmDialog(ThemedDialog):
    def _build(self):
        self._result = False
        self._trusted = False
        lay = QVBoxLayout(self)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(16)
        self._title = QLabel("")
        self._title.setStyleSheet(f"color:{THEME['text']};font-size:15px;font-weight:600;")
        self._detail = QLabel("")
        self._detail.setWordWrap(True)
        self._detail.setStyleSheet(f"color:{THEME['dim']};font-size:13px;line-height:1.6;")
        lay.addWidget(self._title)
        lay.addWidget(self._detail)
        self._trust = QCheckBox("本次会话信任（不再逐个确认）")
        self._trust.setStyleSheet(f"color:{THEME['dim']};font-size:12px;")
        lay.addWidget(self._trust)
        row = QHBoxLayout()
        row.addStretch(1)
        no = QPushButton("取消")
        no.setFixedHeight(34)
        no.setStyleSheet(self._btn_style(THEME["elev"], THEME["dim"]))
        no.clicked.connect(self.reject)
        yes = QPushButton("允许")
        yes.setFixedHeight(34)
        yes.setStyleSheet(self._btn_style(THEME["accent"], "#ffffff"))
        yes.clicked.connect(self._on_yes)
        row.addWidget(no)
        row.addWidget(yes)
        lay.addLayout(row)
        self.adjustSize()

    def set_text(self, title, detail):
        self._title.setText(title)
        self._detail.setText(detail)

    def _on_yes(self):
        self._result = True
        self._trusted = self._trust.isChecked()
        self.accept()

    def result(self):
        return self._result

    def trusted(self):
        return self._trusted


class RenameDialog(ThemedDialog):
    def _build(self):
        self._text = ""
        lay = QVBoxLayout(self)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(16)
        t = QLabel("重命名对话")
        t.setStyleSheet(f"color:{THEME['text']};font-size:15px;font-weight:600;")
        lay.addWidget(t)
        self._edit = QLineEdit()
        self._edit.setPlaceholderText("例如：小说《死亡倒计时》大纲")
        self._edit.setStyleSheet(
            f"QLineEdit{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:8px 10px;font-size:13px;color:{THEME['text']};}}"
            f"QLineEdit:focus{{border:1px solid {THEME['accent']};}}")
        lay.addWidget(self._edit)
        row = QHBoxLayout()
        row.addStretch(1)
        no = QPushButton("取消")
        no.setFixedHeight(34)
        no.setStyleSheet(self._btn_style(THEME["elev"], THEME["dim"]))
        no.clicked.connect(self.reject)
        yes = QPushButton("保存")
        yes.setFixedHeight(34)
        yes.setStyleSheet(self._btn_style(THEME["accent"], "#ffffff"))
        yes.clicked.connect(self._on_ok)
        row.addWidget(no)
        row.addWidget(yes)
        lay.addLayout(row)
        self._edit.returnPressed.connect(self._on_ok)
        self.setMinimumWidth(360)
        self.adjustSize()

    def _on_ok(self):
        self._text = self._edit.text().strip()
        self.accept()

    def text_value(self):
        return self._text


class InfoDialog(ThemedDialog):
    def _build(self):
        lay = QVBoxLayout(self)
        lay.setContentsMargins(24, 24, 24, 24)
        lay.setSpacing(16)
        self._title = QLabel("")
        self._title.setStyleSheet(f"color:{THEME['text']};font-size:15px;font-weight:600;")
        self._detail = QLabel("")
        self._detail.setWordWrap(True)
        self._detail.setStyleSheet(f"color:{THEME['dim']};font-size:13px;line-height:1.6;")
        lay.addWidget(self._title)
        lay.addWidget(self._detail)
        row = QHBoxLayout()
        row.addStretch(1)
        ok = QPushButton("好的")
        ok.setFixedHeight(34)
        ok.setStyleSheet(self._btn_style(THEME["accent"], "#ffffff"))
        ok.clicked.connect(self.accept)
        row.addWidget(ok)
        lay.addLayout(row)
        self.adjustSize()

    def set_text(self, title, detail):
        self._title.setText(title)
        self._detail.setText(detail)


class MultiLineInput(QTextEdit):
    """多行输入框：Enter / Ctrl+Enter 发送，Shift+Enter 换行；随内容自动增高；支持粘贴图片。"""
    sendRequested = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptRichText(False)
        self.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.setFixedHeight(38)
        self.textChanged.connect(self.adjust_height)
        self.image_files = []

    def keyPressEvent(self, e):
        if e.key() in (Qt.Key_Return, Qt.Key_Enter):
            if e.modifiers() & Qt.ShiftModifier:
                # Shift+Enter: 换行
                super().keyPressEvent(e)
            else:
                # Enter: 发送
                self.sendRequested.emit()
                e.accept()
        else:
            super().keyPressEvent(e)

    def adjust_height(self):
        new_h = int(self.document().size().height()) + 16
        self.setFixedHeight(max(38, min(new_h, 160)))

    def insertFromMimeData(self, source):
        if source.hasImage():
            image = source.imageData()
            if image and not image.isNull():
                ts = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                tmp_dir = Path(APP_DIR) / "temp" / "images"
                tmp_dir.mkdir(parents=True, exist_ok=True)
                img_path = tmp_dir / f"paste_{ts}.png"
                image.save(str(img_path))
                self.image_files.append(str(img_path))
                # Show placeholder text so user knows image was captured
                cursor = self.textCursor()
                cursor.insertText("[\u56fe\u7247\u5df2\u7c98\u8d34 " + str(len(self.image_files)) + "]")
                self.textChanged.emit()
        else:
            super().insertFromMimeData(source)

    def get_images(self):
        return self.image_files

    def clear_images(self):
        self.image_files = []


class SessionManagerDialog(QDialog):
    """v4.79：会话管理——置顶 / 分组 / 批量删除 / 筛选。"""

    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.store = parent.store
        self.setWindowTitle("会话管理")
        self.setMinimumSize(540, 560)
        self.setStyleSheet(f"QDialog{{background:{THEME['bg']};}}")

        self._row_sids = []
        self._checks = []

        lay = QVBoxLayout(self)
        lay.setContentsMargins(16, 16, 16, 16)
        lay.setSpacing(12)

        # ---- 搜索 + 分组筛选 ----
        top = QHBoxLayout()
        top.setSpacing(8)
        self.search = QLineEdit()
        self.search.setPlaceholderText("搜索会话标题…")
        self.search.setFixedHeight(34)
        self.search.setStyleSheet(
            f"QLineEdit{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 10px;font-size:13px;color:{THEME['text']};}}"
            f"QLineEdit:focus{{border:1px solid {THEME['accent']};}}")
        self.search.textChanged.connect(self._refresh)
        top.addWidget(self.search, 1)

        self.folder_filter = QComboBox()
        self.folder_filter.setFixedHeight(34)
        self.folder_filter.setMinimumWidth(140)
        self.folder_filter.setStyleSheet(
            f"QComboBox{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 8px;font-size:13px;color:{THEME['text']};}}"
            f"QComboBox:focus{{border:1px solid {THEME['accent']};}}")
        self.folder_filter.currentTextChanged.connect(self._refresh)
        top.addWidget(self.folder_filter)
        lay.addLayout(top)

        # ---- 列表 ----
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setStyleSheet("QScrollArea{border:none;background:transparent;}")
        self.list_widget = QWidget()
        self.list_lay = QVBoxLayout(self.list_widget)
        self.list_lay.setContentsMargins(0, 0, 0, 0)
        self.list_lay.setSpacing(6)
        self.scroll.setWidget(self.list_widget)
        lay.addWidget(self.scroll, 1)

        # ---- 底部：全选 / 批量删除 / 完成 ----
        bottom = QHBoxLayout()
        bottom.setSpacing(8)
        self.sel_all = QCheckBox("全选")
        self.sel_all.setStyleSheet(f"QCheckBox{{color:{THEME['text']};font-size:13px;}}")
        self.sel_all.stateChanged.connect(self._toggle_all)
        bottom.addWidget(self.sel_all)
        bottom.addStretch(1)
        bulk = QPushButton("🗑 批量删除")
        bulk.setFixedHeight(32)
        bulk.setStyleSheet(
            f"QPushButton{{background:{THEME['card']};color:{THEME['danger']};"
            f"border:1px solid {THEME['border']};border-radius:8px;padding:0 14px;font-size:13px;}}"
            f"QPushButton:hover{{border-color:{THEME['danger']};background:#FCE8E6;}}")
        bulk.clicked.connect(self._bulk_delete)
        bottom.addWidget(bulk)
        done = QPushButton("完成")
        done.setDefault(True)
        done.setFixedHeight(32)
        done.setStyleSheet(
            f"QPushButton{{background:{THEME['accent']};color:#FFFFFF;border:none;"
            f"border-radius:8px;padding:0 18px;font-size:13px;font-weight:500;}}"
            f"QPushButton:hover{{background:{THEME['accent_hover']};}}")
        done.clicked.connect(self.accept)
        bottom.addWidget(done)
        lay.addLayout(bottom)

        self._refresh(initial=True)

    # ------------------------------------------------------------------
    def _refresh(self, initial=False, preserve_folder=True):
        if initial:
            cur = self.folder_filter.currentText()
            self.folder_filter.blockSignals(True)
            self.folder_filter.clear()
            self.folder_filter.addItem("全部会话")
            for f in self.store.list_folders():
                self.folder_filter.addItem(f)
            # 还原选择
            idx = self.folder_filter.findText(cur) if preserve_folder and cur else -1
            if idx >= 0:
                self.folder_filter.setCurrentIndex(idx)
            self.folder_filter.blockSignals(False)

        # 清旧行
        while self.list_lay.count():
            it = self.list_lay.takeAt(0)
            w = it.widget()
            if w:
                w.deleteLater()
        self._row_sids = []
        self._checks = []

        folder_sel = self.folder_filter.currentText()
        folder_key = "" if folder_sel in ("全部会话", "") else folder_sel
        items = self.store.all_sorted(query=self.search.text(), folder=folder_key)

        if not items:
            empty = QLabel("没有匹配的会话")
            empty.setStyleSheet(f"color:{THEME['placeholder']};font-size:13px;padding:12px 0;")
            self.list_lay.addWidget(empty)
            return

        for s in items:
            self._build_row(s)

    def _build_row(self, s):
        row = QWidget()
        row.setStyleSheet(
            f"background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:10px;")
        rl = QHBoxLayout(row)
        rl.setContentsMargins(10, 8, 10, 8)
        rl.setSpacing(8)

        chk = QCheckBox()
        chk.setFixedSize(18, 18)
        rl.addWidget(chk)
        self._checks.append(chk)

        pin = QPushButton("⚑" if s.pinned else "☆")
        pin.setFixedSize(30, 30)
        pin.setToolTip("置顶" if not s.pinned else "取消置顶")
        pin.setStyleSheet(
            f"QPushButton{{background:transparent;border:none;font-size:16px;"
            f"color:{'#F4B400' if s.pinned else THEME['placeholder']};border-radius:6px;}}"
            f"QPushButton:hover{{background:{THEME['sidebar_hover']};}}")
        pin.clicked.connect(lambda _, sid=s.sid: self._pin_toggle(sid))
        rl.addWidget(pin)

        title = QPushButton(s.title or "新会话")
        title.setStyleSheet(
            f"QPushButton{{background:transparent;border:none;text-align:left;"
            f"font-size:13px;color:{THEME['text']};padding:0 4px;}}"
            f"QPushButton:hover{{color:{THEME['accent']};}}")
        title.setToolTip("点击打开此会话")
        title.clicked.connect(lambda _, sid=s.sid: self._switch(sid))
        rl.addWidget(title, 1)

        fcombo = QComboBox()
        fcombo.setFixedHeight(28)
        fcombo.setMinimumWidth(96)
        fcombo.setEditable(True)
        fcombo.setStyleSheet(
            f"QComboBox{{background:{THEME['bg']};border:1px solid {THEME['border']};"
            f"border-radius:6px;padding:0 6px;font-size:12px;color:{THEME['text']};}}")
        fcombo.addItem("未分组")
        for f in self.store.list_folders():
            if f != s.folder:
                fcombo.addItem(f)
        fcombo.setCurrentText(s.folder or "未分组")
        fcombo.currentTextChanged.connect(lambda txt, sid=s.sid: self._folder_changed(sid, txt))
        rl.addWidget(fcombo)

        delb = QPushButton("×")
        delb.setFixedSize(26, 26)
        delb.setStyleSheet(
            f"QPushButton{{background:transparent;color:{THEME['placeholder']};"
            f"border:none;font-size:16px;font-weight:600;border-radius:6px;}}"
            f"QPushButton:hover{{color:{THEME['danger']};background:{THEME['sidebar_hover']};}}")
        delb.setToolTip("删除会话")
        delb.clicked.connect(lambda _, sid=s.sid, t=s.title: self._del_one(sid, t))
        rl.addWidget(delb)

        self._row_sids.append(s.sid)
        self.list_lay.addWidget(row)

    def _pin_toggle(self, sid):
        s = self.store.sessions.get(sid)
        if not s:
            return
        self.store.set_pinned(sid, not s.pinned)
        self._refresh()

    def _folder_changed(self, sid, text):
        self.store.set_folder(sid, text)
        # 分组筛选状态下改名可能使该行消失，刷新即可
        self._refresh()

    def _switch(self, sid):
        self.parent._switch_session(sid)
        self.parent._refresh_session_combo()
        self.parent._refresh_recent_on_welcome()
        self.accept()

    def _del_one(self, sid, title):
        if QMessageBox.question(
                self, "删除会话",
                f"确定删除会话「{title or '新会话'}」？此操作不可撤销。",
                QMessageBox.Yes | QMessageBox.No) != QMessageBox.Yes:
            return
        self.store.remove(sid)
        self.parent._refresh_session_combo()
        self.parent._refresh_recent_on_welcome()
        self._refresh(initial=True)

    def _toggle_all(self, state):
        for c in self._checks:
            c.setChecked(bool(state))

    def _bulk_delete(self):
        sids = [sid for sid, c in zip(self._row_sids, self._checks) if c.isChecked()]
        if not sids:
            QMessageBox.information(self, "提示", "请先勾选要删除的会话。")
            return
        if QMessageBox.question(
                self, "批量删除",
                f"确定删除选中的 {len(sids)} 个会话？此操作不可撤销。",
                QMessageBox.Yes | QMessageBox.No) != QMessageBox.Yes:
            return
        n = self.store.remove_many(sids)
        self.parent._refresh_session_combo()
        self.parent._refresh_recent_on_welcome()
        self.sel_all.setChecked(False)
        self._refresh(initial=True)


class _EdgeResizeFilter(QAbstractNativeEventFilter):
    """用 QAbstractNativeEventFilter 给无边框窗口加边缘缩放（避免 nativeEvent 签名冲突）。

    PySide6 冻结模式下 QMainWindow.nativeEvent 虚方法签名与 C++ 不一致
    （expected 2, got 3），覆写必然崩溃。改用 installNativeEventFilter + Filter
    子类，其 nativeEventFilter 固定只收 (eventType, message) 两个参数，无此问题。
    """

    def __init__(self, window):
        super().__init__()
        self._win = window

    def nativeEventFilter(self, eventType, message):
        if eventType == b"windows_generic_MSG":
            try:
                msg = ctypes.cast(message, ctypes.POINTER(ctypes.wintypes.MSG)).contents
                if msg.message == _WM_NCHITTEST and not self._win.isMaximized():
                    # v4.108 H-11：lParam 的屏幕坐标是【物理像素】，而 win.x()/width()
                    # 是【逻辑坐标】——高分屏（125%/150% 缩放）直接相减导致热区偏移、
                    # 边缘缩放失灵。先按 devicePixelRatio 把物理坐标归一为逻辑坐标。
                    dpr = float(self._win.devicePixelRatioF() or 1.0)
                    x = (msg.lParam & 0xFFFF) / dpr
                    y = ((msg.lParam >> 16) & 0xFFFF) / dpr
                    wx = x - self._win.x()
                    wy = y - self._win.y()
                    w = self._win.width()
                    h = self._win.height()
                    m = _EDGE_MARGIN
                    if wx <= m and wy <= m:
                        return True, ctypes.wintypes.LPARAM(_HTTOPLEFT)
                    if wx >= w - m and wy <= m:
                        return True, ctypes.wintypes.LPARAM(_HTTOPRIGHT)
                    if wx <= m and wy >= h - m:
                        return True, ctypes.wintypes.LPARAM(_HT_BOTTOMLEFT)
                    if wx >= w - m and wy >= h - m:
                        return True, ctypes.wintypes.LPARAM(_HT_BOTTOMRIGHT)
                    if wx <= m:
                        return True, ctypes.wintypes.LPARAM(_HTLEFT)
                    if wx >= w - m:
                        return True, ctypes.wintypes.LPARAM(_HTRIGHT)
                    if wy <= m:
                        return True, ctypes.wintypes.LPARAM(_HTTOP)
                    if wy >= h - m:
                        return True, ctypes.wintypes.LPARAM(_HT_BOTTOM)
            except Exception:
                pass
        return False, None


def _flatten_text_content(content_list):
    """把多模态 list content 的文本部分拼成字符串；无文本返回 ''。"""
    texts = []
    for p in content_list:
        if isinstance(p, dict) and p.get("type") == "text":
            t = p.get("text", "")
            if t:
                texts.append(t)
        elif isinstance(p, str) and p:
            texts.append(p)
    return "\n".join(texts).strip()


def _sanitize_filename(name):
    """v4.102 hotfix：清洗 Windows 文件名非法字符（\\ / : * ? " < > |），
    并裁掉首尾空格/点。返回清洗后的安全文件名；空结果回退为 'file'。"""
    if not name:
        return "file"
    safe = re.sub(r'[\\/:*?"<>|]', "_", name)
    safe = safe.strip().strip(".")
    return safe or "file"


def _vision_debug(msg):
    """v4.102 fix5：写视觉链路调试日志到 ~/Documents/小臭玩AI/vision_debug.log，
    用户/我们都能找得到（之前写 APP_DIR，源码与 exe 路径不同导致用户找不到）。"""
    try:
        from datetime import datetime
        log_dir = os.path.expanduser("~/Documents/小臭玩AI")
        os.makedirs(log_dir, exist_ok=True)
        p = os.path.join(log_dir, "vision_debug.log")
        with open(p, "a", encoding="utf-8") as f:
            f.write(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}\n")
    except Exception:
        pass


def _compress_image_for_api(path, max_dim=1568, quality=85):
    """v4.102 fix4+5：图片预处理。
    fix4：避免真实截图多张发图时 payload 超 DeepSeek 视觉模型单请求体限制（>10MB 返 400）。
    fix5：**无论大小先强制 convertToFormat(RGB888) 标准化**，防止 CMYK/灰度/索引色/异常
    PNG 被 DeepSeek 拒为 unsupported image（实测纯 RGB/ARGB 都接受，唯独某些剪贴板格式被拒）。

    策略：读图后立即 RGB888 标准化 → 若宽/高>max_dim 等比缩放 → 转 JPEG 条件：
    原>200KB 或需缩放 或 含 alpha（带透明的图转 JPEG 走 alpha 展平到白底）。
    失败时返 None（send 走直接读 PNG 的 fallback）。"""
    try:
        from PySide6.QtGui import QImage, QPainter
        from PySide6.QtCore import Qt, QBuffer, QIODevice
        import base64 as _b64
        orig_kb = os.path.getsize(path) // 1024
        img = QImage(path)
        if img.isNull():
            _vision_debug(f"compress: QImage.isNull for {os.path.basename(path)}")
            return None
        # 关键：强制 RGB888 标准化，消除 CMYK/灰度/索引色等不被 DeepSeek 接受的格式
        fmt_rgb = getattr(getattr(QImage, "Format", QImage), "Format_RGB888", None) or getattr(QImage, "Format_RGB888", None)
        if fmt_rgb and img.format() != fmt_rgb:
            src_fmt = str(img.format())
            img = img.convertToFormat(fmt_rgb)
            _vision_debug(f"compress: format convert {src_fmt} -> RGB888 for {os.path.basename(path)}")
        w, h = img.width(), img.height()
        need_resize = max(w, h) > max_dim
        if need_resize:
            img = img.scaled(max_dim, max_dim, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        # 转 JPEG 条件：原 >200KB 或需缩放 或 含 alpha（小但异常格式也强制 JPEG 更稳）
        use_jpeg = orig_kb > 200 or need_resize or img.hasAlphaChannel()
        buf = QBuffer()
        buf.open(QIODevice.WriteOnly)
        if use_jpeg and img.hasAlphaChannel():
            # alpha 展平到白底，避免 JPEG 出现黑边
            bg = QImage(img.size(), fmt_rgb)
            bg.fill(0xFFFFFFFF)
            p = QPainter(bg)
            p.drawImage(0, 0, img)
            p.end()
            img = bg
        if use_jpeg:
            img.save(buf, "JPEG", quality)
            mime = "image/jpeg"
        else:
            img.save(buf, "PNG")
            mime = "image/png"
        raw = bytes(buf.data())
        if len(raw) < 100:  # 太小怀疑输出损坏
            _vision_debug(f"compress: output too small ({len(raw)}B) for {path}")
            return None
        final_kb = len(raw) // 1024
        b64 = _b64.b64encode(raw).decode()
        url = f"data:{mime};base64,{b64}"
        _vision_debug(f"compress OK: {os.path.basename(path)} | {w}x{h} | {orig_kb}KB->{final_kb}KB | mime={mime}")
        return (url, mime, orig_kb, final_kb, need_resize or use_jpeg)
    except Exception as e:
        _vision_debug(f"compress EXC for {os.path.basename(path)}: {type(e).__name__}: {e}")
        return None


def _normalize_image_dataurl(url):
    """v4.102 fix6：把任意 image data URL 重编码为 DeepSeek 视觉模型稳接受的 RGB JPEG。
    作为「最后一道关卡」覆盖所有发图来源（贴图/附件/历史消息/兜底回退）——只要进 API
    前统一过一遍，即可规避特殊格式/超大图被拒（HTTP 400 无正文）的情况。无法解码则原样返回。"""
    if not isinstance(url, str) or not url.startswith("data:image/"):
        return url
    try:
        import base64 as _b64
        from PySide6.QtGui import QImage
        from PySide6.QtCore import Qt, QBuffer, QIODevice
        header, _, b64 = url.partition(",")
        if "base64" not in header:
            return url
        raw = _b64.b64decode(b64)
        img = QImage.fromData(raw)
        if img.isNull():
            return url
        fmt_rgb = QImage.Format.Format_RGB888
        if img.format() != fmt_rgb:
            img = img.convertToFormat(fmt_rgb)
        buf = QBuffer(); buf.open(QIODevice.WriteOnly)
        img.save(buf, "JPEG", 85)
        out = bytes(buf.data())
        return f"data:image/jpeg;base64,{_b64.b64encode(out).decode()}"
    except Exception as e:
        _vision_debug(f"_normalize_image_dataurl EXC: {e}")
        return url


def _model_supports_vision(model):
    """模块级：判断模型是否支持图像输入（多模态视觉）。仅这些模型才在
    _sanitize_msg_for_api 中保留 image_url；其余模型图像被归一化为纯文本标签，
    避免把 list content 原样发给不支持视觉的接口导致 400。"""
    if not model:
        return False
    m = str(model).lower()
    return any(k in m for k in (
        "vision", "vl", "gpt-4o", "gpt-4v", "gpt-4.1", "qwen-vl", "qwen2-vl",
        "qwen2.5-vl", "qwen2_5-vl", "glm-4v", "glm-4v-plus", "yi-vl",
        "internvl", "minicpm-v", "deepseek-vl", "step-1v", "moondream",
        "cogvlm", "fuyu", "idefics", "kosmos",
    ))


def _extract_file_image_parts(text, app_dir):
    """v4.102 hotfix：从文本中的 [文件: path] / [file: path] 标记提取图片，
    转为 OpenAI 兼容的 image_url content parts。

    返回 (clean_text, image_parts)。只处理真实存在的图片文件；不存在或非图片文件
    会在原处保留提示文本，避免模型误解。clean_text 已去掉被成功加载图片的标记。
    """
    import base64, mimetypes
    pattern = re.compile(r"\[(?:\u6587\u4ef6|file):\s*([^\]\n]+)\]")
    image_parts = []

    def _replace(m):
        rel = m.group(1).strip()
        # 相对路径以 app_dir 为基准解析；incoming/xxx 即可定位
        path = os.path.join(app_dir, rel) if not os.path.isabs(rel) else rel
        path = os.path.normpath(path)
        if not os.path.isfile(path):
            return f"\n[文件不存在: {rel}]\n"
        ext = os.path.splitext(path)[1].lower().lstrip(".")
        mime = mimetypes.guess_type(path)[0] or ("image/png" if ext == "png" else "image/jpeg")
        if not mime.startswith("image/"):
            return f"\n[非图片文件: {rel}]\n"
        try:
            # v4.102 fix6：附件图片也走 _compress_image_for_api，强制 RGB888 标准化 +
            # JPEG 体积控制，与贴图路径统一，避免特殊格式/超大图被 DeepSeek 视觉模型拒（HTTP 400）。
            comp = _compress_image_for_api(path)
            if comp:
                image_parts.append({"type": "image_url", "image_url": {"url": comp[0]}})
                _vision_debug(f"attach 图片压缩 OK: {os.path.basename(path)} -> {comp[1]} {comp[3]}KB")
            else:
                # 压缩失败回退原始读取（保证不阻断发送）
                with open(path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode()
                image_parts.append({"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}})
            return ""  # 成功加载的图片标记从文本中移除
        except Exception as e:
            return f"\n[图片加载失败: {rel} ({e})]\n"

    clean_text = pattern.sub(_replace, text)
    return clean_text.strip(), image_parts


def _sanitize_msg_for_api(m, vision_ok=False):
    """把一条 session 消息清洗为 OpenAI 兼容接口可接受的 {"role","content"}；不可接受返回 None。

    v4.79 hotfix：session.messages 里混有 UI 展示用的 tool / tool_log 角色、
    content=None、以及多模态 list content——直接发给 DeepSeek 等接口会 400
    （unknown variant `tool_log` / invalid content）。历史消息必须经此过滤。

    vision_ok=True（目标模型支持视觉，如 deepseek-v4-flash-vision-exp）时，若
    content 是含 image_url 的 list，则保留 list 结构（仅 text + image_url 两种合法
    part）原样发视觉模型，让模型真正"看图"；否则仍归一化为纯文本（兼容旧逻辑）。
    """
    if not isinstance(m, dict):
        return None
    role = m.get("role")
    if role not in ("user", "assistant"):
        return None
    c = m.get("content")
    if isinstance(c, list):
        has_img = any(isinstance(p, dict) and p.get("type") == "image_url" for p in c)
        if has_img and vision_ok:
            # 视觉模型：保留 list 结构，只留 text / image_url 两种合法 part
            cleaned = []
            for p in c:
                if not isinstance(p, dict):
                    continue
                t = p.get("type")
                if t == "text":
                    if p.get("text", "").strip():
                        cleaned.append(p)
                elif t == "image_url":
                    # v4.102 fix6：统一重编码图片为 RGB JPEG（兜底所有来源）
                    u = (p.get("image_url") or {}).get("url", "")
                    cleaned.append({"type": "image_url", "image_url": {"url": _normalize_image_dataurl(u)}})
            if cleaned:
                return {"role": role, "content": cleaned}
        # 非视觉（默认或视觉模型但无图）：归一化为纯文本
        c = _flatten_text_content(c)
        if has_img:
            c = (c + "\n[图片]") if c else "[图片]"
    if not isinstance(c, str) or not c.strip():
        return None
    return {"role": role, "content": c}


class ChatWindow(QMainWindow):
    def __init__(self, cfg):
        super().__init__()
        self.cfg = cfg
        # v4.75：记忆加密启动解锁——配置了口令就先设 cipher 再加载记忆，
        # 否则 .enc 文件读不到（_cipher 默认 None 只认明文）
        _enc_pw = self.cfg.get("memory_encryption_passphrase", "")
        if _enc_pw:
            try:
                from memory_store import set_encryption
                set_encryption(_enc_pw)
            except Exception as e:
                log.warning("记忆加密解锁失败: %s", e)
        self.store = SessionStore()
        # v4.88：自动化任务（定时提醒 / 定时执行 Agent 任务）
        self.automation_store = automation.AutomationStore()
        self._auto_timer = QTimer(self)
        self._auto_timer.setInterval(1000)
        self._auto_timer.timeout.connect(self._on_automation_tick)
        self._auto_timer.start()
        # v4.75：对话内搜索 / 单条重生成状态
        self._search_query = ""
        self._search_matches = []
        self._search_pos = -1
        self._edit_target_idx = None
        self._skills = load_skills()
        self._busy = False
        self._busy_timeout = QTimer(self)
        self._busy_timeout.setSingleShot(True)
        self._busy_timeout.timeout.connect(self._reset_busy)
        # v4.58：渲染节流 + 批量保存，避免工具调用期间 UI 线程被信号洪水打满
        self._render_timer = QTimer(self)
        self._render_timer.setSingleShot(True)
        self._render_timer.timeout.connect(self._on_render_tick)
        self._pending_render_force = False
        self._rendered_msg_count = 0  # v4.60：增量渲染消息计数
        self._save_timer = QTimer(self)
        self._save_timer.setSingleShot(True)
        self._save_timer.timeout.connect(self._do_save)
        self._reply = None
        self._streaming = False
        self._streaming_text = ""
        self._sse_buf = ""
        # v4.108 M-25：live 工具卡台账（只在 DOM 中、不进 messages）。
        # agent_index（worker 内配对键）→ 卡片信息；UI 用自增 _live_seq 生成唯一
        # DOM id，彻底避开「agent 每批 index 从 0 起」导致的跨轮/跨任务 id 撞车。
        # 台账同时让「全量重建（会话切换/页面自愈）」能按序恢复进行中的工具卡。
        self._live_tools = {}
        self._live_seq = 0
        self.agent_mode = self.cfg.get("agent_mode", False)
        self.agent_skip_confirm = self.cfg.get("agent_skip_confirm", False)
        # ---- 权限引擎（v4.50，借鉴 andrewyng/openworker 的 permissions.py）----
        _home = os.path.expanduser("~")
        _mode = self.cfg.get("permission_mode") or (
            "auto" if self.agent_skip_confirm else "interactive")
        _auto_allow = set(self.cfg.get("permission_auto_allow", []))
        _external_allow = set(self.cfg.get("permission_external_allow", []))
        _scope = [os.path.join(_home, "Documents"),
                  os.path.join(_home, "Desktop"),
                  config.APP_DIR, _home]
        self.permission_engine = PermissionEngine(_mode, _auto_allow, _scope, _external_allow)
        self.permission_mode = self.permission_engine.mode
        self.session_trusted = False  # 本次会话信任危险/浏览器操作（确认弹窗勾选后置 True）
        self._agent_worker = None
        # ---- 滚动跟随状态（解耦自 setHtml：流式每 chunk 重渲会重置滚动条，读滚动条判断贴底
        #      与 setHtml 抢同一根滚动条会误判，改用粘性标志 + 用户上滑事件驱动）----
        self._follow_bottom = True          # True=新内容来时滚到底；用户上滑看历史置 False
        self._scroll_seq = 0                # 滚底请求序号：每次渲染 +1，过期的延迟滚底请求自动作废
        self._last_view_change = 0.0        # 时间戳哨兵：程序改 chat_view 的时间戳，屏蔽异步 valueChanged 误判
        # ---- 头像素材：Qt 富文本表格单元格内 div 的固定 height 被忽略并撑满行高（实测头像
        #      被拉成贯穿整条消息的蓝条），改用 <img>（replaced element，Qt 尊重固有尺寸），
        #      PNG 在首次启动时生成到 APP_DIR/avatars/，后续直接复用磁盘文件。----
        self._ensure_avatars()
        # ---- 语音（实时语音：ASR 识别 + TTS 朗读，移植自数字分身）----
        self.speech_enabled = bool(self.cfg.get("voice", {}).get("speech_enabled", True))
        self.recorder = voice_mod.Recorder()
        self.audio_player = QMediaPlayer()
        self.audio_out = QAudioOutput()
        self.audio_player.setAudioOutput(self.audio_out)
        self._asr_thread = None
        self._tts_thread = None
        self._recording = False
        self.rec_wav = None
        self._init_ui()
        self.manager = QNetworkAccessManager()
        self.search_manager = QNetworkAccessManager()
        log.info("程序启动，模型=%s", cfg["model"])
        # v4.36 MCP 服务器异步初始化（不阻塞 UI 启动）
        # v4.59：初始化完成后更新状态栏显示 MCP 工具数
        def _init_mcp():
            config.init_mcp_clients(self.cfg)
            total_tools = sum(len(c.tools) for c in config.mcp_clients)
            if total_tools:
                names = ", ".join(c.name for c in config.mcp_clients)
                self.status_label.setText(f"🔌 MCP: {names} ({total_tools} 工具)")
        QTimer.singleShot(100, _init_mcp)

    # ============ UI 搭建 ============
    def _init_ui(self):
        self.setWindowTitle("小臭玩AI")
        self.setWindowIcon(get_app_icon())
        self._fit_window_to_screen()
        self.setMinimumSize(480, 520)  # 允许缩得更小，但避免窗口碎成不可用
        self.setWindowFlags(Qt.FramelessWindowHint)

        # 全局样式
        self.setStyleSheet(f"""
            QMainWindow {{ background-color: {THEME['bg']}; }}
            /* ---- ???? QSS (Codex ??) ---- */
            QWidget {{
                font-family: -apple-system, "Segoe UI Variable", "Segoe UI", "Microsoft YaHei", system-ui, sans-serif;
                font-size: 13px;
            }}
            QWidget::selection {{
                background: rgba(26,115,232,0.20);
                color: {THEME['text']};
            }}
            QToolTip {{
                background: {THEME['card']}; color: {THEME['text']};
                border: 1px solid {THEME['border']}; border-radius: 8px;
                padding: 6px 10px; font-size: 12px;
            }}
            QScrollBar:vertical {{
                background: transparent; width: 6px; margin: 0; border-radius: 3px;
            }}
            QScrollBar::handle:vertical {{
                background: rgba(148,163,184,0.25); min-height: 40px; border-radius: 4px;
            }}
            QScrollBar::handle:vertical:hover {{
                background: rgba(148,163,184,0.45);
            }}
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
                height: 0;
            }}
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {{
                background: transparent;
            }}
            QScrollBar:horizontal {{
                background: transparent; height: 6px; margin: 0; border-radius: 3px;
            }}
            QScrollBar::handle:horizontal {{
                background: rgba(148,163,184,0.25); min-width: 40px; border-radius: 4px;
            }}
            QScrollBar::handle:horizontal:hover {{
                background: rgba(148,163,184,0.45);
            }}
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{
                width: 0;
            }}
        """)

        central = QWidget()
        central.setStyleSheet(f"background:{THEME['bg']};")
        self.setCentralWidget(central)

        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # ===== 自定义标题栏（48）=====
        self._build_title_bar()
        main_layout.addWidget(self.title_bar)

        # ===== 内容区域：侧栏(256) + 右区 =====
        content = QWidget()
        content.setStyleSheet(f"background:{THEME['bg']};")
        outer = QHBoxLayout(content)
        outer.setContentsMargins(0, 0, 0, 0)
        outer.setSpacing(0)

        self._build_sidebar()
        outer.addWidget(self.sidebar)

        # ===== 右区：主栈(1) + 状态条(24) =====
        right_area = QWidget()
        right_area.setStyleSheet(f"background:{THEME['bg']};")
        ra_lay = QVBoxLayout(right_area)
        ra_lay.setContentsMargins(0, 0, 0, 0)
        ra_lay.setSpacing(0)

        self.main_stack = QStackedWidget()
        self.main_stack.setStyleSheet(f"background:{THEME['bg']};")
        ra_lay.addWidget(self.main_stack, 1)

        # 页0：首页
        self.welcome_page = QWidget()
        self.welcome_page.setStyleSheet(f"background:{THEME['bg']};")
        self._build_welcome_page()
        self.main_stack.addWidget(self.welcome_page)

        # 页1：对话
        self.chat_page = QWidget()
        self.chat_page.setStyleSheet(f"background:{THEME['bg']};")
        self._build_chat_page()
        self.main_stack.addWidget(self.chat_page)

        # 页2：编排
        self.orchestrate_page = QWidget()
        self.orchestrate_page.setStyleSheet(f"background:{THEME['bg']};")
        self._build_orchestrate_page()
        self.main_stack.addWidget(self.orchestrate_page)

        # 页3：生图
        self.image_page = QWidget()
        self.image_page.setStyleSheet(f"background:{THEME['bg']};")
        self._build_image_page()
        self.main_stack.addWidget(self.image_page)

        # 页4：生视频
        self.video_page = QWidget()
        self.video_page.setStyleSheet(f"background:{THEME['bg']};")
        self._build_video_page()
        self.main_stack.addWidget(self.video_page)

        # 页5：数字人分身（整合工作台之一）
        self.twin_page = QWidget()
        self.twin_page.setStyleSheet(f"background:{THEME['bg']};")
        self._build_twin_page()
        self.main_stack.addWidget(self.twin_page)

        # 页6：导演台（整合工作台之二）
        self.director_page = QWidget()
        self.director_page.setStyleSheet(f"background:{THEME['bg']};")
        self._build_director_page()
        self.main_stack.addWidget(self.director_page)

        # 页7：工具
        self.tools_page = QWidget()
        self.tools_page.setStyleSheet(f"background:{THEME['bg']};")
        self._build_tools_page()
        self.main_stack.addWidget(self.tools_page)

        # 页8：自动化任务
        self.automation_page = QWidget()
        self.automation_page.setStyleSheet(f"background:{THEME['bg']};")
        self._build_automation_page()
        self.main_stack.addWidget(self.automation_page)

        # 页9：设置
        self.settings_page = QWidget()
        self.settings_page.setStyleSheet(f"background:{THEME['bg']};")
        self._build_settings_page()
        self.main_stack.addWidget(self.settings_page)

        self.main_stack.setCurrentIndex(0)

        # ===== 状态条(24) =====
        self._build_status_bar()
        ra_lay.addWidget(self.status_bar)

        outer.addWidget(right_area, 1)
        main_layout.addWidget(content, 1)

        # ===== 弹层（保留，工具/设置旧入口）=====
        self._skill_popup, skill_inner_layout = self._build_skill_popup()
        self.skill_inner_layout = skill_inner_layout  # v4.84：审核通过后热重载技能库用
        self._settings_popup = self._build_settings_popup()
        self.skill_buttons = {}
        self._populate_skill_lib(skill_inner_layout)

        # 模型选择
        profiles = self.cfg.get("model_profiles", {})
        if profiles:
            names = list(profiles.keys())
            self.model_combo.addItems(names)
            self.model_combo.currentTextChanged.connect(self._on_model_change)
            for i, nm in enumerate(names):
                p = profiles[nm]
                if p.get("base_url") == self.cfg["base_url"] and p.get("model") == self.cfg["model"]:
                    self.model_combo.blockSignals(True)
                    self.model_combo.setCurrentIndex(i)
                    self.model_combo.blockSignals(False)
                    self._on_model_change(nm)
                    break
            else:
                if names:
                    self._on_model_change(names[0])

        # 初始化
        self._render_messages(force_bottom=True)
        self._refresh_recent_on_welcome()
        self._refresh_skill_buttons()
        self._update_skill_bar()
        self._refresh_deliverables()
        self._refresh_session_combo()
        self.input_box.setFocus()
        self._scan_agent_resume()  # v4.101：启动时检测本会话是否有暂停的 Agent 任务可继续

        # ===== 模块1：剪贴板自动监听 =====
        from clipboard_monitor import ClipboardMonitor
        self.clipboard_monitor = ClipboardMonitor(self.cfg)
        self.clipboard_monitor.clipboard_event.connect(self._on_clipboard_event)
        if self.cfg.get("clipboard_enabled", True):
            self.clipboard_monitor.start()

        # ---- 边缘缩放过滤器（用 QAbstractNativeEventFilter 避免 nativeEvent 签名冲突）----
        self._edge_filter = _EdgeResizeFilter(self)
        QApplication.instance().installNativeEventFilter(self._edge_filter)

    def _fit_window_to_screen(self):
        """初始尺寸自适应屏幕可用区域并居中，并限制在舒适范围内便于用户后续手动缩放。

        无边框窗口默认取屏幕可用区的 ~82%（宽上限 1280、高上限 820），既不写死满屏，
        也留出边缘便于用户拖拽四边/四角缩放（见 _EdgeResizeFilter）。
        """
        try:
            scr = QApplication.primaryScreen()
            if scr is not None:
                avail = scr.availableGeometry()
                aw, ah = avail.width(), avail.height()
                if aw > 0 and ah > 0:
                    w = min(1280, max(960, int(aw * 0.82)))
                    h = min(820, max(620, int(ah * 0.85)))
                    self.resize(w, h)
                    x = avail.x() + max(0, (aw - w) // 2)
                    y = avail.y() + max(0, (ah - h) // 2)
                    self.move(x, y)
                    return
        except Exception as e:
            log.warning("自适应屏幕尺寸失败，使用默认尺寸: %s", e)
        # 兜底：无屏幕信息或异常时回退到固定尺寸
        self.resize(1280, 820)

    # ============ 对话页（会话头 + 聊天视图 + 输入区 + 交付物）============
    def _build_chat_page(self):
        """对话页：左聊天列（会话头/视图/输入）+ 右交付物(260)。"""
        split = QSplitter(Qt.Horizontal)
        split.setStyleSheet("QSplitter{background:transparent;border:none;}")
        split.setHandleWidth(0)

        # ---- 左：聊天列 ----
        chat_col = QWidget()
        chat_col.setStyleSheet(f"background:{THEME['bg']};")
        cc_lay = QVBoxLayout(chat_col)
        cc_lay.setContentsMargins(0, 0, 0, 0)
        cc_lay.setSpacing(0)

        # 会话头
        header = QWidget()
        header.setFixedHeight(48)
        header.setStyleSheet(f"background:{THEME['bg']};border-bottom:1px solid {THEME['border']};")
        hb = QHBoxLayout(header)
        hb.setContentsMargins(10, 0, 10, 0)
        hb.setSpacing(6)

        self.session_combo = QComboBox()
        self.session_combo.setFixedHeight(32)
        self.session_combo.setMinimumWidth(160)
        self.session_combo.setStyleSheet(
            f"QComboBox{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 10px;font-size:13px;color:{THEME['text']};}}"
            f"QComboBox:hover{{border-color:{THEME['border_hover']};}}"
            f"QComboBox:focus{{border:1px solid {THEME['accent']};}}"
            f"QComboBox::drop-down{{border:none;width:20px;}}")
        self.session_combo.currentIndexChanged.connect(self._on_session_combo)
        hb.addWidget(self.session_combo)

        new_btn = QPushButton("＋")
        new_btn.setFixedSize(32, 32)
        new_btn.setToolTip("新建对话")
        new_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['accent']};color:#FFFFFF;border:none;"
            f"border-radius:8px;padding:0;font-size:15px;font-weight:500;}}"
            f"QPushButton:hover{{background:{THEME['accent_hover']};}}")
        new_btn.clicked.connect(self._new_session)
        hb.addWidget(new_btn)

        rename_btn = QPushButton("✏️")
        rename_btn.setFixedSize(32, 32)
        rename_btn.setToolTip("重命名当前对话")
        rename_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['card']};color:{THEME['text']};"
            f"border:1px solid {THEME['border']};border-radius:8px;padding:0;font-size:13px;}}"
            f"QPushButton:hover{{border-color:{THEME['accent']};}}")
        rename_btn.clicked.connect(self._rename_session)
        hb.addWidget(rename_btn)

        del_btn = QPushButton("🗑")
        del_btn.setFixedSize(32, 32)
        del_btn.setToolTip("删除当前对话")
        del_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['card']};color:{THEME['danger']};"
            f"border:1px solid {THEME['border']};border-radius:8px;padding:0;font-size:13px;}}"
            f"QPushButton:hover{{border-color:{THEME['danger']};background:#FCE8E6;}}")
        del_btn.clicked.connect(self._delete_active_session)
        hb.addWidget(del_btn)

        mgr_btn = QPushButton("📋")
        mgr_btn.setFixedSize(32, 32)
        mgr_btn.setToolTip("会话管理")
        mgr_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['card']};color:{THEME['text']};"
            f"border:1px solid {THEME['border']};border-radius:8px;padding:0;font-size:13px;}}"
            f"QPushButton:hover{{border-color:{THEME['accent']};}}")
        mgr_btn.clicked.connect(self._open_session_manager)
        hb.addWidget(mgr_btn)

        # 麦克风选择
        self.mic_combo = QComboBox()
        self.mic_combo.setFixedHeight(32)
        self.mic_combo.setMinimumWidth(150)
        self.mic_combo.setStyleSheet(
            f"QComboBox{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 8px;font-size:12px;color:{THEME['dim']};}}"
            f"QComboBox:hover{{border-color:{THEME['border_hover']};}}")
        # v4.96：麦克风选择从 header 移除，节省空间（实例保留供 _fill_mics/_get_selected_mic 使用）
        # hb.addWidget(self.mic_combo)

        # 刷新麦克风按钮（实例保留，不在 header 显示）
        self.mic_refresh_btn = QPushButton("🔄")
        self.mic_refresh_btn.setFixedHeight(32)
        self.mic_refresh_btn.setFixedWidth(32)
        self.mic_refresh_btn.setToolTip("刷新麦克风列表")
        self.mic_refresh_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;color:{THEME['dim']};font-size:13px;}}"
            f"QPushButton:hover{{border-color:{THEME['border_hover']};color:{THEME['text']};}}")
        self.mic_refresh_btn.clicked.connect(self._fill_mics)
        # hb.addWidget(self.mic_refresh_btn)

        # 语音开关（🔊：开/关，控制回复朗读）
        self.speech_chk = QPushButton("🔊")
        self.speech_chk.setFixedSize(32, 32)
        self.speech_chk.setToolTip("语音朗读开关")
        self.speech_chk.setCheckable(True)
        self.speech_chk.setChecked(self.speech_enabled)
        self.speech_chk.clicked.connect(self.on_toggle_speech)
        hb.addWidget(self.speech_chk)
        self._style_voice_switch(self.speech_chk, self.speech_enabled)

        # 技能管理器按钮（🧩）：呼出可视化技能面板
        self.skill_mgr_btn = QPushButton("🧩")
        self.skill_mgr_btn.setFixedSize(32, 32)
        self.skill_mgr_btn.setToolTip("技能管理器（Ctrl+Alt+S）")
        self.skill_mgr_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['card']};color:{THEME['text']};"
            f"border:1px solid {THEME['border']};border-radius:8px;padding:0;font-size:13px;}}"
            f"QPushButton:hover{{border-color:{THEME['accent']};}}")
        self.skill_mgr_btn.clicked.connect(self._open_skill_manager)
        hb.addWidget(self.skill_mgr_btn)

        # 技能市场按钮（🛍）：呼出技能市场发现 UI
        self.market_btn = QPushButton("🛍")
        self.market_btn.setFixedSize(32, 32)
        self.market_btn.setToolTip("技能市场（发现并安装新技能）")
        self.market_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['card']};color:{THEME['text']};"
            f"border:1px solid {THEME['border']};border-radius:8px;padding:0;font-size:13px;}}"
            f"QPushButton:hover{{border-color:{THEME['accent']};}}")
        self.market_btn.clicked.connect(self._open_skill_market)
        hb.addWidget(self.market_btn)

        # 工作流模板按钮（⚙️）：呼出可视化工作流面板
        self.wf_mgr_btn = QPushButton("⚙️")
        self.wf_mgr_btn.setFixedSize(32, 32)
        self.wf_mgr_btn.setToolTip("工作流模板（Ctrl+Alt+W）")
        self.wf_mgr_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['card']};color:{THEME['text']};"
            f"border:1px solid {THEME['border']};border-radius:8px;padding:0;font-size:13px;}}"
            f"QPushButton:hover{{border-color:{THEME['accent']};}}")
        self.wf_mgr_btn.clicked.connect(self._open_workflow_manager)
        hb.addWidget(self.wf_mgr_btn)

        hb.addStretch(1)

        # 导出按钮（📤）：呼出富格式导出（MD/HTML/PDF/DOCX）
        self.export_btn = QPushButton("📤")
        self.export_btn.setFixedSize(32, 32)
        self.export_btn.setToolTip("导出当前对话（Markdown/HTML/PDF/Word）")
        self.export_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['card']};color:{THEME['text']};"
            f"border:1px solid {THEME['border']};border-radius:8px;padding:0;font-size:13px;}}"
            f"QPushButton:hover{{border-color:{THEME['accent']};}}")
        self.export_btn.clicked.connect(self.export_session)
        hb.addWidget(self.export_btn)

        self.dv_expand_btn = QPushButton("⟨ 交付物")
        self.dv_expand_btn.setFixedHeight(32)
        self.dv_expand_btn.setVisible(False)
        self.dv_expand_btn.setStyleSheet(
            f"QPushButton{{background:transparent;border:1px solid {THEME['border']};"
            f"border-radius:8px;color:{THEME['dim']};font-size:12px;padding:0 10px;}}"
            f"QPushButton:hover{{background:{THEME['panel2']};color:{THEME['text']};}}")
        self.dv_expand_btn.clicked.connect(self._toggle_deliverables)
        hb.addWidget(self.dv_expand_btn)

        cc_lay.addWidget(header)

        # 聊天视图（v4.104：QWebEngineView 真浏览器渲染——圆角/Markdown/流式局部更新）
        self.chat_view = chat_web.ChatWebView(THEME)
        # v4.75：对话内搜索跳转 + 单条「重新生成 / 改写问题」链接（app:// 协议拦截）
        self.chat_view.anchorActivated.connect(self._on_anchor_clicked)
        # v4.104：页面意外重载（DOM 清空）→ 全量重渲染自愈
        self.chat_view.pageReloaded.connect(self._on_chat_page_reloaded)
        cc_lay.addWidget(self.chat_view, 1)

        # 输入区
        cc_lay.addWidget(self._build_input_area())

        # v4.109：模型下拉填充（必须在 _build_input_area 建好 combo 之后）
        self._fill_model_combo()

        # 麦克风列表填充（必须在 status_label 建好之后，否则 __init__ 崩溃）
        self._fill_mics()

        split.addWidget(chat_col)
        split.setStretchFactor(0, 1)

        # ---- 右：交付物 ----
        self._build_deliverables()
        split.addWidget(self.deliverables)
        split.setStretchFactor(1, 0)
        split.setSizes([980, 260])

        chat_layout = QVBoxLayout(self.chat_page)
        chat_layout.setContentsMargins(0, 0, 0, 0)
        chat_layout.setSpacing(0)
        chat_layout.addWidget(split)

    # ---------- v4.109 模型选择下拉（Auto 智能路由 / 手动锁定档位） ----------
    def _fill_model_combo(self):
        """填充下拉：Auto + 主模型 + 各 model_profiles 档位。

        - Auto（默认）：走原有智能路由，行为与 v4.108 完全一致。
        - 主模型：锁定 cfg 的 base_url/model/api_key。
        - 各 profile：锁定该档位；未填 api_key 的档位列出但**禁用**，
          避免选中后静默失败（也能让大哥看到"去设置里填 key 就能用"）。
        选择持久化到 config.json 的 model_lock，重启保持。
        """
        # 注意：必须叫 chat_model_combo。设置弹层里还有一个同名的 self.model_combo
        # （_build_settings_popup 内创建，__init__ 在其后执行会覆盖本引用），
        # 复用会导致输入区下拉变孤儿控件、且路由读到设置弹层的值。
        if not hasattr(self, "chat_model_combo"):
            return
        cfg = self.cfg or {}
        self.chat_model_combo.blockSignals(True)
        try:
            self.chat_model_combo.clear()
            self.chat_model_combo.addItem("Auto · 智能路由", "")
            self.chat_model_combo.addItem("主模型 · %s" % (cfg.get("model") or "未设置"),
                                          "__main__")
            for name, prof in (cfg.get("model_profiles") or {}).items():
                prof = prof or {}
                ok = bool(prof.get("api_key")) and bool(prof.get("model")) and bool(prof.get("base_url"))
                label = "%s · %s" % (name, prof.get("model") or "未设置")
                if not ok:
                    label += "（未配置 Key）"
                idx = self.chat_model_combo.addItem(label, name)
                if not ok:
                    try:
                        self.chat_model_combo.model().item(idx).setEnabled(False)
                    except Exception:
                        pass
            # 恢复上次选择（锁定档位若已不存在则安全回落 Auto）
            saved = cfg.get("model_lock", "") or ""
            target = 0
            for i in range(self.chat_model_combo.count()):
                if self.chat_model_combo.itemData(i) == saved:
                    target = i
                    break
            self.chat_model_combo.setCurrentIndex(target)
            self._model_lock = self.chat_model_combo.itemData(target) or ""
        finally:
            self.chat_model_combo.blockSignals(False)

    def _on_model_combo_changed(self, idx):
        """下拉切换：写入 self._model_lock 并持久化。Auto 时清空锁定。"""
        try:
            val = self.chat_model_combo.itemData(idx)
        except Exception:
            return
        val = val or ""
        self._model_lock = val
        try:
            self.cfg["model_lock"] = val
            import config as _cfg
            _cfg.save_config(self.cfg)
        except Exception as e:
            log.warning("保存 model_lock 失败: %s", e)
        try:
            name = (self.chat_model_combo.itemText(idx) or "").strip()
            if val:
                self.status_label.setText("已锁定模型：%s（Auto 路由已停用）" % name)
            else:
                self.status_label.setText("模型：Auto 智能路由")
        except Exception:
            pass

    def _build_input_area(self):
        """输入区卡片：状态/技能条 + 附件 + 输入框 + 发送。"""
        input_area = QWidget()
        input_area.setStyleSheet(f"background:{THEME['bg']};")
        ia_lay = QVBoxLayout(input_area)
        ia_lay.setContentsMargins(24, 8, 24, 20)
        ia_lay.setSpacing(8)

        # v4.109：状态行 = 左状态文字 + 右模型选择（Auto 智能路由 / 手动锁定档位）
        _top_row = QWidget()
        _top_lay = QHBoxLayout(_top_row)
        _top_lay.setContentsMargins(0, 0, 0, 0)
        _top_lay.setSpacing(8)
        self.status_label = QLabel("")
        self.status_label.setStyleSheet(
            f"color:{THEME['faint']};font-size:11px;min-height:12px;")
        _top_lay.addWidget(self.status_label, 1)
        self.chat_model_combo = QComboBox()
        self.chat_model_combo.setFixedHeight(24)
        self.chat_model_combo.setMinimumWidth(150)
        self.chat_model_combo.setMaximumWidth(280)
        self.chat_model_combo.setToolTip(
            "Auto = 智能路由（默认：轻量走 Agnes，工具/复杂任务自动升舱 DeepSeek）\n"
            "其他选项 = 手动锁定，本轮起全程只用该模型，不再自动切换")
        self.chat_model_combo.setStyleSheet(
            f"QComboBox{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:6px;padding:2px 6px;font-size:11px;color:{THEME['dim']};}}"
            f"QComboBox:hover{{border-color:{THEME['border_hover']};}}"
            f"QComboBox:focus{{border:1px solid {THEME['accent']};}}"
            f"QComboBox::drop-down{{border:none;width:18px;}}"
            f"QComboBox QAbstractItemView{{background:{THEME['card']};"
            f"border:1px solid {THEME['border']};"
            f"selection-background-color:{THEME['accent']};}}")
        self.chat_model_combo.currentIndexChanged.connect(self._on_model_combo_changed)
        _top_lay.addWidget(self.chat_model_combo, 0)
        ia_lay.addWidget(_top_row)

        self.skill_bar = QWidget()
        self.skill_bar.setVisible(False)
        skb_lay = QHBoxLayout(self.skill_bar)
        skb_lay.setContentsMargins(4, 0, 4, 0)
        skb_lay.setSpacing(8)
        self.skill_name_label = QLabel("")
        self.skill_name_label.setStyleSheet(
            f"color:{THEME['accent2']};font-size:12px;font-weight:600;")
        skb_lay.addWidget(self.skill_name_label)
        skb_lay.addStretch(1)
        skill_clear_btn = QPushButton("取消技能")
        skill_clear_btn.setStyleSheet(
            f"QPushButton{{background:transparent;border:none;color:{THEME['faint']};font-size:12px;}}"
            f"QPushButton:hover{{color:{THEME['danger']};}}")
        skill_clear_btn.clicked.connect(self._clear_skill)
        skb_lay.addWidget(skill_clear_btn)
        ia_lay.addWidget(self.skill_bar)

        input_frame = QFrame()
        input_frame.setStyleSheet(
            f"QFrame#inputCard{{"
            f"background:{THEME['panel2']};"
            f"border:1px solid {THEME['border']};"
            f"border-radius:10px;"
            f"padding:6px 8px;"
            f"margin:0;"
            f"}}"
            f"QFrame#inputCard:focus-within{{"
            f"border-color:{THEME['accent']};"
            f"border-width:1.5px;"
            f"}}")
        input_frame.setObjectName("inputCard")
        ifl = QHBoxLayout(input_frame)
        ifl.setContentsMargins(4, 4, 4, 4)
        ifl.setSpacing(8)

        attach_btn = QPushButton("+")
        attach_btn.setFixedSize(32, 32)
        attach_btn.setToolTip("选择文件（拖拽也支持）")
        attach_btn.setStyleSheet(
            f"QPushButton{{background:transparent;color:{THEME['dim']};"
            f"border:none;font-size:15px;}}"
            f"QPushButton:hover{{color:{THEME['text']};}}")
        attach_btn.clicked.connect(self._on_attach_file)
        ifl.addWidget(attach_btn, 0, Qt.AlignBottom)

        self.input_box = MultiLineInput()
        self.input_box.setPlaceholderText("输入自然语言指令 | Enter发送 | Shift+Enter换行")
        self.input_box.sendRequested.connect(self.send)
        self.input_box.setStyleSheet(
            f"QTextEdit{{background:transparent;border:none;padding:8px 4px 8px 14px;"
            f"font-size:14px;color:{THEME['text']};line-height:1.5;}}"
            f"QTextEdit:focus{{outline:none;}}")
        self.input_box.textChanged.connect(self._on_input_changed)
        ifl.addWidget(self.input_box, 1)

        # 按住说话（实时语音输入）
        self.talk_btn = QPushButton("🎤 按住说话")
        self.talk_btn.setFixedHeight(34)
        self.talk_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['card']};color:{THEME['dim']};"
            f"border:1px solid {THEME['border']};border-radius:6px;font-size:13px;}}"
            f"QPushButton:hover{{border-color:{THEME['border_hover']};color:{THEME['text']};}}"
            f"QPushButton:pressed{{background:{THEME['accent']};color:#FFFFFF;"
            f"border-color:{THEME['accent']};}}")
        self.talk_btn.pressed.connect(self._start_talk)
        self.talk_btn.released.connect(self._stop_talk)
        ifl.addWidget(self.talk_btn, 0, Qt.AlignBottom)

        self.send_btn = QPushButton("→")
        self.send_btn.setFixedSize(34, 34)
        self.send_btn.setEnabled(False)
        self.send_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['accent']};color:#FFFFFF;"
            f"border:none;border-radius:6px;font-size:16px;font-weight:600;}}"
            f"QPushButton:hover{{background:{THEME['accent_hover']};}}"
            f"QPushButton:disabled{{background:{THEME['border']};color:{THEME['faint']};}}")
        self.send_btn.clicked.connect(self.send)
        ifl.addWidget(self.send_btn, 0, Qt.AlignBottom)

        # 停止按钮（Agent 执行中可见，随时中止等待中的轮次/工具）
        self.stop_btn = QPushButton("⏹ 停止")
        self.stop_btn.setFixedHeight(34)
        self.stop_btn.setFixedWidth(60)
        self.stop_btn.setVisible(False)
        self.stop_btn.setStyleSheet(
            f"QPushButton{{background:{THEME.get('danger', '#d93025')};color:#FFFFFF;"
            f"border:none;border-radius:6px;font-size:13px;font-weight:600;}}"
            f"QPushButton:hover{{background:#b3261e;}}"
            f"QPushButton:disabled{{background:{THEME['border']};color:{THEME['faint']};}}")
        self.stop_btn.clicked.connect(self._request_agent_stop)
        ifl.addWidget(self.stop_btn, 0, Qt.AlignBottom)

        # v4.101：任务暂停后「继续上次任务」入口（普通 Agent 断点续传）
        self.resume_agent_btn = QPushButton("▶ 继续上次任务")
        self.resume_agent_btn.setFixedHeight(34)
        self.resume_agent_btn.setFixedWidth(120)
        self.resume_agent_btn.setVisible(False)
        self.resume_agent_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['accent']};color:#FFFFFF;"
            f"border:none;border-radius:6px;font-size:13px;font-weight:600;}}"
            f"QPushButton:hover{{background:{THEME['accent_hover']};}}")
        self.resume_agent_btn.clicked.connect(self._resume_agent_task)
        ifl.addWidget(self.resume_agent_btn, 0, Qt.AlignBottom)

        ia_lay.addWidget(input_frame)

        input_hint = QLabel("Enter 发送 · Shift+Enter 换行")
        input_hint.setStyleSheet(f"color:{THEME['placeholder']};font-size:11px;padding-left:4px;")
        ia_lay.addWidget(input_hint)

        return input_area

    # ============ 实时语音（移植自数字分身）============
    def _style_voice_switch(self, btn, on):
        """语音开关视觉态：开=蓝底白字，关=灰底弱字。"""
        if on:
            btn.setStyleSheet(
                f"QPushButton{{background:{THEME['accent']};color:#FFFFFF;border:none;"
                f"border-radius:8px;padding:0;font-size:15px;font-weight:600;}}")
        else:
            btn.setStyleSheet(
                f"QPushButton{{background:{THEME['card']};color:{THEME['dim']};"
                f"border:1px solid {THEME['border']};border-radius:8px;"
                f"padding:0;font-size:15px;font-weight:500;}}")

    def _fill_mics(self):
        import tempfile as _tf
        diag = {}
        try:
            mics = voice_mod.detect_mics(diag=diag)
        except Exception as e:
            mics = []
            diag["exception"] = f"{type(e).__name__}: {e}"
        self.mic_combo.clear()
        # 写诊断日志到 %TEMP%/xiaochou_voice.log 便于排查
        try:
            logp = _tf.path.join(_tf.gettempdir(), "xiaochou_voice.log")
            with open(logp, "a", encoding="utf-8") as _lf:
                import datetime as _dt
                _lf.write(f"\n[{_dt.datetime.now():%H:%M:%S}] _fill_mics\n")
                for k, v in diag.items():
                    _lf.write(f"  {k}: {v}\n")
                _lf.write(f"  result: {mics}\n")
        except Exception:
            pass
        if mics:
            self.mic_combo.addItems(mics)
            self.status_label.setText(f"麦克风就绪：{mics[0]}")
        else:
            self.mic_combo.addItem("未检测到麦克风")
            # 在状态栏显示失败原因，方便用户/排查
            why = diag.get("exception") or (
                "ffmpeg 不存在" if not diag.get("ffmpeg_exists") else
                f"ffmpeg 返回空列表(rc={diag.get('returncode')}, stderr={diag.get('stderr_len')}B)")
            self.status_label.setText(f"⚠ 未检测到麦克风：{why}")

    def _current_mic(self):
        return self.mic_combo.currentText() if self.mic_combo.currentIndex() >= 0 else None

    def on_toggle_speech(self, checked):
        self.speech_enabled = bool(checked)
        self._style_voice_switch(self.speech_chk, self.speech_enabled)
        self.cfg.setdefault("voice", {})["speech_enabled"] = self.speech_enabled
        self._save_cfg()
        if not self.speech_enabled:
            self.audio_player.stop()  # 立即消声
        self.status_label.setText("🔊 语音朗读：开" if self.speech_enabled else "🔇 语音朗读：关")

    def _open_skill_manager(self):
        """呼出可视化技能管理器面板（按钮 / 托盘菜单 / Ctrl+Alt+S 共用）"""
        try:
            from skill_manager_ui import open_skill_manager
            open_skill_manager(self.cfg)
        except Exception as e:
            self.status_label.setText(f"技能管理器打开失败：{e}")

    def _open_skill_market(self):
        """呼出技能市场发现 UI（按钮 / 托盘菜单共用）"""
        try:
            from skill_market_ui import open_skill_market
            open_skill_market(self.cfg)
        except Exception as e:
            self.status_label.setText(f"技能市场打开失败：{e}")

    def _open_workflow_manager(self):
        """呼出工作流模板面板（按钮 / 托盘菜单 / Ctrl+Alt+W 共用）"""
        try:
            from workflow_manager_ui import open_workflow_manager
            open_workflow_manager(self.cfg, self)
        except Exception as e:
            self.status_label.setText(f"工作流模板打开失败：{e}")

    def export_diagnostic(self):
        """导出调试诊断包（设置页按钮 / 托盘菜单共用）"""
        try:
            from diagnostic_export import export_diagnostic_package
            export_diagnostic_package(self)
        except Exception as e:
            self.status_label.setText(f"诊断包导出失败：{e}")

    def _start_talk(self):
        if self._busy:
            self.status_label.setText("上一条还在处理，稍等…")
            return
        if self._recording:
            return
        mic = self._current_mic()
        if not mic or mic == "未检测到麦克风":
            self.status_label.setText("没有可用麦克风，先选麦克风或直接打字")
            return
        self._recording = True
        try:
            self.rec_wav = self.recorder.start(mic=mic)
        except Exception as e:
            self._recording = False
            self.rec_wav = None
            self.status_label.setText(f"麦克风启动失败：{e}")
            return
        self.talk_btn.setText("🔴 聆听中…（松开结束）")
        self.status_label.setText(f"🔴 聆听中… 麦克风：{mic}")

    def _stop_talk(self):
        if self._busy:
            # 上一条还在处理：丢弃本次录音，避免叠加
            self.rec_wav = None
            self._recording = False
            self.talk_btn.setText("🎤 按住说话")
            return
        if not self._recording:
            return
        self._recording = False
        self.talk_btn.setText("🎤 按住说话")
        if not self.rec_wav:
            # 按下即松开 / 没真正开始录音：直接恢复，不报错
            return
        self.status_label.setText("⏳ 语音识别中…")
        wav = self.recorder.stop()
        self.rec_wav = None
        if not wav:
            self.status_label.setText(f"没录到声音：{self.recorder.error or '未知'}")
            return
        if self._asr_thread and self._asr_thread.isRunning():
            self._asr_thread.terminate()
        sf = self.cfg.get("siliconflow", {})
        self._asr_thread = _ASRWorker(wav, sf)
        self._asr_thread.sig_text.connect(self._on_asr_done)
        self._asr_thread.sig_error.connect(self._on_asr_error)
        self._asr_thread.start()

    def _on_asr_done(self, text):
        if not text:
            self.status_label.setText("没听清，再说一次 / 或直接打字")
            return
        self.input_box.setPlainText(text)
        self.send()

    def _on_asr_error(self, msg):
        self.status_label.setText(f"语音识别失败：{msg}")

    def _speak(self, text):
        """若语音开关开，把助手回复合成为语音并播放（后台线程合成，主线程播放）。"""
        if not self.speech_enabled:
            return
        if not text or not text.strip():
            return
        if self._tts_thread and self._tts_thread.isRunning():
            self._tts_thread.terminate()
        tts = self.cfg.get("tts", {})
        self._tts_thread = _TTSWorker(text, tts)
        self._tts_thread.sig_mp3.connect(self.play_audio)
        self._tts_thread.start()

    def play_audio(self, mp3):
        if not self.speech_enabled or not mp3 or not os.path.exists(mp3):
            return
        self.audio_player.setSource(QUrl.fromLocalFile(mp3))
        self.audio_player.play()

    def _notify_task_done(self, title, message):
        """v4.95 任务完成通知：托盘弹窗 + 语音播报（语音独立于对话语音开关，因是提醒非朗读）。"""
        try:
            tray = getattr(getattr(self, "tray_app", None), "tray", None)
            if tray is not None:
                tray.showMessage(title, message, QSystemTrayIcon.Information, 6000)
        except Exception as e:
            log.error("托盘通知失败: %s", e)
        try:
            tts = self.cfg.get("tts", {})
            self._tts_thread = _TTSWorker(message, tts)
            self._tts_thread.sig_mp3.connect(self._play_notify_audio)
            self._tts_thread.start()
        except Exception as e:
            log.error("语音提醒失败: %s", e)

    def _play_notify_audio(self, mp3):
        """播放提醒语音（不受对话语音开关限制）。"""
        if not mp3 or not os.path.exists(mp3):
            return
        try:
            self.audio_player.setSource(QUrl.fromLocalFile(mp3))
            self.audio_player.play()
        except Exception as e:
            log.error("播放提醒语音失败: %s", e)

    # ============ 状态条（24）============
    def _build_status_bar(self):
        self.status_bar = QWidget()
        self.status_bar.setFixedHeight(24)
        self.status_bar.setStyleSheet(
            f"background:{THEME['card']};border-top:1px solid {THEME['border']};")
        sb_lay = QHBoxLayout(self.status_bar)
        sb_lay.setContentsMargins(16, 0, 16, 0)
        sb_lay.setSpacing(12)
        self.conn_status = QLabel("● 已连接")
        self.conn_status.setStyleSheet(f"color:{THEME['ok']};font-size:11px;")
        sb_lay.addWidget(self.conn_status)
        sb_lay.addStretch(1)
        token_label = QLabel("Agnes 免费 · DeepSeek 已订阅")
        token_label.setStyleSheet(f"color:{THEME['faint']};font-size:11px;")
        sb_lay.addWidget(token_label)
        hint = QLabel("Enter 发送 · Shift+Enter 换行")
        hint.setStyleSheet(f"color:{THEME['faint']};font-size:11px;")
        sb_lay.addWidget(hint)

    # ============ 按钮样式 ============
    def _primary_btn_style(self):
        return (f"QPushButton{{background:{THEME['accent']};color:#FFFFFF;border:none;"
                f"border-radius:8px;padding:0 16px;font-size:13px;font-weight:600;}}"
                f"QPushButton:hover{{background:{THEME['accent_hover']};}}"
                f"QPushButton:disabled{{background:{THEME['accent_disabled']};color:#FFFFFF;}}")

    def _secondary_btn_style(self):
        return (f"QPushButton{{background:{THEME['card']};border:1px solid {THEME['border']};"
                f"border-radius:8px;padding:0 16px;font-size:13px;font-weight:500;color:{THEME['dim']};}}"
                f"QPushButton:hover{{background:{THEME['panel2']};color:{THEME['text']};"
                f"border-color:{THEME['border_highlight']};}}")

    # ============ 浏览器扩展辅助 ============
    def _copy_ext_token(self):
        tok = self.cfg.get("browser_bridge_token", "")
        if not tok:
            self.status_label.setText("配对码为空，请重启小臭以生成")
            return
        try:
            from PySide6.QtWidgets import QApplication
            QApplication.clipboard().setText(tok)
            self.status_label.setText("✅ 配对码已复制到剪贴板")
        except Exception as e:
            self.status_label.setText("复制失败：" + str(e))

    def _open_path(self, path):
        """用系统默认程序打开文件/目录。"""
        import subprocess
        try:
            os.startfile(path) if os.name == "nt" else subprocess.run(
                ["open" if os.uname().sysname == "Darwin" else "xdg-open", path])
        except Exception as e:
            self.status_label.setText("打开失败：" + str(e))

    # ============ 编排页（小说一条龙）============
    def _build_orchestrate_page(self):
        page = self.orchestrate_page
        lay = QVBoxLayout(page)
        lay.setContentsMargins(32, 24, 32, 24)
        lay.setSpacing(16)

        head = QLabel("小说一条龙 · 编排")
        head.setStyleSheet(f"font-size:20px;font-weight:700;color:{THEME['text']};")
        lay.addWidget(head)
        sub = QLabel("精简流水线（3Phase+2检查）：爆款雷达 → 选题验证 → 写手成稿 → 虚拟编辑审稿 → 终稿定稿"
                     "（节点间传递上文，结果写入下方日志）。短篇按「字数」一次性写满；长篇按章生成，"
                     "跑完点「续写下一章」出下一章。\n⚠️ 爆款雷达出 3 个切入点后【会停下等你选定】，运行中可随时点「暂停/插入意见」改方向。")
        sub.setStyleSheet(f"font-size:12px;color:{THEME['dim']};")
        lay.addWidget(sub)

        cfg_row = QHBoxLayout()
        self.orch_topic = QLineEdit()
        self.orch_topic.setPlaceholderText("输入小说主题，如：重生回到高考前")
        self.orch_topic.setFixedHeight(34)
        self.orch_topic.setStyleSheet(
            f"QLineEdit{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 12px;font-size:13px;color:{THEME['text']};}}"
            f"QLineEdit:focus{{border:1px solid {THEME['accent']};}}")
        cfg_row.addWidget(self.orch_topic, 1)

        self.orch_len_type = QComboBox()
        self.orch_len_type.addItems(["短篇", "长篇"])
        self.orch_len_type.setFixedHeight(34)
        self.orch_len_type.setToolTip("短篇：一次性写完目标字数；长篇：按章生成，用「续写」出下一章")
        self.orch_len_type.setStyleSheet(
            f"QComboBox{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 10px;font-size:13px;color:{THEME['text']};}}")
        cfg_row.addWidget(self.orch_len_type)

        self.orch_words = QSpinBox()
        self.orch_words.setRange(300, 200000)
        self.orch_words.setSingleStep(500)
        self.orch_words.setValue(2000)
        self.orch_words.setSuffix(" 字")
        self.orch_words.setFixedHeight(34)
        self.orch_words.setToolTip("短篇=全文目标字数；长篇=每章目标字数")
        self.orch_words.setStyleSheet(
            f"QSpinBox{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 10px;font-size:13px;color:{THEME['text']};}}"
            f"QSpinBox:focus{{border:1px solid {THEME['accent']};}}")
        cfg_row.addWidget(self.orch_words)

        self.orch_platform = QComboBox()
        self.orch_platform.addItems(["番茄小说", "知乎", "公众号", "抖音", "头条"])
        self.orch_platform.setFixedHeight(34)
        self.orch_platform.setStyleSheet(
            f"QComboBox{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 10px;font-size:13px;color:{THEME['text']};}}")
        cfg_row.addWidget(self.orch_platform)

        self.orch_run_btn = QPushButton("开始生成")
        self.orch_run_btn.setFixedHeight(34)
        self.orch_run_btn.setStyleSheet(self._primary_btn_style())
        self.orch_run_btn.clicked.connect(self._run_orchestrate)
        cfg_row.addWidget(self.orch_run_btn)

        self.orch_continue_btn = QPushButton("续写下一章")
        self.orch_continue_btn.setFixedHeight(34)
        self.orch_continue_btn.setStyleSheet(self._secondary_btn_style())
        self.orch_continue_btn.setEnabled(False)
        self.orch_continue_btn.setToolTip("基于已生成底稿继续写（长篇出下一章 / 短篇扩写）")
        self.orch_continue_btn.clicked.connect(self._continue_orchestrate)
        cfg_row.addWidget(self.orch_continue_btn)

        self.orch_pause_btn = QPushButton("暂停/插入意见")
        self.orch_pause_btn.setFixedHeight(34)
        self.orch_pause_btn.setStyleSheet(self._secondary_btn_style())
        self.orch_pause_btn.setEnabled(False)
        self.orch_pause_btn.setToolTip("运行中请求在下一个阶段边界暂停，可插入修改意见")
        self.orch_pause_btn.clicked.connect(self._request_orch_pause)
        cfg_row.addWidget(self.orch_pause_btn)

        clear_btn = QPushButton("清空")
        clear_btn.setFixedHeight(34)
        clear_btn.setStyleSheet(self._secondary_btn_style())
        clear_btn.clicked.connect(lambda: self.orch_log.clear())
        cfg_row.addWidget(clear_btn)
        lay.addLayout(cfg_row)

        node_colors = {"blue": THEME['accent'], "green": THEME['ok'],
                       "orange": THEME['warn'], "purple": THEME['accent2'],
                       "red": THEME['danger']}
        self.orch_node_status = []
        node_grid = QGridLayout()
        node_grid.setContentsMargins(0, 0, 0, 0)
        node_grid.setSpacing(10)
        for i, (name, color, brief) in enumerate(OrchestrateWorker.STAGES):
            card = QWidget()
            card.setStyleSheet(f"background:{THEME['card']};border:1px solid {THEME['border']};"
                               f"border-radius:10px;padding:10px 12px;")
            cl = QHBoxLayout(card)
            cl.setContentsMargins(10, 8, 10, 8)
            cl.setSpacing(10)
            dot = QLabel("●")
            dot.setFixedSize(14, 14)
            dot.setStyleSheet(f"color:{node_colors[color]};font-size:12px;background:transparent;")
            cl.addWidget(dot)
            tl = QLabel(name)
            tl.setStyleSheet(f"font-size:14px;font-weight:600;color:{THEME['text']};background:transparent;")
            cl.addWidget(tl)
            st = QLabel("待运行")
            st.setStyleSheet(f"font-size:12px;color:{THEME['faint']};background:transparent;")
            cl.addStretch(1)
            cl.addWidget(st)
            node_grid.addWidget(card, i // 3, i % 3)
            self.orch_node_status.append(st)
        lay.addLayout(node_grid)

        # —— 选题闸门：爆款雷达出选项后，展示让你选定（隐藏，需要时显示）——
        self.orch_choice_box = QGroupBox("① 爆款雷达已给出切入点，请选定（或自填）")
        self.orch_choice_box.setStyleSheet(
            f"QGroupBox{{background:{THEME['panel2']};border:1px solid {THEME['border_highlight']};"
            f"border-radius:10px;font-size:13px;font-weight:600;color:{THEME['text']};"
            f"padding:12px 14px;margin-top:8px;}}")
        cb_lay = QVBoxLayout(self.orch_choice_box)
        cb_lay.setSpacing(8)
        self.orch_choice_area = QWidget()
        self.orch_choice_area_lay = QVBoxLayout(self.orch_choice_area)
        self.orch_choice_area_lay.setContentsMargins(0, 0, 0, 0)
        self.orch_choice_area_lay.setSpacing(6)
        cb_lay.addWidget(self.orch_choice_area)
        custom_row = QHBoxLayout()
        self.orch_choice_custom = QLineEdit()
        self.orch_choice_custom.setPlaceholderText("都不满意？在此自填切入点方向，再点右侧按钮")
        self.orch_choice_custom.setFixedHeight(32)
        self.orch_choice_custom.setStyleSheet(
            f"QLineEdit{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 10px;font-size:13px;color:{THEME['text']};}}")
        self.orch_choice_custom_btn = QPushButton("用我的输入")
        self.orch_choice_custom_btn.setFixedHeight(32)
        self.orch_choice_custom_btn.setStyleSheet(self._primary_btn_style())
        self.orch_choice_custom_btn.clicked.connect(self._on_orch_choice_custom)
        custom_row.addWidget(self.orch_choice_custom, 1)
        custom_row.addWidget(self.orch_choice_custom_btn)
        cb_lay.addLayout(custom_row)
        self.orch_choice_box.setVisible(False)
        lay.addWidget(self.orch_choice_box)

        # —— 中途暂停：插入修改意见的反馈框（隐藏，暂停时显示）——
        self.orch_pause_box = QGroupBox("已暂停 · 可插入修改意见，再继续")
        self.orch_pause_box.setStyleSheet(
            f"QGroupBox{{background:{THEME['panel2']};border:1px solid {THEME['warn']};"
            f"border-radius:10px;font-size:13px;font-weight:600;color:{THEME['text']};"
            f"padding:12px 14px;margin-top:8px;}}")
        pb_lay = QVBoxLayout(self.orch_pause_box)
        pb_lay.setSpacing(8)
        self.orch_pause_hint = QLabel("")
        self.orch_pause_hint.setStyleSheet(f"font-size:12px;color:{THEME['dim']};")
        pb_lay.addWidget(self.orch_pause_hint)
        pf_row = QHBoxLayout()
        self.orch_feedback = QLineEdit()
        self.orch_feedback.setPlaceholderText("例如：第二主角动机再强点 / 换个大女主 / 节奏放慢…（留空则直接继续）")
        self.orch_feedback.setFixedHeight(32)
        self.orch_feedback.setStyleSheet(
            f"QLineEdit{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 10px;font-size:13px;color:{THEME['text']};}}")
        self.orch_resume_btn = QPushButton("继续 ▶")
        self.orch_resume_btn.setFixedHeight(32)
        self.orch_resume_btn.setStyleSheet(self._primary_btn_style())
        self.orch_resume_btn.clicked.connect(self._on_orch_resume)
        pf_row.addWidget(self.orch_feedback, 1)
        pf_row.addWidget(self.orch_resume_btn)
        pb_lay.addLayout(pf_row)
        self.orch_pause_box.setVisible(False)
        lay.addWidget(self.orch_pause_box)

        self.orch_log = QTextEdit()
        self.orch_log.setReadOnly(True)
        self.orch_log.setStyleSheet(
            f"QTextEdit{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:10px;padding:12px;font-size:13px;line-height:1.6;color:{THEME['text']};}}")
        lay.addWidget(self.orch_log, 1)

        # 断点恢复横幅（崩溃/强杀遗留的检查点，启动时提示续跑）
        self._orch_resume_banner = None
        self._scan_orch_resume()
        # 技能审核队列扫描（模型自动创建的技能待人工通过才生效）
        self._scan_skill_review()

    def _run_orchestrate(self):
        topic = self.orch_topic.text().strip()
        if not topic:
            self.orch_log.setPlainText("请先填写小说主题。")
            return
        if getattr(self, "_orch_running", False):
            return
        self._orch_state = None  # 新小说，清空续写上下文
        worker = OrchestrateWorker(
            self, topic, self.orch_platform.currentText(),
            self.orch_len_type.currentText(), self.orch_words.value())
        self._start_orch_worker(worker)

    def _start_orch_worker(self, worker):
        """统一的 worker 接线 + 启动（开始生成 / 续写 / 断点恢复 共用）。"""
        self._orch_running = True
        self.orch_continue_btn.setEnabled(False)
        self.orch_run_btn.setText("取消生成")
        try:
            self.orch_run_btn.clicked.disconnect()
        except Exception:
            pass
        self.orch_run_btn.clicked.connect(self._cancel_orchestrate)
        self._orch_worker = worker
        self._orch_worker.stage.connect(lambda i, n: self.orch_node_status[i].setText(n))
        self._orch_worker.node_status.connect(self._on_orch_node_status)
        self._orch_worker.log.connect(lambda t: self.orch_log.insertPlainText(t))
        self._orch_worker.done.connect(self._on_orch_done)
        self._orch_worker.need_choice.connect(self._on_orch_need_choice)
        self._orch_worker.paused.connect(self._on_orch_paused)
        self.orch_pause_btn.setEnabled(True)
        self.orch_choice_box.setVisible(False)
        self.orch_pause_box.setVisible(False)
        self._orch_worker.start()

    def _on_orch_node_status(self, i, status):
        colors = {"running": THEME['tool_running'], "done": THEME['ok'], "error": THEME['danger']}
        text = {"running": "运行中…", "done": "完成", "error": "失败"}
        self.orch_node_status[i].setText(text.get(status, status))
        self.orch_node_status[i].setStyleSheet(
            f"font-size:12px;color:{colors.get(status, THEME['faint'])};background:transparent;")

    def _on_orch_done(self, msg):
        self._orch_running = False
        # v4.95：小说编排成功完成 → 托盘+语音通知
        if msg and "生成完成" in str(msg):
            self._notify_task_done("生成完成", "小说一条龙已生成完成")
        try:
            self.orch_run_btn.clicked.disconnect()
        except Exception:
            pass
        self.orch_run_btn.clicked.connect(self._run_orchestrate)
        self.orch_run_btn.setText("开始生成")
        self.orch_pause_btn.setEnabled(False)
        self.orch_choice_box.setVisible(False)
        self.orch_pause_box.setVisible(False)
        # 保存底稿上下文，供「续写」使用
        st = getattr(self._orch_worker, "final_state", None)
        if st:
            self._orch_state = st
            self.orch_continue_btn.setEnabled(True)
            self.orch_continue_btn.setText(
                "续写下一章" if st["length_type"] == "长篇" else "续写/扩写")
        # v4.101 断点续传：正常完成 → 删除检查点；用户取消 → 标记 paused 保留（可继续）
        try:
            tid = getattr(self._orch_worker, "task_id", None)
            if tid:
                if getattr(self._orch_worker, "_cancelled", False):
                    task_resume.mark_paused(self.cfg, tid)
                    self.orch_log.insertPlainText(
                        "⏸ 已取消（检查点已保留，重开 APP 或再次进入本页可继续，也可在横幅里丢弃）\n")
                else:
                    task_resume.mark_done(self.cfg, tid)
        except Exception:
            pass
        # D 项（轨迹记忆）：成功跑通则记录轨迹；并按配置（默认关）保守自动提炼经验库
        try:
            import trace_log
            w = self._orch_worker
            if w and msg and "生成完成" in msg:
                trace_log.capture_trace(self.cfg, {
                    "topic": w.topic,
                    "platform": w.platform,
                    "length_type": w.length_type,
                    "target_words": w.target_words,
                    "chosen_direction": (w._chosen_direction or "")[:300],
                    "chapter": w.chapter,
                    "stages": [s[0] for s in w.STAGES],
                    "stage_durations": getattr(w, "_stage_durations", {}),
                    "retry_count": getattr(w, "_retry_count", 0),
                    "final_words": len(w.full_draft or ""),
                    "outcome": "success",
                })
                if self.cfg.get("orch_trace_auto_refine", False):
                    trace_log.auto_refine_harness(
                        self.cfg, min_count=self.cfg.get("orch_trace_min_count", 3))
                trace_log.prune(self.cfg)
        except Exception:
            pass
        self.orch_log.insertPlainText(f"\n✅ {msg}\n")

    def _continue_orchestrate(self):
        """基于已生成底稿继续写：长篇出下一章，短篇扩写。"""
        if getattr(self, "_orch_running", False):
            return
        st = getattr(self, "_orch_state", None)
        if not st:
            self.orch_log.insertPlainText("⚠️ 还没有可续写的底稿，请先点「开始生成」。\n")
            return
        worker = OrchestrateWorker(
            self, st["topic"], st["platform"], st["length_type"],
            st["target_words"], prev_state=st)
        self._start_orch_worker(worker)

    def _resume_from_checkpoint(self, task_id):
        """断点恢复：从崩溃/强杀遗留的检查点继续跑（reattach）。"""
        cp = task_resume.load_checkpoint(self.cfg, task_id)
        if not cp:
            return
        if getattr(self, "_orch_running", False):
            return
        start = int(cp.get("stage", 0)) + 1
        worker = OrchestrateWorker(
            self, cp.get("topic", ""), cp.get("platform", "番茄小说"),
            cp.get("length_type", "短篇"), cp.get("target_words", 2000),
            prev_state=cp, task_id=task_id, start_stage=start)
        self._clear_resume_banner()
        self._start_orch_worker(worker)

    def _scan_orch_resume(self):
        """启动/进入编排页时扫描未完成的长任务检查点。

        若开启「自动续跑」(orch_auto_resume，默认 True)：重开 APP 即自动从断点继续，
        无需手动点「继续」（借鉴 Prime-Agent daemon 自动 reattach）。
        否则：弹「继续/丢弃」横幅等用户确认。
        """
        try:
            active = [c for c in task_resume.list_active(self.cfg)
                      if c.get("task_type") in (None, "orchestrate")]
        except Exception:
            return
        if not active:
            return
        if self.cfg.get("orch_auto_resume", True):
            tid = active[0].get("_task_id")
            self.orch_log.insertPlainText(
                "🔄 已自动恢复上次未完成的『小说一条龙』任务，从断点继续…\n")
            self._resume_from_checkpoint(tid)
        else:
            self._show_resume_banner(active[0])

    def _show_resume_banner(self, cp):
        self._clear_resume_banner()
        stage_idx = int(cp.get("stage", 0))
        stage_name = "未知"
        try:
            stage_name = OrchestrateWorker.STAGES[stage_idx][0]
        except Exception:
            pass
        updated = cp.get("updated", "")
        banner = QGroupBox("⏯ 检测到上次『小说一条龙』任务未跑完")
        banner.setStyleSheet(
            f"QGroupBox{{background:{THEME['panel2']};border:1px solid {THEME['accent']};"
            f"border-radius:10px;font-size:13px;font-weight:600;color:{THEME['text']};"
            f"padding:10px 14px;margin-top:8px;}}")
        bl = QHBoxLayout(banner)
        bl.setContentsMargins(10, 6, 10, 6)
        bl.setSpacing(10)
        info = QLabel(f"上次进行到「{stage_name}」阶段（{updated}），是否从断点继续？"
                      f"（已完成的阶段不会重跑）")
        info.setStyleSheet(f"font-size:12px;color:{THEME['text']};background:transparent;")
        bl.addWidget(info, 1)
        auto_cb = QCheckBox("总是自动续跑")
        auto_cb.setChecked(bool(self.cfg.get("orch_auto_resume", True)))
        auto_cb.setStyleSheet(f"font-size:12px;color:{THEME['text']};background:transparent;")
        auto_cb.setToolTip("开启后，崩溃/强杀重开 APP 会自动从断点继续，不再弹此确认")
        auto_cb.toggled.connect(lambda on: self._set_auto_resume(bool(on)))
        bl.addWidget(auto_cb)
        tid = cp.get("_task_id")
        resume_btn = QPushButton("继续 ▶")
        resume_btn.setFixedHeight(30)
        resume_btn.setStyleSheet(self._primary_btn_style())
        resume_btn.clicked.connect(lambda: self._resume_from_checkpoint(tid))
        bl.addWidget(resume_btn)
        discard_btn = QPushButton("丢弃")
        discard_btn.setFixedHeight(30)
        discard_btn.setStyleSheet(self._secondary_btn_style())
        discard_btn.clicked.connect(lambda: self._discard_resume(tid))
        bl.addWidget(discard_btn)
        self._orch_resume_banner = banner
        try:
            lay = self.orchestrate_page.layout()
            lay.insertWidget(1, banner)
        except Exception:
            pass

    def _discard_resume(self, task_id):
        try:
            task_resume.mark_done(self.cfg, task_id)
        except Exception:
            pass
        self._clear_resume_banner()

    def _set_auto_resume(self, on):
        """横幅勾选「总是自动续跑」→ 写入并持久化配置。"""
        try:
            self.cfg["orch_auto_resume"] = on
            config.save_config(self.cfg)
        except Exception:
            pass

    def _clear_resume_banner(self):
        b = getattr(self, "_orch_resume_banner", None)
        if b is not None:
            try:
                b.setParent(None)
                b.deleteLater()
            except Exception:
                pass
            self._orch_resume_banner = None

    def _cancel_orchestrate(self):
        """取消编排：唤醒所有挂起等待并请求 worker 停止。
        v4.101：取消 ≠ 丢弃——保留检查点（标记 paused），重开 APP/进编排页可继续。"""
        if getattr(self, "_orch_worker", None):
            self._orch_worker._cancelled = True
            self._orch_worker.release_locks()
        self.orch_pause_btn.setEnabled(False)
        self.orch_choice_box.setVisible(False)
        self.orch_pause_box.setVisible(False)
        self.orch_log.insertPlainText("\n⏹ 正在取消…\n")

    # ---- 选题闸门 / 中途暂停 交互 ----
    def _request_orch_pause(self):
        """请求 worker 在下一个阶段边界暂停，等待插入意见。"""
        w = getattr(self, "_orch_worker", None)
        if w and getattr(self, "_orch_running", False):
            w.request_pause()
            self.orch_log.insertPlainText("\n⏸ 已请求暂停，将在下一阶段边界挂起…\n")

    def _on_orch_need_choice(self, options, raw):
        """爆款雷达出选项：展示选项按钮让你选定（或自填）。"""
        # 清空旧按钮
        while self.orch_choice_area_lay.count():
            item = self.orch_choice_area_lay.takeAt(0)
            wgt = item.widget()
            if wgt:
                wgt.deleteLater()
        self._orch_choice_opts = list(options)
        for idx, opt in enumerate(options):
            btn = QPushButton(f"选项{idx+1}：{opt['label']}")
            btn.setStyleSheet(self._secondary_btn_style())
            btn.setFixedHeight(34)
            btn.clicked.connect(lambda _checked=False, i=idx: self._on_orch_choice(i))
            self.orch_choice_area_lay.addWidget(btn)
        self.orch_choice_box.setVisible(True)
        self.orch_pause_btn.setEnabled(False)
        self.orch_log.insertPlainText("\n🔘 请选择爆款雷达给出的切入点（或自填后点「用我的输入」）。\n")

    def _on_orch_choice(self, idx):
        w = getattr(self, "_orch_worker", None)
        if w:
            w.choose(idx)
        self.orch_choice_box.setVisible(False)
        self.orch_pause_btn.setEnabled(True)

    def _on_orch_choice_custom(self):
        text = self.orch_choice_custom.text().strip()
        if not text:
            self.orch_log.insertPlainText("⚠️ 自填内容为空，请先输入切入点方向。\n")
            return
        w = getattr(self, "_orch_worker", None)
        if w:
            w.choose_custom(text)
        self.orch_choice_box.setVisible(False)
        self.orch_pause_btn.setEnabled(True)

    def _on_orch_paused(self, label):
        """worker 在阶段边界暂停，展示反馈框。"""
        self.orch_pause_hint.setText(f"当前阶段：{label}")
        self.orch_feedback.setText("")
        self.orch_pause_box.setVisible(True)
        self.orch_pause_btn.setEnabled(False)

    def _on_orch_resume(self):
        w = getattr(self, "_orch_worker", None)
        if w:
            w.resume_with_feedback(self.orch_feedback.text())
        self.orch_pause_box.setVisible(False)
        self.orch_pause_btn.setEnabled(True)

    # ============ 生图页 ============
    def _build_image_page(self):
        page = self.image_page
        lay = QVBoxLayout(page)
        lay.setContentsMargins(32, 24, 32, 24)
        lay.setSpacing(16)
        head = QLabel("生图 · Agnes Image")
        head.setStyleSheet(f"font-size:20px;font-weight:700;color:{THEME['text']};")
        lay.addWidget(head)
        sub = QLabel(f"模型：{self.cfg.get('image_gen_model','agnes-image-2.1-flash')}　尺寸可下方选择")
        sub.setStyleSheet(f"font-size:12px;color:{THEME['dim']};")
        lay.addWidget(sub)

        # ---- 参数行：尺寸下拉 ----
        opt_row = QHBoxLayout()
        opt_row.setSpacing(10)
        size_lbl = QLabel("尺寸")
        size_lbl.setStyleSheet(f"font-size:13px;color:{THEME['text']};")
        opt_row.addWidget(size_lbl)
        self.image_size_combo = QComboBox()
        self.image_size_combo.setFixedHeight(34)
        self.image_size_combo.setMinimumWidth(150)
        for s in ["1024x1024 方形", "1024x768 横版 4:3", "768x1024 竖版 3:4",
                  "1280x720 横版 16:9", "720x1280 竖版 9:16",
                  "1536x1024 横版 3:2", "1024x1536 竖版 2:3"]:
            self.image_size_combo.addItem(s)
        cur = self.cfg.get("image_gen_size", "1024x768")
        idx = self.image_size_combo.findText(cur, Qt.MatchStartsWith)
        if idx >= 0:
            self.image_size_combo.setCurrentIndex(idx)
        self.image_size_combo.setStyleSheet(
            f"QComboBox{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:4px 10px;font-size:13px;color:{THEME['text']};}}"
            f"QComboBox:focus{{border:1px solid {THEME['accent']};}}"
            f"QComboBox::drop-down{{border:none;width:18px;}}"
            f"QComboBox QAbstractItemView{{background:{THEME['card']};"
            f"border:1px solid {THEME['border']};color:{THEME['text']};"
            f"selection-background-color:{THEME['accent']};}}")
        opt_row.addWidget(self.image_size_combo)
        opt_row.addStretch(1)
        lay.addLayout(opt_row)

        self.image_prompt = QTextEdit()
        self.image_prompt.setFixedHeight(90)
        self.image_prompt.setPlaceholderText("描述你想生成的画面…")
        self.image_prompt.setStyleSheet(
            f"QTextEdit{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:10px;padding:10px 12px;font-size:13px;color:{THEME['text']};}}"
            f"QTextEdit:focus{{border:1px solid {THEME['accent']};}}")
        lay.addWidget(self.image_prompt)

        row = QHBoxLayout()
        gen = QPushButton("生成图片")
        gen.setFixedHeight(36)
        gen.setStyleSheet(self._primary_btn_style())
        gen.clicked.connect(self._gen_image)
        row.addWidget(gen)
        self.image_status = QLabel("")
        self.image_status.setStyleSheet(f"color:{THEME['dim']};font-size:12px;")
        row.addWidget(self.image_status)
        row.addStretch(1)
        lay.addLayout(row)

        self._image_paths = []
        self.image_result = QListWidget()
        self.image_result.setStyleSheet(
            f"QListWidget{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:10px;padding:8px;font-size:13px;color:{THEME['text']};}}")
        self.image_result.itemDoubleClicked.connect(self._open_image_result)
        lay.addWidget(self.image_result, 1)

    def _gen_image(self):
        prompt = self.image_prompt.toPlainText().strip()
        if not prompt:
            self.image_status.setText("请输入画面描述")
            return
        size = self.image_size_combo.currentText().split()[0]  # "1024x768 横版 4:3" -> "1024x768"
        self.image_status.setText(f"生成中…（{size}）")
        self._image_thread = _GenThread(tools_mod.tool_image_gen, self.cfg, APP_DIR, prompt, size)
        self._image_thread.result.connect(self._on_image_result)
        self._image_thread.start()

    def _on_image_result(self, res):
        if isinstance(res, str):
            self.image_status.setText(res)
            return
        rel, kind, name = res
        self.image_status.setText(f"已生成：{name}")
        self._image_paths.append(os.path.join(APP_DIR, rel))
        self.image_result.addItem(name)
        self.store.active().deliverables.append(
            {"rel": rel, "kind": kind, "name": name, "desc": rel})
        self.store.save()
        self._refresh_deliverables()

    def _open_image_result(self, item):
        idx = self.image_result.row(item)
        if 0 <= idx < len(self._image_paths):
            QDesktopServices.openUrl(QUrl.fromLocalFile(self._image_paths[idx]))

    # ============ 生视频页 ============
    def _build_video_page(self):
        page = self.video_page
        lay = QVBoxLayout(page)
        lay.setContentsMargins(32, 24, 32, 24)
        lay.setSpacing(16)
        head = QLabel("生视频 · Agnes Video")
        head.setStyleSheet(f"font-size:20px;font-weight:700;color:{THEME['text']};")
        lay.addWidget(head)
        sub = QLabel("文生视频 / 图生视频（Agnes 直连，免费）。生成可能需数分钟，请耐心等待。")
        sub.setStyleSheet(f"font-size:12px;color:{THEME['dim']};")
        lay.addWidget(sub)

        self.video_prompt = QTextEdit()
        self.video_prompt.setFixedHeight(90)
        self.video_prompt.setPlaceholderText("描述视频画面与镜头…（口播台词请填下方「台词/口播」框，不要写这里）")
        self.video_prompt.setStyleSheet(
            f"QTextEdit{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:10px;padding:10px 12px;font-size:13px;color:{THEME['text']};}}"
            f"QTextEdit:focus{{border:1px solid {THEME['accent']};}}")
        lay.addWidget(self.video_prompt)

        # ---- 选项行：时长 / 横竖 / 生成 ----
        opt = QHBoxLayout()
        self.video_duration = QSpinBox()
        # 新版 agnes-video-2.5-flash 时长合法范围 4~12 秒（旧版可到 16s），
        # UI 上限同步收窄，避免用户选了 13~16 却被静默钳到 12 而困惑。
        self.video_duration.setRange(4, 12)
        self.video_duration.setValue(6)
        self.video_duration.setSuffix(" 秒")
        self.video_duration.setFixedHeight(34)
        self.video_duration.setStyleSheet(
            f"QSpinBox{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 10px;font-size:13px;color:{THEME['text']};}}")
        opt.addWidget(self.video_duration)

        self.video_resolution = QComboBox()
        # 预设均经 2026-08-17 实测：Agnes 视频接受任意 WxH（无白名单），至少支持到 4K。
        for label, val in [
            ("竖屏 1080×1920 (9:16)", "1080x1920"),
            ("竖屏 720×1280 (9:16)", "720x1280"),
            ("竖屏 768×1152 (3:4)", "768x1152"),
            ("横屏 1920×1080 (16:9)", "1920x1080"),
            ("横屏 1280×720 (16:9)", "1280x720"),
            ("横屏 1152×768 (4:3)", "1152x768"),
            ("横屏 1088×832 (4:3)", "1088x832"),
            ("方形 1024×1024 (1:1)", "1024x1024"),
        ]:
            self.video_resolution.addItem(label, val)
        self.video_resolution.setCurrentIndex(2)  # 默认竖屏 768×1152（保持原默认）
        self.video_resolution.setFixedHeight(34)
        self.video_resolution.setStyleSheet(
            f"QComboBox{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 10px;font-size:13px;color:{THEME['text']};}}")
        opt.addWidget(self.video_resolution)

        gen = QPushButton("生成视频")
        gen.setFixedHeight(34)
        gen.setStyleSheet(self._primary_btn_style())
        gen.clicked.connect(self._gen_video)
        opt.addWidget(gen)
        lay.addLayout(opt)

        # ---- 首帧 / 尾帧上传（关键帧模式）----
        def _frame_row(label_text, line_edit, tooltip):
            row = QHBoxLayout()
            lbl = QLabel(label_text)
            lbl.setFixedWidth(36)
            lbl.setStyleSheet(f"font-size:13px;color:{THEME['text']};")
            row.addWidget(lbl)
            line_edit.setPlaceholderText(tooltip)
            line_edit.setFixedHeight(34)
            line_edit.setStyleSheet(
                f"QLineEdit{{background:{THEME['card']};border:1px solid {THEME['border']};"
                f"border-radius:8px;padding:0 12px;font-size:12px;color:{THEME['text']};}}")
            row.addWidget(line_edit, 1)
            browse = QPushButton("浏览…")
            browse.setFixedHeight(34)
            browse.setFixedWidth(64)
            browse.setStyleSheet(
                f"QPushButton{{background:{THEME['card']};border:1px solid {THEME['border']};"
                f"border-radius:8px;font-size:13px;color:{THEME['text']};}}"
                f"QPushButton:hover{{background:{THEME['elev']};}}")
            browse.clicked.connect(lambda _, le=line_edit: self._pick_video_frame(le))
            row.addWidget(browse)
            return row

        self.video_first = QLineEdit()
        lay.addLayout(_frame_row("首帧", self.video_first,
                                 "首帧图 URL，或点“浏览”选本地图（可选；仅首帧=首帧锁定）"))
        self.video_last = QLineEdit()
        lay.addLayout(_frame_row("尾帧", self.video_last,
                                 "尾帧图 URL，或点“浏览”选本地图（与首帧同选=首尾帧过渡）"))

        hint = QLabel("提示：仅选首帧→单图首帧锁定；首尾都选→精确首尾帧过渡；都不选→纯文生视频。")
        hint.setStyleSheet(f"font-size:11px;color:{THEME['dim']};")
        lay.addWidget(hint)

        self.video_status = QLabel("")
        self.video_status.setStyleSheet(f"color:{THEME['dim']};font-size:12px;")
        lay.addWidget(self.video_status)

        self._video_paths = []
        self.video_result = QListWidget()
        self.video_result.setStyleSheet(
            f"QListWidget{{background:{THEME['card']};border:1px solid {THEME['border']};"
            f"border-radius:10px;padding:8px;font-size:13px;color:{THEME['text']};}}")
        self.video_result.itemDoubleClicked.connect(self._open_video_result)
        lay.addWidget(self.video_result, 1)

    def _gen_video(self):
        prompt = self.video_prompt.toPlainText().strip()
        if not prompt:
            self.video_status.setText("请输入视频描述")
            return
        res = self.video_resolution.currentData() or "768x1152"
        first = self.video_first.text().strip() or None
        last = self.video_last.text().strip() or None
        self.video_status.setText("提交任务中…（可能需数分钟）")
        self._video_thread = _GenThread(
            tools_mod.tool_video_gen, self.cfg, APP_DIR, prompt,
            self.video_duration.value(), None, resolution=res,
            first_frame=first, last_frame=last)
        self._video_thread.result.connect(self._on_video_result)
        self._video_thread.start()

    def _pick_video_frame(self, line_edit):
        """选择首尾帧本地图片，读成 base64 data URI 填入输入框（无需图床）。"""
        path, _ = QFileDialog.getOpenFileName(
            self, "选择帧图片", "", "图片 (*.png *.jpg *.jpeg *.webp *.bmp)")
        if not path:
            return
        data_uri = self._file_to_datauri(path)
        if data_uri:
            line_edit.setText(data_uri)
            self.video_status.setText(f"已载入帧图：{os.path.basename(path)}")
        else:
            self.video_status.setText("图片读取失败")

    @staticmethod
    def _file_to_datauri(path):
        """读取本地图片为 base64 data URI（优先用 PIL 缩放到最长边 1280 以减小体积）。"""
        try:
            raw = None
            mime = "image/jpeg"
            try:
                from PIL import Image
                import io
                with Image.open(path) as _im:
                    im = _im.convert("RGB")
                    w, h = im.size
                    max_edge = 1280
                    if max(w, h) > max_edge:
                        scale = max_edge / max(w, h)
                        im = im.resize((int(w * scale), int(h * scale)))
                    buf = io.BytesIO()
                    ext = os.path.splitext(path)[1].lower()
                    fmt = "PNG" if ext in (".png", ".bmp") else "JPEG"
                    im.save(buf, fmt)
                    raw = buf.getvalue()
                    mime = "image/png" if fmt == "PNG" else "image/jpeg"
            except Exception:
                with open(path, "rb") as f:
                    raw = f.read()
                mime = "image/png" if path.lower().endswith((".png", ".bmp")) else "image/jpeg"
            b64 = base64.b64encode(raw).decode("ascii")
            return f"data:{mime};base64,{b64}"
        except Exception:
            return None

    def _on_video_result(self, res):
        if isinstance(res, str):
            self.video_status.setText(res)
            return
        rel, kind, name = res
        self.video_status.setText(f"已生成：{name}")
        self._video_paths.append(os.path.join(APP_DIR, rel))
        self.video_result.addItem(name)
        self.store.active().deliverables.append(
            {"rel": rel, "kind": kind, "name": name, "desc": rel})
        self.store.save()
        self._refresh_deliverables()

    def _open_video_result(self, item):
        idx = self.video_result.row(item)
        if 0 <= idx < len(self._video_paths):
            QDesktopServices.openUrl(QUrl.fromLocalFile(self._video_paths[idx]))

    # ============ 数字人分身页（整合工作台）============
    def _build_twin_page(self):
        """委托给 digital_twin_panel 模块构建（懒导入，避免拖累启动与打包）。"""
        from digital_twin_panel import build_twin_panel
        build_twin_panel(self)

    # ============ 导演台页（整合工作台）============
    def _build_director_page(self):
        """委托给 director_panel 模块构建（懒导入）。"""
        from director_panel import build_director_panel
        build_director_panel(self)

    # ============ 自动化任务页（v4.88）============
    def _build_automation_page(self):
        """委托给 automation_panel 模块构建（懒导入）。"""
        from automation_panel import build_automation_panel
        build_automation_panel(self)

    # ============ 自动化任务调度（v4.88）============
    def _reload_automation_if_changed(self):
        """Agent 工具（子线程）写 automation_tasks.json 后，主窗口按 mtime 感知并重载。

        面板与 Agent 工具都落盘到同一文件；面板操作同一 store 实例即时可见，
        Agent 工具则在子线程新开实例写文件，靠这里每秒对比 mtime 拉新，1 秒内同步。
        """
        try:
            mtime = os.path.getmtime(self.automation_store.path)
        except OSError:
            return
        if mtime == getattr(self, "_auto_mtime", None):
            return
        self._auto_mtime = mtime
        self.automation_store._load()
        # 面板若正在显示，同步刷新列表
        try:
            if hasattr(self, "_refresh_automation_list"):
                self._refresh_automation_list()
        except Exception:
            pass

    def _on_automation_tick(self):
        """每秒检查到期任务：提醒→弹窗；执行→交给 Agent 后台跑。"""
        try:
            self._reload_automation_if_changed()
            now = datetime.now()
            for t in self.automation_store.list_enabled():
                if not automation.is_due(t, now):
                    continue
                action = t.get("action", automation.ACT_REMIND)
                if action == automation.ACT_REMIND:
                    automation.mark_fired(t, now)
                    self.automation_store.save()
                    self._fire_reminder(t.get("message", ""))
                else:
                    # 执行任务：App 忙则跳过，等下一 tick 重试（不标记 fired）
                    if self._busy:
                        continue
                    automation.mark_fired(t, now)
                    self.automation_store.save()
                    self._fire_automation_run(t)
        except Exception as e:
            log.error("自动化任务调度异常: %s", e)

    def _fire_automation_run(self, task):
        """把任务的执行指令作为一条 user 消息交给 Agent 跑（结果流式显示在对话页）。"""
        try:
            msg = (task.get("message") or "").strip()
            if not msg:
                return
            session = self.store.active()
            session.messages.append({"role": "user", "content": f"【自动化任务】{msg}"})
            self.store.save()
            # 切到对话页（nav index 0）让用户看到结果
            self._switch_nav(0)
            self._render_messages(force_bottom=True)
            # v4.95：标记本次 run 是自动化任务触发，完成时托盘+语音通知
            self._pending_done_notify = (task.get("name") or "").strip() or "自动化任务"
            self._agent_run()
        except Exception as e:
            log.error("自动化任务执行失败: %s", e)

    # ============ 工具箱页 ============
    def _build_tools_page(self):
        page = self.tools_page
        lay = QVBoxLayout(page)
        lay.setContentsMargins(32, 24, 32, 24)
        lay.setSpacing(16)
        head = QLabel("工具箱")
        head.setStyleSheet(f"font-size:20px;font-weight:700;color:{THEME['text']};")
        lay.addWidget(head)
        sub = QLabel("点选技能后将应用到当前对话，并自动切到对话页。")
        sub.setStyleSheet(f"font-size:12px;color:{THEME['dim']};")
        lay.addWidget(sub)

        # 技能市场入口
        market_row = QHBoxLayout()
        market_btn = QPushButton("🛍 浏览技能市场")
        market_btn.setFixedHeight(34)
        market_btn.setCursor(Qt.PointingHandCursor)
        market_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['accent']};color:#FFFFFF;border:none;"
            f"border-radius:8px;padding:0 16px;font-size:13px;font-weight:500;}}"
            f"QPushButton:hover{{background:{THEME['accent_hover']};}}")
        market_btn.clicked.connect(self._open_skill_market)
        market_row.addWidget(market_btn)
        market_row.addStretch(1)
        lay.addLayout(market_row)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea{background:transparent;border:none;}")
        inner = QWidget()
        grid = QGridLayout(inner)
        grid.setContentsMargins(0, 0, 0, 0)
        grid.setSpacing(10)
        scroll.setWidget(inner)
        lay.addWidget(scroll, 1)

        cats = {}
        order = []
        for sk in self._skills:
            c = sk.get("category", "其他")
            cats.setdefault(c, []).append(sk)
            if c not in order:
                order.append(c)
        r = 0
        for c in order:
            cl = QLabel(c)
            cl.setStyleSheet(f"font-size:12px;font-weight:600;color:{THEME['faint']};")
            grid.addWidget(cl, r, 0, 1, 3)
            r += 1
            col = 0
            for sk in cats[c]:
                btn = QPushButton(f"{sk.get('emoji','')}  {sk.get('name','')}")
                btn.setFixedHeight(64)
                btn.setCursor(Qt.PointingHandCursor)
                btn.setStyleSheet(
                    f"QPushButton{{background:{THEME['card']};border:1px solid {THEME['border']};"
                    f"border-radius:10px;padding:8px 12px;font-size:13px;font-weight:500;"
                    f"color:{THEME['text']};text-align:left;}}"
                    f"QPushButton:hover{{border-color:{THEME['accent']};background:{THEME['surface']};}}")
                btn.setToolTip(sk.get("desc", ""))
                btn.clicked.connect(lambda _=False, sid=sk["id"]: self._on_tool_pick(sid))
                grid.addWidget(btn, r, col)
                col += 1
                if col >= 3:
                    col = 0
                    r += 1
            if col != 0:
                r += 1
        grid.setRowStretch(r, 1)

    def _on_tool_pick(self, sid):
        self._on_skill_pick(sid)
        self.main_stack.setCurrentIndex(1)
        self._update_nav_styles(0)

    # ============ 设置页 ============
    def _build_settings_page(self):
        page = self.settings_page
        page.setFont(QFont("Microsoft YaHei", 13))
        lay = QVBoxLayout(page)
        lay.setContentsMargins(32, 24, 32, 24)
        lay.setSpacing(16)
        head = QLabel("设置")
        head.setStyleSheet(f"font-size:20px;font-weight:700;color:{THEME['text']};background:transparent;")
        lay.addWidget(head)
        sub = QLabel("当前配置摘要；点击按钮打开详细设置弹层。")
        sub.setStyleSheet(f"font-size:12px;color:{THEME['dim']};background:transparent;")
        lay.addWidget(sub)

        info = QWidget()
        info.setObjectName("settingsInfoCard")
        info.setStyleSheet(f"QWidget#settingsInfoCard{{background:{THEME['card']};border:1px solid {THEME['border']};"
                           f"border-radius:10px;padding:16px;}}")
        il = QVBoxLayout(info)
        il.setSpacing(8)
        prof = self.cfg.get("model_profiles", {})
        cur = next((n for n, p in prof.items()
                    if p.get("base_url") == self.cfg.get("base_url")
                    and p.get("model") == self.cfg.get("model")), "—")
        il.addWidget(self._kv("当前模型", cur))
        il.addWidget(self._kv("Agent 模式", "开" if self.cfg.get("agent_mode") else "关"))
        il.addWidget(self._kv("联网搜索", "开" if self.cfg.get("search_enabled", True) else "关"))
        key = self.cfg.get("api_key", "")
        il.addWidget(self._kv("API Key",
                     ("●" * min(8, len(key)) or "未设置") + ("（已隐藏）" if key else "")))
        lay.addWidget(info)

        btn_row = QHBoxLayout()
        open_btn = QPushButton("打开设置弹层")
        open_btn.setFixedHeight(36)
        open_btn.setStyleSheet(self._primary_btn_style())
        open_btn.clicked.connect(lambda: self._show_popup(self._settings_popup))
        btn_row.addWidget(open_btn)

        export_btn = QPushButton("导出当前对话")
        export_btn.setFixedHeight(36)
        export_btn.setStyleSheet(self._secondary_btn_style())
        export_btn.clicked.connect(self.export_session)
        btn_row.addWidget(export_btn)
        diag_btn = QPushButton("导出诊断包")
        diag_btn.setFixedHeight(36)
        diag_btn.setStyleSheet(self._secondary_btn_style())
        diag_btn.clicked.connect(self.export_diagnostic)
        btn_row.addWidget(diag_btn)
        btn_row.addStretch(1)
        lay.addLayout(btn_row)

        # ===== 我的记忆（跨对话长期记忆）=====
        mem_card = QWidget()
        mem_card.setObjectName("settingsMemCard")
        mem_card.setStyleSheet(f"QWidget#settingsMemCard{{background:{THEME['card']};border:1px solid {THEME['border']};"
                               f"border-radius:10px;padding:16px;}}")
        ml = QVBoxLayout(mem_card)
        ml.setSpacing(10)
        mhead = QLabel("我的记忆（跨对话长期记忆）")
        mhead.setStyleSheet(f"font-size:15px;font-weight:600;color:{THEME['text']};background:transparent;")
        ml.addWidget(mhead)
        msub = QLabel("小臭会在对话中自动记住你的稳定偏好、约定与身份，并在新对话里沿用。"
                      "可在此查看；清空会删除全部记忆（不可恢复）。")
        msub.setWordWrap(True)
        msub.setStyleSheet(f"font-size:12px;color:{THEME['dim']};background:transparent;")
        ml.addWidget(msub)
        self.mem_view = QTextEdit()
        self.mem_view.setReadOnly(True)
        self.mem_view.setMaximumHeight(150)
        self.mem_view.setStyleSheet(f"background:{THEME['bg']};border:1px solid {THEME['border']};"
                                    f"border-radius:8px;font-size:12px;color:{THEME['text']};")
        ml.addWidget(self.mem_view)
        mbtns = QHBoxLayout()
        mrefresh = QPushButton("刷新")
        mrefresh.setFixedHeight(32)
        mrefresh.setStyleSheet(self._secondary_btn_style())
        mrefresh.clicked.connect(self._refresh_memory_view)
        mclear = QPushButton("清空记忆")
        mclear.setFixedHeight(32)
        mclear.setStyleSheet("QPushButton{background:#dc2626;color:#fff;border:none;"
                             "border-radius:8px;padding:0 14px;font-size:13px;font-weight:600;}"
                             "QPushButton:hover{background:#b91c1c;}")
        mclear.clicked.connect(self._clear_memory)
        mbtns.addWidget(mrefresh)
        mbtns.addWidget(mclear)
        mbtns.addStretch(1)
        ml.addLayout(mbtns)
        lay.addWidget(mem_card)
        self._refresh_memory_view()

        # ===== 浏览器扩展（v4.103：抓取网页进对话）=====
        ext_card = QWidget()
        ext_card.setObjectName("settingsExtCard")
        ext_card.setStyleSheet(f"QWidget#settingsExtCard{{background:{THEME['card']};border:1px solid {THEME['border']};"
                               f"border-radius:10px;padding:16px;}}")
        el = QVBoxLayout(ext_card)
        el.setSpacing(10)
        ehead = QLabel("浏览器扩展（抓网页进对话）")
        ehead.setStyleSheet(f"font-size:15px;font-weight:600;color:{THEME['text']};background:transparent;")
        el.addWidget(ehead)
        esub = QLabel("装好「小臭抓网页」扩展后，在任意网页一键把正文/选中文字发到小臭，"
                      "让 AI 帮你总结、提取、分析。配对码只需复制一次到扩展里。")
        esub.setWordWrap(True)
        esub.setStyleSheet(f"font-size:12px;color:{THEME['dim']};background:transparent;")
        el.addWidget(esub)
        etok_row = QHBoxLayout()
        etok_lbl = QLabel("配对码：")
        etok_lbl.setStyleSheet(f"font-size:12px;color:{THEME['text']};background:transparent;")
        etok_row.addWidget(etok_lbl)
        tok_val = self.cfg.get("browser_bridge_token", "")
        self.ext_token_edit = QLineEdit(tok_val)
        self.ext_token_edit.setReadOnly(True)
        self.ext_token_edit.setStyleSheet(f"background:{THEME['bg']};border:1px solid {THEME['border']};"
                                          f"border-radius:8px;padding:6px 8px;font-size:13px;"
                                          f"color:{THEME['text']};font-family:'Microsoft YaHei','ui-monospace','Menlo','Consolas',monospace;")
        etok_row.addWidget(self.ext_token_edit, 1)
        ecopy = QPushButton("复制")
        ecopy.setFixedHeight(32)
        ecopy.setStyleSheet(self._secondary_btn_style())
        ecopy.clicked.connect(self._copy_ext_token)
        etok_row.addWidget(ecopy)
        el.addLayout(etok_row)
        estatus = QLabel("桥接服务：本机 127.0.0.1:9100（已自动启动）")
        estatus.setStyleSheet(f"font-size:12px;color:{THEME['dim']};background:transparent;")
        el.addWidget(estatus)
        einstall = QPushButton("打开扩展安装说明")
        einstall.setFixedHeight(34)
        einstall.setStyleSheet(self._secondary_btn_style())
        # v4.108 M-27：优先运行目录（用户自装/自改），缺失时回退到 exe 随附的
        # 扩展副本——PyInstaller 6.x onedir 的 datas 落 _internal/ 下，开发态在仓库根。
        install_path = os.path.join(
            os.path.expanduser("~"), "Documents", "小臭玩AI", "browser_extension", "README.md")
        if not os.path.exists(install_path):
            try:
                _exe_dir = os.path.dirname(sys.executable)
                for _cand in (os.path.join(_exe_dir, "_internal", "browser_extension",
                                           "README.md"),
                              os.path.join(_exe_dir, "browser_extension", "README.md")):
                    if os.path.exists(_cand):
                        install_path = _cand
                        break
            except Exception:
                pass
        einstall.clicked.connect(lambda: self._open_path(install_path))
        el.addWidget(einstall)
        lay.addWidget(ext_card)

        lay.addStretch(1)

    def _kv(self, k, v):
        w = QWidget()
        hl = QHBoxLayout(w)
        hl.setContentsMargins(0, 0, 0, 0)
        k_l = QLabel(k)
        k_l.setStyleSheet(f"font-size:13px;color:{THEME['dim']};background:transparent;")
        v_l = QLabel(str(v))
        v_l.setStyleSheet(f"font-size:13px;color:{THEME['text']};font-weight:500;background:transparent;")
        hl.addWidget(k_l)
        hl.addStretch(1)
        hl.addWidget(v_l)
        return w

    # ============ 跨对话长期记忆（我的记忆）============
    def _refresh_memory_view(self):
        """刷新设置页的『我的记忆』面板内容。"""
        try:
            from memory_store import load_memory, memory_stats
            mem = load_memory()
            n, sz = memory_stats()
            self.mem_view.setPlainText(mem if mem else "（暂无长期记忆，对话中小臭会自动积累）")
            self.mem_view.append(f"\n— 共 {n} 条 · {sz} 字节 —")
        except Exception as e:
            self.mem_view.setPlainText(f"读取记忆失败：{e}")

    def _clear_memory(self):
        """清空全部长期记忆（带二次确认）。"""
        dlg = ConfirmDialog(self)
        dlg.set_text("清空长期记忆", "确定要删除全部跨对话记忆吗？\n此操作不可恢复。")
        dlg.exec()
        if dlg.result():
            try:
                from memory_store import clear_memory
                if clear_memory():
                    self._refresh_memory_view()
                    self.status_label.setText("长期记忆已清空")
            except Exception as e:
                self.status_label.setText(f"清空失败：{e}")

    # ============ 自定义标题栏 ============
    def _build_title_bar(self):
        """构建顶栏(48)：Logo/应用名 + 搜索框(480) + 头像 + 系统按钮，支持拖动。"""
        self.title_bar = QWidget()
        self.title_bar.setFixedHeight(48)
        self.title_bar.setObjectName("titleBar")
        self.title_bar.setStyleSheet(
            f"QWidget#titleBar{{background:{THEME['card']};border-bottom:1px solid {THEME['border']};}}"
        )
        tb_layout = QHBoxLayout(self.title_bar)
        tb_layout.setContentsMargins(16, 0, 12, 0)
        tb_layout.setSpacing(12)

        # ---- 左侧 Logo + 应用名 ----
        logo_icon = QLabel()
        logo_icon.setFixedSize(28, 28)
        logo_icon.setScaledContents(True)
        logo_icon.setPixmap(QPixmap(resource_path("images/logo.png")).scaled(
            28, 28, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        tb_layout.addWidget(logo_icon)
        logo_name = QLabel("小臭玩AI")
        logo_name.setStyleSheet(
            f"font-size:16px;font-weight:600;color:{THEME['text']};background:transparent;")
        logo_name.setCursor(Qt.PointingHandCursor)
        logo_name.mousePressEvent = lambda e: self._go_home()
        tb_layout.addWidget(logo_name)
        tb_layout.addSpacing(16)

        # ---- 搜索框 (480×36) ----
        self.search_box = QLineEdit()
        self.search_box.setFixedSize(480, 36)
        self.search_box.setPlaceholderText("搜索对话、文件、工具…")
        self.search_box.setStyleSheet(
            f"QLineEdit{{background:{THEME['bg']};border:1px solid {THEME['border']};"
            f"border-radius:18px;padding:0 16px;font-size:13px;color:{THEME['text']};}}"
            f"QLineEdit:hover{{border-color:{THEME['border_hover']};}}"
            f"QLineEdit:focus{{border:1px solid {THEME['accent']};}}")
        tb_layout.addWidget(self.search_box)
        # v4.75：对话内搜索（回车搜索 / Shift+回车上一处 / Esc 清除）
        self.search_box.installEventFilter(self)

        tb_layout.addStretch(1)

        # ---- 右侧头像 ----
        self.user_avatar = QLabel("xyb")
        self.user_avatar.setFixedSize(34, 34)
        self.user_avatar.setAlignment(Qt.AlignCenter)
        self.user_avatar.setStyleSheet(
            f"QLabel{{background:{THEME['accent']};color:#FFFFFF;border-radius:17px;"
            f"font-size:13px;font-weight:600;}}")
        tb_layout.addWidget(self.user_avatar)

        # ---- 系统按钮 ----
        btn_base = (
            "QPushButton{{background:transparent;color:{dim};border:none;font-size:13px;}}"
            "QPushButton:hover{{background:rgba(32,33,36,0.06);color:{text};}}"
        ).format(dim=THEME["dim"], text=THEME["text"])
        close_base = (
            "QPushButton{{background:transparent;color:{dim};border:none;font-size:14px;}}"
            "QPushButton:hover{{background:#EA4335;color:#FFFFFF;}}"
        ).format(dim=THEME["dim"])
        btn_w, btn_h = 44, 48
        self.min_btn = QPushButton("—")
        self.min_btn.setFixedSize(btn_w, btn_h)
        self.min_btn.setStyleSheet(btn_base)
        self.min_btn.clicked.connect(self.showMinimized)
        tb_layout.addWidget(self.min_btn)
        self.max_btn = QPushButton("□")
        self.max_btn.setFixedSize(btn_w, btn_h)
        self.max_btn.setStyleSheet(btn_base)
        self.max_btn.clicked.connect(self._toggle_maximize)
        tb_layout.addWidget(self.max_btn)
        self.close_btn = QPushButton("✕")
        self.close_btn.setFixedSize(btn_w, btn_h)
        self.close_btn.setStyleSheet(close_base)
        self.close_btn.clicked.connect(self.close)
        tb_layout.addWidget(self.close_btn)

        # ---- 窗口拖动 ----
        self._drag_pos = None
        self.title_bar.mousePressEvent = self._title_bar_mouse_press
        self.title_bar.mouseMoveEvent = self._title_bar_mouse_move
        self.title_bar.mouseReleaseEvent = self._title_bar_mouse_release
        self.title_bar.mouseDoubleClickEvent = self._title_bar_double_click

    def _go_home(self):
        """点击 Logo 返回首页。"""
        self.main_stack.setCurrentIndex(0)
        self._update_nav_styles(-1)

    def _toggle_maximize(self):
        if self.isMaximized():
            self.showNormal()
            self.max_btn.setText("\u25a1")
        else:
            self.showMaximized()
            self.max_btn.setText("\u2750")

    def _title_bar_mouse_press(self, event):
        if event.button() == Qt.LeftButton:
            self._drag_pos = event.globalPosition().toPoint()

    def _title_bar_mouse_move(self, event):
        if self._drag_pos is not None and event.buttons() & Qt.LeftButton:
            delta = event.globalPosition().toPoint() - self._drag_pos
            self.move(self.pos() + delta)
            self._drag_pos = event.globalPosition().toPoint()

    def _title_bar_mouse_release(self, event):
        self._drag_pos = None

    def _title_bar_double_click(self, event):
        if event.button() == Qt.LeftButton:
            self._toggle_maximize()

    # ============ 欢迎页 ============
    def _build_welcome_page(self):
        """构建欢迎页：问候语 + 标题 + 3 张功能卡片 + 最近对话列表。"""
        wl = QVBoxLayout(self.welcome_page)
        wl.setContentsMargins(32, 36, 32, 28)
        wl.setSpacing(0)

        # 滚动区域包裹
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea{background:transparent;border:none;}")
        scroll_content = QWidget()
        scroll_content.setStyleSheet("background:transparent;")
        scl = QVBoxLayout(scroll_content)
        scl.setContentsMargins(0, 0, 0, 0)
        scl.setSpacing(0)

        # ---- 问候语 ----
        hour = datetime.now().hour
        if hour < 6:
            greeting_text = "夜深了"
        elif hour < 12:
            greeting_text = "上午好"
        elif hour < 14:
            greeting_text = "中午好"
        elif hour < 18:
            greeting_text = "下午好"
        else:
            greeting_text = "晚上好"

        greeting = QLabel(f"{greeting_text}")
        greeting.setStyleSheet(
            f"font-size:14px;font-weight:600;color:{THEME['accent']};margin-bottom:4px;")
        scl.addWidget(greeting)

        # ---- 标题 ----
        title = QLabel("今天想做什么？")
        title.setStyleSheet(
            f"font-size:28px;font-weight:700;color:{THEME['text']};"
            f"letter-spacing:-0.5px;margin-bottom:28px;")
        scl.addWidget(title)

        # ---- 3 张功能卡片 ----
        cards_row = QHBoxLayout()
        cards_row.setSpacing(16)

        card_data = [
            ("blue",   "✏️", "写文章",
             "知乎短故事、头条评论、小红书笔记，AI帮你写"),
            ("green",  "💻", "写代码",
             "修Bug、加功能、调样式，交给Agent搞定"),
            ("orange", "🎨", "做设计",
             "生成配图、改页面UI、出HTML排版"),
        ]

        self.welcome_cards = []
        for color_key, icon_str, card_title, card_desc in card_data:
            card = QPushButton()
            card.setFixedSize(190, 140)
            card.setCursor(Qt.PointingHandCursor)

            # 卡片配色
            bg_key = f"card_{color_key}_bg"
            icon_key = f"card_{color_key}_icon"
            card_bg = THEME.get(bg_key, THEME["surface_raised"])
            card_icon_color = THEME.get(icon_key, THEME["accent"])

            card.setStyleSheet(
                f"QPushButton{{"
                f"background:{card_bg};"
                f"border:1px solid {THEME['border']};"
                f"border-radius:10px;"
                f"padding:14px 20px 16px 18px;text-align:left;"
                f"}}"
                f"QPushButton:hover{{"
                f"border-color:{card_icon_color};"
                f"background:{THEME['elev']};"
                f"margin:-2px;"
                f"}}"
            )

            # 卡片内部布局
            card_lay = QVBoxLayout(card)
            card_lay.setContentsMargins(0, 0, 0, 0)
            card_lay.setSpacing(8)

            icon_lbl = QLabel(icon_str)
            icon_lbl.setStyleSheet(
                f"font-size:22px;background:transparent;border:none;color:{card_icon_color};")
            card_lay.addWidget(icon_lbl)

            title_lbl = QLabel(card_title)
            title_lbl.setStyleSheet(
                f"font-size:16px;font-weight:600;color:{THEME['text']};"
                f"background:transparent;border:none;")
            card_lay.addWidget(title_lbl)

            desc_lbl = QLabel(card_desc)
            desc_lbl.setWordWrap(True)
            desc_lbl.setMinimumHeight(40)
            desc_lbl.setStyleSheet(
                f"font-size:12px;color:{THEME['dim']};line-height:1.55;"
                f"background:transparent;border:none;word-wrap:break-word;")
            card_lay.addWidget(desc_lbl)
            card_lay.addStretch()

            card.clicked.connect(self._on_welcome_card)
            self.welcome_cards.append(card)
            cards_row.addWidget(card)

        scl.addLayout(cards_row)
        scl.addSpacing(32)

        # ---- 最近对话区域 ----
        recent_header = QHBoxLayout()
        recent_label = QLabel("最近对话")
        recent_label.setStyleSheet(
            f"font-size:12px;font-weight:600;color:{THEME['faint']};"
            f"text-transform:uppercase;letter-spacing:0.5px;")
        recent_header.addWidget(recent_label)
        recent_header.addStretch(1)
        view_all = QLabel("查看全部")
        view_all.setStyleSheet(
            f"font-size:12px;color:{THEME['accent']};font-weight:400;")
        view_all.setCursor(Qt.PointingHandCursor)
        view_all.mousePressEvent = lambda e: self._open_session_manager()
        recent_header.addWidget(view_all)
        scl.addLayout(recent_header)
        scl.addSpacing(10)

        # 最近对话列表容器
        self.welcome_recent_container = QWidget()
        self.welcome_recent_container.setStyleSheet("background:transparent;")
        self.welcome_recent_layout = QVBoxLayout(self.welcome_recent_container)
        self.welcome_recent_layout.setContentsMargins(0, 0, 0, 0)
        self.welcome_recent_layout.setSpacing(0)
        scl.addWidget(self.welcome_recent_container)

        scl.addStretch(1)
        scroll.setWidget(scroll_content)
        wl.addWidget(scroll)

    def _refresh_recent_on_welcome(self):
        """刷新欢迎页的最近对话列表。"""
        # 清空旧项
        while self.welcome_recent_layout.count():
            item = self.welcome_recent_layout.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()

        # v4.79：置顶优先 + 创建时间倒序，取最多 5 个（含置顶、排除当前激活）
        recent = [s.sid for s in self.store.all_sorted()
                  if s.messages and s.sid != self.store.active_sid][:5]

        if not recent:
            empty_lbl = QLabel("暂无对话记录")
            empty_lbl.setStyleSheet(
                f"color:{THEME['placeholder']};font-size:13px;padding:12px 0;")
            self.welcome_recent_layout.addWidget(empty_lbl)
            return

        dot_colors = [THEME["accent"], THEME["ok"], THEME["tool_running"]]
        for i, sid in enumerate(recent):
            s = self.store.sessions[sid]
            title = s.title or "新会话"
            display = title if len(title) <= 28 else title[:27] + "…"
            # 时间
            ts = getattr(s, 'updated_at', None) or s.created or datetime.now()
            if isinstance(ts, str):
                try:
                    ts = datetime.fromisoformat(ts)
                except ValueError:
                    ts = datetime.now()
            time_str = ts.strftime("%H:%M")

            row = QWidget()
            row.setCursor(Qt.PointingHandCursor)
            row.setFixedHeight(40)
            rl = QHBoxLayout(row)
            rl.setContentsMargins(0, 0, 0, 0)
            rl.setSpacing(10)

            dot = QLabel("")
            dot.setFixedSize(6, 6)
            dot.setStyleSheet(
                f"background:{dot_colors[i % len(dot_colors)]};"
                f"border-radius:3px;")
            rl.addWidget(dot)

            lbl = QLabel(display)
            lbl.setStyleSheet(
                f"color:{THEME['dim']};font-size:13px;background:transparent;")
            rl.addWidget(lbl, 1)

            time_lbl = QLabel(time_str)
            time_lbl.setStyleSheet(
                f"color:{THEME['placeholder']};font-size:12px;background:transparent;")
            rl.addWidget(time_lbl)

            # 悬停删除按钮
            del_btn = QPushButton("×")
            del_btn.setFixedSize(22, 22)
            del_btn.setCursor(Qt.PointingHandCursor)
            del_btn.setStyleSheet(
                f"QPushButton{{background:transparent;color:{THEME['placeholder']};"
                f"border:none;font-size:16px;font-weight:600;border-radius:4px;}}"
                f"QPushButton:hover{{color:{THEME['danger']};background:{THEME['sidebar_hover']};}}")
            del_btn.clicked.connect(
                lambda _, s=sid, t=title: self._request_delete_session(s, t))
            rl.addWidget(del_btn)

            # 整行点击切换（点到删除按钮除外）
            def _make_press(sid, del_btn, container_ref):
                def _press(e):
                    if container_ref.childAt(e.pos()) is del_btn:
                        return
                    self._switch_session(sid)
                return _press

            # 底部分隔线（最后一项不加）
            if i < len(recent) - 1:
                sep_w = QFrame()
                sep_w.setFrameShape(QFrame.HLine)
                sep_w.setStyleSheet(f"QFrame{{color:{THEME['separator']};max-height:1px;}}")
                # 用嵌套 layout 实现分隔线
                wrapper = QVBoxLayout()
                wrapper.setContentsMargins(0, 0, 0, 0)
                wrapper.setSpacing(0)
                wrapper.addWidget(row)
                wrapper.addWidget(sep_w)
                container = QWidget()
                container.setLayout(wrapper)
                container.setStyleSheet("background:transparent;")
                container.mousePressEvent = _make_press(sid, del_btn, container)
                self.welcome_recent_layout.addWidget(container)
            else:
                row.mousePressEvent = _make_press(sid, del_btn, row)
                self.welcome_recent_layout.addWidget(row)

    # ============ 侧边栏 (256px 导航) ============
    def _build_sidebar(self):
        """构建侧栏(256)：Logo + 6 导航 pill + 用户区。"""
        self.sidebar = QWidget()
        self.sidebar.setFixedWidth(256)
        self.sidebar.setStyleSheet(
            f"QWidget#sidebarWidget{{background:{THEME['sidebar']};border-right:1px solid {THEME['border']};}}")
        self.sidebar.setObjectName("sidebarWidget")
        sb = QVBoxLayout(self.sidebar)
        sb.setContentsMargins(16, 16, 16, 16)
        sb.setSpacing(8)

        # ---- Logo 行 ----
        logo_row = QHBoxLayout()
        logo_icon = QLabel()
        logo_icon.setFixedSize(28, 28)
        logo_icon.setScaledContents(True)
        logo_icon.setPixmap(QPixmap(resource_path("images/logo.png")).scaled(
            28, 28, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        logo_row.addWidget(logo_icon)
        logo_name = QLabel("小臭玩AI")
        logo_name.setStyleSheet(
            f"font-size:16px;font-weight:600;color:{THEME['text']};padding-left:8px;background:transparent;")
        logo_row.addWidget(logo_name, 1)
        sb.addLayout(logo_row)
        sb.addSpacing(12)

        # ---- 导航 pill ×6 ----
        self.nav_buttons = []
        # 面板挂载范式（整合工作台）：
        #   nav_defs 顺序 == main_stack 页面顺序（首页占 index 0，不在 nav 中）。
        #   新增面板只需：① 在此按序插入 (label, icon)；② 在 _init_ui 的 main_stack
        #   中按同一相对位置 addWidget 一个 QWidget 并调用其 _build_*_page；
        #   ③ _build_*_page 内懒导入对应面板模块（如 digital_twin_panel）委托构建。
        #   _switch_nav(idx) -> main_stack.setCurrentIndex(idx+1) 自动保持同步。
        nav_defs = [
            ("对话", "💬"),
            ("编排", "🔀"),
            ("生图", "🖼"),
            ("生视频", "🎬"),
            ("数字人", "🧑"),
            ("导演台", "🎥"),
            ("工具", "🛠"),
            ("任务", "⏰"),
            ("设置", "⚙"),
        ]
        nav_container = QWidget()
        nav_lay = QVBoxLayout(nav_container)
        nav_lay.setContentsMargins(0, 0, 0, 0)
        nav_lay.setSpacing(4)
        for i, (label, icon) in enumerate(nav_defs):
            btn = QPushButton(f"{icon}  {label}")
            btn.setFixedHeight(44)
            btn.setCursor(Qt.PointingHandCursor)
            btn.clicked.connect(lambda _, idx=i: self._switch_nav(idx))
            nav_lay.addWidget(btn)
            self.nav_buttons.append(btn)
        sb.addWidget(nav_container)

        sb.addStretch(1)

        # ---- 用户区 ----
        user_row = QHBoxLayout()
        avatar = QLabel("xyb")
        avatar.setFixedSize(34, 34)
        avatar.setAlignment(Qt.AlignCenter)
        avatar.setStyleSheet(
            f"QLabel{{background:{THEME['accent']};color:#FFFFFF;border-radius:17px;"
            f"font-size:13px;font-weight:600;}}")
        user_row.addWidget(avatar)
        uname = QLabel("xyb")
        uname.setStyleSheet(
            f"font-size:14px;font-weight:500;color:{THEME['text']};background:transparent;")
        user_row.addWidget(uname, 1)
        online = QLabel("● 在线")
        online.setStyleSheet(
            f"font-size:12px;color:{THEME['ok']};background:transparent;")
        user_row.addWidget(online)
        sb.addLayout(user_row)

        self._update_nav_styles(-1)

    def _update_nav_styles(self, active):
        """刷新导航 pill 选中态：active=-1 表示首页(无选中)。"""
        for i, btn in enumerate(self.nav_buttons):
            if i == active:
                btn.setStyleSheet(
                    f"QPushButton{{background:{THEME['accent']};color:#FFFFFF;border:none;"
                    f"border-radius:24px;padding:0 16px;font-size:14px;font-weight:500;text-align:left;}}"
                    f"QPushButton:hover{{background:{THEME['accent_hover']};color:#FFFFFF;}}")
            else:
                btn.setStyleSheet(
                    f"QPushButton{{background:transparent;color:{THEME['dim']};border:none;"
                    f"border-radius:24px;padding:0 16px;font-size:14px;font-weight:500;text-align:left;}}"
                    f"QPushButton:hover{{background:{THEME['blue_hover']};color:{THEME['text']};}}")

    def _switch_nav(self, index):
        """导航切换：nav 序号 -> 主栈页面(nav+1，首页占0)。"""
        self.main_stack.setCurrentIndex(index + 1)
        self._update_nav_styles(index)

    def _open_session_manager(self):
        """v4.79：打开会话管理对话框（置顶/分组/批量删除/筛选）。"""
        try:
            dlg = SessionManagerDialog(self)
            dlg.exec()
        except Exception as e:
            log.warning("打开会话管理失败: %s", e)
        self._refresh_session_combo()
        self._refresh_recent_on_welcome()

    def _refresh_session_combo(self):
        """刷新对话页顶部的会话下拉框。"""
        if not hasattr(self, "session_combo"):
            return
        self.session_combo.blockSignals(True)
        self.session_combo.clear()
        active_sid = self.store.active_sid
        idx = 0
        for i, (sid, s) in enumerate(self.store.sessions.items()):
            title = s.title or "新会话"
            display = title if len(title) <= 24 else title[:23] + "…"
            self.session_combo.addItem(display, sid)
            if sid == active_sid:
                idx = i
        self.session_combo.setCurrentIndex(idx)
        self.session_combo.blockSignals(False)

    def _on_session_combo(self, index):
        sid = self.session_combo.itemData(index)
        if sid:
            self._switch_session(sid)

    # ============ 欢迎页卡片点击 ============
    def _on_welcome_card(self):
        sender = self.sender()
        if not sender:
            return
        # 通过卡片索引确定
        idx = -1
        for i, card in enumerate(self.welcome_cards):
            if card is sender:
                idx = i
                break
        prompts = [
            "帮我写一篇知乎短故事，主题自拟",
            "帮我写一段 Python 代码，实现",
            "帮我设计一个页面 UI，要求",
        ]
        if 0 <= idx < len(prompts):
            self.input_box.setPlainText(prompts[idx])
        self.input_box.setFocus()

    def _on_attach_file(self):
        paths, _ = QFileDialog.getOpenFileNames(self, "选择文件")
        if not paths:
            return
        for p in paths:
            ext = os.path.splitext(p)[1].lower()
            if ext in (".wav", ".mp3", ".m4a", ".aac", ".flac", ".ogg"):
                self._attach_audio(p)
            else:
                # v4.66：把附件复制进工作区 incoming/ 目录，并把相对路径写进消息，
                # 这样 read_file 能按相对路径直接找到它。否则只剩文件名，模型会去
                # dist 目录瞎找 → "文件不存在"，进而用 run_command 反复搜、刷屏。
                # v4.102 hotfix：Windows 文件名禁止 \ / : * ? " < > | 等字符，用户原
                # 文件名（如「涉密岗用AI | 一定要死守的3道脱敏底线✅.png」含竖线）会
                # 导致 shutil.copy2 抛非法字符异常、复制失败、图根本没进 incoming，
                # 标记指向不存在的文件 → 视觉模型收不到图、静默无响应。复制前先清洗
                # 非法字符，并重名去重避免覆盖；标记里用清洗后的名字，保证能对应上。
                inc_dir = os.path.join(APP_DIR, "incoming")
                try:
                    os.makedirs(inc_dir, exist_ok=True)
                except Exception:
                    pass
                raw_base = os.path.basename(p)
                base = _sanitize_filename(raw_base) or "file"
                dst = os.path.join(inc_dir, base)
                # 重名去重：避免不同原文件清洗后同名互相覆盖
                if os.path.exists(dst):
                    stem, sufx = os.path.splitext(base)
                    i = 1
                    while os.path.exists(dst):
                        dst = os.path.join(inc_dir, f"{stem}_{i}{sufx}")
                        i += 1
                    base = os.path.basename(dst)
                try:
                    import shutil
                    shutil.copy2(p, dst)
                except Exception as e:
                    self.input_box.insertPlainText(
                        f"\n[文件: {base}]（复制失败：{e}，请确认文件可访问）\n")
                    self._on_input_changed()
                    continue
                self.input_box.insertPlainText(f"\n[文件: incoming/{base}]\n")
                self._on_input_changed()

    def _attach_audio(self, path):
        """用户添加音频文件：自动 ASR 转写并填入输入框（复用 _ASRWorker）。"""
        sf = self.cfg.get("siliconflow", {})
        if not sf or not sf.get("api_key"):
            self.input_box.insertPlainText(
                f"\n[音频文件: {os.path.basename(path)}]"
                f"（未配置硅基流动 key，无法自动转写，请在设置填写后重试）\n"
            )
            self._on_input_changed()
            return
        self.status_label.setText(f"🎧 音频识别中：{os.path.basename(path)}")
        worker = _ASRWorker(path, sf)
        worker.sig_text.connect(
            lambda txt, p=path: self._on_attach_asr_done(p, txt)
        )
        worker.sig_error.connect(
            lambda msg, p=path: self._on_attach_asr_error(p, msg)
        )
        worker.start()

    def _on_attach_asr_done(self, path, text):
        base = os.path.basename(path)
        if not text:
            self.input_box.insertPlainText(f"\n[音频文件: {base}]（未识别出语音内容）\n")
            self.status_label.setText("音频未识别出内容")
        else:
            self.input_box.insertPlainText(f"\n[音频转写: {base}]\n{text}\n")
            self.status_label.setText("音频已转写，可直接发送")
        self._on_input_changed()

    def _on_attach_asr_error(self, path, msg):
        self.input_box.insertPlainText(
            f"\n[音频文件: {os.path.basename(path)}]（转写失败：{msg}）\n"
        )
        self.status_label.setText("音频转写失败")
        self._on_input_changed()

    # ============ 会话切换 ============
    def _switch_session(self, sid):
        self.store.switch(sid)
        self._rendered_msg_count = 0  # v4.60：切换会话重置增量渲染计数
        self._refresh_session_combo()
        self._render_messages(force_bottom=True)
        self._refresh_skill_buttons()
        self._update_skill_bar()
        self._refresh_deliverables()
        self._refresh_recent_on_welcome()
        # 跳转到对话页并高亮导航
        self.main_stack.setCurrentIndex(1)
        self._update_nav_styles(0)
        self.input_box.setFocus()
        self._scan_agent_resume()  # v4.101：切到有暂停任务的会话时提示「继续」

    def _new_session(self):
        self.store.new_session()
        self._rendered_msg_count = 0  # v4.60：新建会话重置增量渲染计数
        self._refresh_session_combo()
        self._render_messages(force_bottom=True)
        self._refresh_skill_buttons()
        self._update_skill_bar()
        self._refresh_deliverables()
        self._refresh_recent_on_welcome()
        self.input_box.setFocus()

    def _close_session(self, sid):
        if self._busy:
            self.status_label.setText("正在处理，稍后再关闭会话")
            return
        self.store.remove(sid)
        # v4.73：删除会话时一并清理其独立的上下文摘要文件，避免残留串台
        try:
            import os as _os
            _ud = _os.path.expanduser("~/Documents/小臭玩AI")
            for _f in (f"context_summary_{sid}.json", f"key_info_{sid}.json"):
                _p = _os.path.join(_ud, _f)
                if _os.path.exists(_p):
                    _os.remove(_p)
        except Exception:
            pass
        self._refresh_session_combo()
        self._render_messages(force_bottom=True)
        self._refresh_skill_buttons()
        self._update_skill_bar()
        self._refresh_deliverables()
        self._refresh_recent_on_welcome()
        self.input_box.setFocus()

    def _delete_active_session(self):
        """删除当前对话：先确认，避免误删历史。"""
        sid = self.store.active_sid
        s = self.store.sessions.get(sid)
        if not s:
            return
        if self._busy:
            self.status_label.setText("正在处理，稍后再删除对话")
            return
        title = s.title or "新会话"
        n = len(self.store.sessions)
        extra = "（这是最后一个对话，将清空内容而非删除）" if n <= 1 else ""
        box = QMessageBox(self)
        box.setIcon(QMessageBox.Question)
        box.setWindowTitle("删除对话")
        box.setText(f"确定删除「{title}」吗？{extra}\n删除后不可恢复。")
        box.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        box.setDefaultButton(QMessageBox.No)
        # 中文按钮
        yes_btn = box.button(QMessageBox.Yes)
        no_btn = box.button(QMessageBox.No)
        if yes_btn:
            yes_btn.setText("删除")
        if no_btn:
            no_btn.setText("取消")
        if box.exec() == QMessageBox.Yes:
            self._close_session(sid)
            self.status_label.setText(f"已删除「{title}」")

    def _request_delete_session(self, sid, title):
        """删除指定对话（欢迎页最近列表用）：先确认。"""
        if self._busy:
            self.status_label.setText("正在处理，稍后再删除对话")
            return
        if not sid or sid not in self.store.sessions:
            return
        box = QMessageBox(self)
        box.setIcon(QMessageBox.Question)
        box.setWindowTitle("删除对话")
        box.setText(f"确定删除「{title}」吗？\n删除后不可恢复。")
        box.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        box.setDefaultButton(QMessageBox.No)
        yes_btn = box.button(QMessageBox.Yes)
        no_btn = box.button(QMessageBox.No)
        if yes_btn:
            yes_btn.setText("删除")
        if no_btn:
            no_btn.setText("取消")
        if box.exec() == QMessageBox.Yes:
            self._close_session(sid)
            self.status_label.setText(f"已删除「{title}」")

    def _rename_session(self):
        sid = self.store.active_sid
        s = self.store.sessions.get(sid)
        if not s:
            return
        dlg = RenameDialog(self)
        cur = s.title if s.title and s.title != "新会话" else ""
        dlg._edit.setText(cur)
        if dlg.exec():
            new_title = dlg.text_value()
            if new_title:
                s.title = new_title[:40]
                self.store.save()
                self._refresh_session_combo()
                self._refresh_recent_on_welcome()
                self.status_label.setText(f"已重命名为「{s.title}」")

    # ============ 聊天气泡渲染 ============
    def _fmt_bubble(self, role, text, idx=None):
        """v4.104：flex 布局气泡（Chromium 渲染，圆角/阴影/动画真实生效）。
        assistant 正文走 data-md 通道，由页面内 marked.js 渲染 Markdown（代码块/列表/表格）。
        idx：消息序号，assistant 气泡附带「重新生成/改写问题」操作链接（app:// 协议）。"""
        if role == "user":
            esc = html_mod.escape(text).replace("\n", "<br>")
            esc = re.sub(r'(https?://[^\s\u4e00-\u9fff，。、；：！？（）()【】<>"]+)',
                         r'<a href="\1" target="_blank">\1</a>', esc)
            if not esc.strip():
                esc = "…"
            return (
                '<div class="msg-row user">'
                '<div class="col"><div class="who">You</div>'
                f'<div class="bubble">{esc}</div></div>'
                f'{self._avatar_img_html("user")}'
                '</div>'
            )
        # assistant：Markdown 交给页面内 marked.js 渲染（HTML 属性转义即可，
        # 浏览器 getAttribute 解码回原文；属性内合法保留换行）
        raw = text if (text and str(text).strip()) else "…"
        md_attr = html_mod.escape(raw, quote=True)
        actions = ""
        if idx is not None:
            actions = (
                '<div class="actions">'
                f'<a href="app://regen/{idx}">🔄 重新生成</a>'
                f'<a href="app://edit/{idx}">✏️ 改写问题</a>'
                f'<span class="sep">|</span>'
                f'<a href="app://fb_up/{idx}">👍</a>'
                f'<a href="app://fb_down/{idx}">👎</a>'
                '</div>'
            )
        return (
            '<div class="msg-row ai">'
            f'{self._avatar_img_html("ai")}'
            '<div class="col"><div class="who">Agent</div>'
            f'<div class="bubble"><div class="md" data-md="{md_attr}"></div></div>'
            f'{actions}'
            '</div></div>'
        )

    def _hl(self, escaped):
        """v4.75：对话内搜索高亮——在转义后的 HTML 文本中，对标签外可见文本做大小写不敏感
        的 <mark> 包裹；标签（含 <a href>、<br>）原样保留，避免破坏链接/结构。"""
        q = getattr(self, "_search_query", "")
        if not q:
            return escaped
        try:
            qe = re.escape(html_mod.escape(q))
        except Exception:
            return escaped
        if not qe:
            return escaped
        parts = re.split(r'(<[^>]+>)', escaped)
        for i in range(0, len(parts), 2):  # 偶数段=可见文本
            parts[i] = re.sub(
                qe,
                r'<span style="background-color:#ffd54f;color:#000;'
                r'border-radius:2px;padding:0 1px;">\g<0></span>',
                parts[i], flags=re.IGNORECASE)
        return "".join(parts)

    def _wrap_msg(self, bubble, idx):
        """v4.75：给每条气泡外层包一个带 id + 命名锚点的 div，供对话内搜索 scrollToAnchor 跳转定位。"""
        if idx is None:
            return bubble
        return f'<div id="msg_{idx}"><a name="msg_{idx}"></a>{bubble}</div>'

    # ============ v4.75：对话内搜索 / 单条重生成 / 改写问题 ============
    def _extract_text(self, m):
        """从一条消息取出纯文本（兼容 str 与多模态 list 内容）。"""
        c = m.get("content", "")
        if isinstance(c, str):
            return c
        if isinstance(c, list):
            return "".join(p.get("text", "") for p in c
                           if isinstance(p, dict) and p.get("type") == "text")
        return ""

    def _on_anchor_clicked(self, url):
        """v4.75：处理气泡内的 regen:/edit: 链接；其余（http 等）交给系统浏览器。
        v4.108 H-06 修复：QWebEngineView 的链接信号传 str（QTextBrowser 时代是 QUrl），
        统一按 str 处理，恢复「重新生成/改写/👍/👎」四个气泡按钮。"""
        s = url if isinstance(url, str) else url.toString()
        if s.startswith("regen:"):
            try:
                self._regen_message(int(s[len("regen:"):]))
            except Exception:
                pass
            return
        if s.startswith("edit:"):
            try:
                self._edit_message(int(s[len("edit:"):]))
            except Exception:
                pass
            return
        if s.startswith("fb_up:"):
            try:
                self._on_feedback(int(s[len("fb_up:"):]), True)
            except Exception:
                pass
            return
        if s.startswith("fb_down:"):
            try:
                self._on_feedback(int(s[len("fb_down:"):]), False)
            except Exception:
                pass
            return
        QDesktopServices.openUrl(QUrl(s) if isinstance(s, str) else url)

    def _regen_message(self, idx):
        """v4.75：重新生成 idx 处的 assistant 回复。截断该条及之后所有消息，按原上文路由重跑。"""
        session = self.store.active()
        msgs = session.messages
        if not (0 <= idx < len(msgs)) or msgs[idx].get("role") != "assistant":
            return
        if self._busy:
            self.status_label.setText("上一条还在处理，稍等…")
            return
        prompt_idx = idx - 1
        while prompt_idx >= 0 and msgs[prompt_idx].get("role") != "user":
            prompt_idx -= 1
        if prompt_idx < 0:
            self.status_label.setText("找不到对应的上文，无法重生成")
            return
        del msgs[idx:]
        self.store.save()
        self._rendered_msg_count = 0
        self._render_messages()
        self._run_for_trigger(self._extract_text(msgs[prompt_idx]), "重新生成")

    def _edit_message(self, idx):
        """v4.75：改写 idx 处的用户问题。原文载入输入框并标记编辑目标，改完回车即重开该分支。"""
        session = self.store.active()
        msgs = session.messages
        if not (0 <= idx < len(msgs)) or msgs[idx].get("role") != "user":
            return
        if self._busy:
            self.status_label.setText("上一条还在处理，稍等…")
            return
        self._edit_target_idx = idx
        self.input_box.setPlainText(self._extract_text(msgs[idx]))
        self.input_box.setFocus()
        self._on_input_changed()
        self.status_label.setText("已载入该问题到输入框，修改后回车即重新生成此分支")

    # ============ v4.76：反馈闭环（👍/👎） ============
    def _feedback_path(self):
        return os.path.join(os.path.expanduser("~/Documents"), "小臭玩AI", "feedback.jsonl")

    def _record_feedback(self, idx, positive, reason=None):
        import json
        try:
            fp = self._feedback_path()
            os.makedirs(os.path.dirname(fp), exist_ok=True)
            sid = getattr(self.store.active(), "id", "")
            entry = {
                "ts": datetime.now().isoformat(timespec="seconds"),
                "session_id": sid,
                "idx": idx,
                "rating": "up" if positive else "down",
                "reason": reason or "",
            }
            with open(fp, "a", encoding="utf-8") as f:
                f.write(json.dumps(entry, ensure_ascii=False) + "\n")
        except Exception as e:
            log.warning("记录反馈失败: %s", e)

    def _refresh_feedback_count(self):
        import json
        try:
            fp = self._feedback_path()
            if not os.path.exists(fp):
                return
            up = down = 0
            with open(fp, encoding="utf-8") as f:
                for line in f:
                    try:
                        d = json.loads(line)
                        if d.get("rating") == "up":
                            up += 1
                        elif d.get("rating") == "down":
                            down += 1
                    except Exception:
                        pass
            self.status_label.setText(f"反馈统计：👍 {up} · 👎 {down}")
        except Exception:
            pass

    def _on_feedback(self, idx, positive):
        """v4.76：记录用户对某条助手回复的反馈，形成闭环。👎 可补充原因并记入长期偏好。"""
        session = self.store.active()
        msgs = session.messages
        if not (0 <= idx < len(msgs)):
            return
        snippet = self._extract_text(msgs[idx])[:40].replace("\n", " ")
        self._record_feedback(idx, positive)
        if positive:
            self.status_label.setText(f"👍 已记录你的喜欢（{snippet}…）")
        else:
            from PySide6.QtWidgets import QInputDialog
            reason, ok = QInputDialog.getText(
                self, "不太满意？",
                "可补充一句原因（选填），将记入长期偏好以改进后续回复：")
            if ok and reason.strip():
                self._record_feedback(idx, positive, reason.strip())
                try:
                    from memory_store import append_memory
                    append_memory(
                        f"[反馈偏好] 用户不喜欢：{reason.strip()}",
                        topic="用户反馈偏好", category="feedback")
                except Exception:
                    pass
                self.status_label.setText(f"👎 已记录，并记入偏好：{reason.strip()}")
            else:
                self.status_label.setText("👎 已记录，谢谢反馈")
        self._refresh_feedback_count()

    def _run_for_trigger(self, trigger, tag=""):
        """v4.75：复用 send() 的自动路由，用触发词触发一次生成（不新增 user 消息）。"""
        use_agent = self.agent_mode or self._message_needs_agent(trigger)
        advice_only = self._message_is_advice_only(trigger)
        if use_agent and advice_only:
            use_agent = False
        statement_only = self._message_is_statement_only(trigger)
        topic_only = self._message_is_topic_only(trigger)
        self._busy = True
        self.send_btn.setEnabled(False)
        self._busy_timeout.start(120000)
        tag_suf = f"·{tag}" if tag else ""
        if use_agent and not statement_only:
            self.status_label.setText(f"Agent 启动…{tag_suf}")
            self._agent_run()
        elif statement_only:
            self.status_label.setText(f"（日常闲聊）{tag_suf}")
            self._start_stream(trigger, None)
        elif advice_only:
            self.status_label.setText(f"（给建议中…）{tag_suf}")
            self._start_stream(trigger, None)
        elif topic_only:
            self.status_label.setText(f"选题生成中…（用 AI 常识列方向）{tag_suf}")
            self._start_stream(trigger, None)
        elif self.cfg.get("search_enabled", True):
            self.status_label.setText(f"搜索中…{tag_suf}")
            self._do_search(trigger)
        else:
            self._start_stream(trigger, None)

    def _search_in_chat(self):
        """v4.75：对话内搜索——高亮全部匹配并跳到首个。
        v4.104：高亮由页面内 JS 完成（jsHighlight 遍历文本节点包 <mark>），
        无需再重渲染整个对话。"""
        q = self.search_box.text().strip()
        self._search_query = q
        self.chat_view.highlight(q)  # JS 侧高亮/清除，不重建 DOM
        if not q:
            self._search_matches = []
            self._search_pos = -1
            self.status_label.setText("")
            return
        ql = q.lower()
        session = self.store.active()
        matches = [i for i, m in enumerate(session.messages)
                   if ql in self._extract_text(m).lower()]
        self._search_matches = matches
        if matches:
            self._jump_to_match(0)
        else:
            self._search_pos = -1
            self.status_label.setText(f"未找到「{q}」")

    def _jump_to_match(self, pos):
        if not self._search_matches:
            return
        n = len(self._search_matches)
        pos %= n
        self._search_pos = pos
        self.chat_view.scroll_to(self._search_matches[pos])
        self.status_label.setText(
            f"找到 {n} 处匹配（{pos + 1}/{n}）")

    def _search_nav(self, delta):
        if not self._search_matches:
            self._search_in_chat()
            return
        self._jump_to_match(self._search_pos + delta)

    def eventFilter(self, obj, ev):
        """v4.75：搜索框 Enter=搜索/下一处，Shift+Enter=上一处，Esc=清除。"""
        from PySide6.QtCore import QEvent
        if obj is getattr(self, "search_box", None) and ev.type() == QEvent.KeyPress:
            key = ev.key()
            if key in (Qt.Key_Return, Qt.Key_Enter):
                self._search_nav(1 if not (ev.modifiers() & Qt.ShiftModifier) else -1)
                return True
            if key == Qt.Key_Escape:
                self.search_box.clear()
                self._search_in_chat()
                return True
        return super().eventFilter(obj, ev)

    def _save_enc_passphrase(self):
        """v4.75：设置/修改/关闭记忆加密口令，并迁移已有记忆。"""
        pw = self.enc_pw_edit.text()
        prev = self.cfg.get("memory_encryption_passphrase", "")
        self.cfg["memory_encryption_passphrase"] = pw
        self._save_cfg()
        try:
            from memory_store import encrypt_existing, decrypt_existing, set_encryption
            if pw:
                if prev:
                    decrypt_existing()  # 改口令：先用旧密钥解密
                encrypt_existing(pw)
                set_encryption(pw)
                self.status_label.setText(
                    "记忆加密已开启" + ("（口令已更新）" if prev else "，已有记忆已加密迁移"))
            else:
                if prev:
                    decrypt_existing()
                    set_encryption("")
                    self.status_label.setText("记忆加密已关闭，已解密为明文")
                else:
                    self.status_label.setText("记忆加密未启用（未填写口令）")
        except Exception as e:
            self.status_label.setText(f"加密操作失败：{e}")

    # ============ v4.76：自动备份 / 版本更新 处理器 ============
    def _save_autobackup_settings(self):
        """保存自动备份计划：写入 config，并用 schtasks 创建/更新/删除系统任务。"""
        import subprocess
        freq = self.ab_freq_combo.currentData()
        t = self.ab_time_edit.time().toString("HH:mm")
        self.cfg["autobackup_freq"] = freq
        self.cfg["autobackup_time"] = t
        self._save_cfg()
        exe = sys.executable
        task = "小臭玩AI自动备份"
        if not freq:
            try:
                subprocess.run(["schtasks", "/Delete", "/TN", task, "/F"],
                               capture_output=True, text=True, timeout=20)
            except Exception:
                pass
            self.status_label.setText("已关闭自动备份（系统计划任务已移除）")
            return
        cmd = f'"{exe}" --autobackup'
        try:
            r = subprocess.run(
                ["schtasks", "/Create", "/TN", task, "/TR", cmd,
                 "/SC", freq.upper(), "/ST", t, "/F", "/RL", "HIGHEST"],
                capture_output=True, text=True, timeout=30)
            ok = r.returncode == 0
            msg = (r.stdout or r.stderr or "").strip()
        except Exception as e:
            ok, msg = False, str(e)
        if ok:
            self.status_label.setText(f"自动备份已设置：{freq} {t}（系统任务计划程序）")
        else:
            self.status_label.setText(f"自动备份设置失败：{msg}")

    def _run_backup_now(self):
        """立即执行一次备份（调用本 exe 的 --autobackup）。"""
        import subprocess
        try:
            r = subprocess.run([sys.executable, "--autobackup"],
                               capture_output=True, text=True, timeout=60)
            msg = (r.stdout or r.stderr or "备份完成").strip()
        except Exception as e:
            msg = f"备份失败：{e}"
        QMessageBox.information(self, "立即备份", msg)
        self.status_label.setText(msg)

    def _check_update(self):
        """v4.76：检查更新——本地构建说明 或 访问 update_check_url 比对版本。"""
        import json
        from config import APP_VERSION, APP_BUILD_DATE, UPDATE_CHECK_URL
        url = self.cfg.get("update_check_url", "") or UPDATE_CHECK_URL
        if not url:
            QMessageBox.information(
                self, "检查更新",
                f"当前版本：{APP_VERSION}（构建 {APP_BUILD_DATE}）\n\n"
                "本程序为本地构建版本，无在线更新通道。\n"
                "如需更新，请联系构建者重新打包新版即可。")
            return
        try:
            import urllib.request
            req = urllib.request.Request(url, headers={"User-Agent": "小臭玩AI"})
            with urllib.request.urlopen(req, timeout=8) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            latest = data.get("version", "")
            if not latest or latest == APP_VERSION:
                QMessageBox.information(self, "检查更新", f"已是最新版本：{APP_VERSION}")
            else:
                dlg = QMessageBox(self)
                dlg.setWindowTitle("发现新版本")
                dlg.setText(f"发现新版本：{latest}\n当前：{APP_VERSION}\n\n"
                           f"{data.get('notes', '')}")
                dlg.setInformativeText("是否打开下载页？")
                dlg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
                if dlg.exec() == QMessageBox.Yes and data.get("url"):
                    QDesktopServices.openUrl(QUrl(data["url"]))
        except Exception as e:
            QMessageBox.warning(
                self, "检查更新",
                f"检查失败：{e}\n（网络受限时属正常，本地构建无需在线更新）")

    def _run_perf(self, save_baseline=False):
        """v4.78：后台跑性能基线，避免阻塞 UI。"""
        if getattr(self, "_perf_worker", None) and self._perf_worker.isRunning():
            return
        w = PerfWorker(save_baseline=save_baseline)
        w.done.connect(self._on_perf_done)
        self._perf_worker = w
        w.start()

    def _on_perf_done(self, res):
        """v4.78：性能基线结果弹窗（标注 REGRESSION / IMPROVED）。"""
        from PySide6.QtWidgets import QMessageBox
        self._perf_worker = None
        if res.get("error"):
            QMessageBox.warning(self, "性能基线", "运行失败：%s" % res["error"])
            return
        m = res["metrics"]; cmp = res["comparison"]; verdict = res["verdict"]
        lines = []
        if res.get("saved_baseline"):
            lines.append("已保存为新基线")
        lines.append("判定：%s" % verdict)
        lines.append("")
        for k, v in m.items():
            if k in ("ts", "frozen"):
                continue
            lines.append("%s = %s" % (k, v))
        if cmp:
            lines.append("")
            lines.append("— 与基线比对 —")
            for k, c in cmp.items():
                if c.get("status") == "n/a":
                    lines.append("%s：%s（无基线值）" % (k, c.get("value")))
                else:
                    lines.append("%s：%s (%+g%%) %s" % (k, c["value"], c.get("delta_pct") or 0, c["status"]))
        else:
            lines.append("")
            lines.append("（无基线可比对，建议先『设为当前基线』）")
        lines.append("")
        lines.append("记录：%s" % res.get("run_path", ""))
        QMessageBox.information(self, "性能基线报告", "\n".join(lines))

    def _ensure_avatars(self):
        """v4.97：生成 26×26 圆形头像 PNG 到 APP_DIR/avatars/。AI=🤖蓝紫、User=👤灰。
        气泡 HTML 用 <img> 引用。文件已存在且版本标记匹配则跳过；版本升级时强制刷新。
        生成失败不抛异常（降级到无头像）。"""
        try:
            d = os.path.join(APP_DIR, "avatars")
            os.makedirs(d, exist_ok=True)
            ver_file = os.path.join(d, ".version")
            # 版本升级时强制删除旧头像重新生成
            if os.path.isfile(ver_file):
                with open(ver_file, "r", encoding="utf-8") as f:
                    cached_ver = f.read().strip()
            else:
                cached_ver = ""
            if cached_ver != APP_VERSION:
                for old in ("avatar_ai.png", "avatar_user.png"):
                    p = os.path.join(d, old)
                    if os.path.isfile(p):
                        os.remove(p)
                with open(ver_file, "w", encoding="utf-8") as f:
                    f.write(APP_VERSION)
            ai_path = os.path.join(d, "avatar_ai.png")
            user_path = os.path.join(d, "avatar_user.png")
            if os.path.isfile(ai_path) and os.path.isfile(user_path):
                return
            for path, bg, label in (
                (ai_path, "#6366F1", "🤖"),   # Indigo 柔和蓝紫
                (user_path, "#6B7280", "👤"),  # 现代灰
            ):
                if os.path.isfile(path):
                    continue
                pm = QPixmap(26, 26)
                pm.fill(Qt.transparent)
                p = QPainter(pm)
                p.setRenderHint(QPainter.Antialiasing)
                p.setBrush(QColor(bg))
                p.setPen(Qt.NoPen)
                # v4.97：正圆（半径=13，刚好 26×26）
                p.drawEllipse(0, 0, 26, 26)
                p.setPen(QColor("#FFFFFF"))
                font = QFont("Segoe UI Emoji", 14)
                # fallback：如果 Segoe UI Emoji 不可用，用默认字体
                if not QFont("Segoe UI Emoji").exactMatch():
                    font = QFont()
                    font.setPixelSize(14)
                p.setFont(font)
                p.drawText(QRect(0, 0, 26, 26), Qt.AlignCenter, label)
                p.end()
                pm.save(path, "PNG")
        except Exception as e:
            log.warning("生成头像 PNG 失败：%s", e)

    def _avatar_img_html(self, who):
        """生成头像 <img> 标签 HTML。who='ai'|'user'。文件不存在返回空串（降级无头像）。"""
        name = "avatar_ai.png" if who == "ai" else "avatar_user.png"
        p = os.path.join(APP_DIR, "avatars", name)
        if not os.path.isfile(p):
            return ""
        return ('<img src="file:///' + p.replace(os.sep, "/") + '" '
                'width="26" height="26" '
                'style="vertical-align:top;border-radius:13px;display:block;">')

    def _on_scroll_changed(self, _value):
        """v4.104：滚动跟随已迁移到页面内 JS（scroll 监听 + 粘性 stick），
        此方法保留为空实现以防旧引用；Python 侧不再维护滚动条哨兵。"""
        pass

    def _on_chat_page_reloaded(self):
        """v4.104：聊天页意外重载后 DOM 已清空 → 重置计数全量重渲染。"""
        self._rendered_msg_count = 0
        self._render_messages()

    def _request_scroll_bottom(self):
        """v4.104：滚底走 JS（jsForceBottom：置跟随 + 滚底一次到位），
        Chromium 布局同步完成，不再需要 Qt 异步 layout 的四次重试。"""
        if not self._follow_bottom:
            return
        self.chat_view.force_bottom()

    def _render_throttled(self, force_bottom=False):
        """v4.58：节流版渲染——50ms 内多次调用合并为一次，防工具调用/流式渲染风暴。"""
        self._pending_render_force = self._pending_render_force or force_bottom
        if not self._render_timer.isActive():
            self._render_timer.start(50)

    def _on_render_tick(self):
        """节流定时器触发的实际渲染。"""
        self._render_messages(self._pending_render_force)
        self._pending_render_force = False

    def _flush_render(self):
        """v4.60o：冲刷可能仍在排队的节流渲染，并保证末条消息不漏渲染、不重复。

        根因：_render_throttled 把实际渲染推迟到 50ms 后的 QTimer，而 _on_agent_done /
        _on_stream_finished 又同步调用 _render_messages。若最后一次节流还没触发、
        _rendered_msg_count 滞后，最终同步渲染会把末 1~2 条消息重复插入 —— 表现为
        同一个回答气泡出现两三次。
        处理：① 若有待触发定时器，停掉并立即 tick（渲染全部新消息、更新计数）；
        ② 兜底：若仍有未渲染消息（如 _on_stream_finished 刚追加的末条），补一次
        增量渲染。两步都按 _rendered_msg_count 增量提交，不重复、不遗漏。
        """
        if self._render_timer.isActive():
            self._render_timer.stop()
            self._on_render_tick()
        if self._rendered_msg_count < len(self.store.active().messages):
            self._render_messages()

    def _save_throttled(self):
        """v4.58：节流版 store.save——1 秒内多次触发合并为一次磁盘写入。"""
        if not self._save_timer.isActive():
            self._save_timer.start(1000)

    def _do_save(self):
        """批量保存定时器触发的实际写入。"""
        try:
            self.store.save()
        except Exception:
            pass

    def _render_messages(self, force_bottom=False):
        session = self.store.active()
        has_content = bool(session.messages) or self._streaming
        if has_content and self.main_stack.currentIndex() != 1:
            self.main_stack.setCurrentIndex(1)
        elif not has_content and self.main_stack.currentIndex() != 0:
            self.main_stack.setCurrentIndex(0)

        # v4.104：WebEngine 增量渲染——新消息 insertAdjacentHTML 追加（零重排），
        # 流式只替换 #stream-bubble 的 innerHTML；全量重建仅在会话切换/重生成时发生。
        msgs = session.messages
        rendered = getattr(self, "_rendered_msg_count", 0)
        if rendered == 0 or not msgs or rendered > len(msgs):
            # 首次渲染 / 会话切换 / 重生成 → 全量重建
            parts = self._build_parts(msgs)
            # v4.104 fix：全量重建也要带上流式内容，否则流式首帧（count=0 时）
            # 不显示，等第一帧文本来了才冒出来，观感突兀。
            if self._streaming and self._streaming_text:
                # v4.108 M-25：流式部件带 #stream-bubble id（与 jsStream 创建结构一致），
                # commit 后 end_stream 能移除——否则无 id 残留与增量 append 正文形成
                # 同段双气泡（pageReloaded 自愈路径曾触发）。
                parts.append(
                    '<div class="msg-row ai" id="stream-bubble">'
                    + self._fmt_bubble("assistant", self._streaming_text)
                    + '</div>')
            self.chat_view.render_all("".join(parts))
            self._rendered_msg_count = len(msgs)
            # v4.108 M-25：jsRenderAll 清空了 DOM，live 工具卡（不进 messages）随之消失。
            # Agent 仍运行（页面自愈/切换回会话的全量重建）时按台账恢复 running/done 卡，
            # 否则 replace_live 找不到目标、工具过程在界面上直接蒸发。
            if (getattr(self, "_agent_active", False)
                    and getattr(self, "_live_tools", None)):
                for _rec in sorted(self._live_tools.values(),
                                   key=lambda r: r.get("ui_seq", 0)):
                    if _rec.get("status") == "done":
                        self.chat_view.append(self._tool_card_html(
                            _rec["ui_seq"], _rec.get("display_name", ""),
                            done=True, result_preview=_rec.get("result", "")))
                    else:
                        self.chat_view.append(self._tool_card_html(
                            _rec["ui_seq"], _rec.get("display_name", ""),
                            _rec.get("args_preview", "")))
        else:
            # 增量追加：只渲染渲染计数之后的新消息
            new_msgs = msgs[rendered:]
            if new_msgs:
                agent_live = getattr(self, "_agent_active", False)
                for i, m in enumerate(new_msgs, start=rendered):
                    # Agent 运行期间 tool_log 不从消息管线渲染——实时 running/done 卡
                    # 由 _on_tool_started/_on_tool_finished 直接注入（带动画），
                    # 避免同一工具调用出现 2~3 张重复卡（旧 UI 的视觉噪音来源之一）。
                    if agent_live and m.get("role") == "tool_log":
                        continue
                    bubble = self._fmt_single_message(m, i)
                    if bubble:
                        self.chat_view.append(bubble)
                self._rendered_msg_count = len(msgs)
            # 流式文本：只更新流式气泡
            if self._streaming:
                self._replace_last_streaming()
            else:
                # v4.104 fix：增量路径必须清残留流式气泡——否则「流式结束走早退分支
                # / agent 收尾 flush」等 count 未重置场景下，最终回答气泡 + #stream-bubble
                # 残留 = 同一段话重复 2 个气泡。
                self.chat_view.end_stream()

        if force_bottom:
            self._follow_bottom = True
        if self._follow_bottom:
            self._request_scroll_bottom()

    def _render_director_redirect(self, text):
        """v4.107：主对话框里的导演指令不进主控 Agent、不写主会话（零交集），
        只在聊天区渲染一条引导气泡，提示去导演台底部「导演对话」条下指令。
        该气泡不入库——下次渲染主会话会从 session 重建，引导自然消失（一次性提示）。"""
        bubble = self._fmt_bubble(
            "assistant",
            "🎬 这条像给导演台的指令（「%s」）。导演台已独立出来，请在"
            "**导演台底部「导演对话」条**里直接下指令；主对话框不处理导演相关操作。"
            % text)
        if self.main_stack.currentIndex() != 1:
            self.main_stack.setCurrentIndex(1)
        self.chat_view.append(bubble)
        self._request_scroll_bottom()

    def _replace_last_streaming(self):
        """v4.104：流式更新——只替换 #stream-bubble 的 innerHTML（页面内局部重排），
        其余消息 DOM 不动。旧 QTextBrowser 时代需 setHtml 全量重建 + 滚动比例恢复，
        是「越聊越卡 + 滚动条跳动」的根因，现已根治。"""
        if not self._streaming_text:
            return
        bubble = self._fmt_bubble("assistant", self._streaming_text)
        self.chat_view.update_stream(bubble)

    def _build_parts(self, msgs):
        """构建完整 HTML 片段列表（仅 setHtml 时使用）。"""
        parts = []
        for i, m in enumerate(msgs):
            bubble = self._fmt_single_message(m, i)
            if bubble:
                parts.append(bubble)
        return parts

    def _fmt_single_message(self, m, idx=None):
        """v4.60：格式化单条消息为 HTML 片段，供增量渲染复用。
        idx：消息序号，传入后外层包 <div id="msg_{idx}"> 供搜索跳转定位。"""
        # v4.60o：内部注入指令（如"记住能力"时要求调 remember）不渲染成气泡
        if m.get("_internal"):
            return ""
        if m.get("role") == "tool_log":
            name = m.get("name", "")
            result = m.get("result", "")
            if name == "image_gen" and re.search(r"\.(png|jpe?g|webp|gif)$", result, re.I):
                img_abs = os.path.join(APP_DIR, result).replace("/", os.sep)
                if os.path.isfile(img_abs):
                    bubble = (
                        '<div class="tool-wrap">'
                        f'<img src="file:///{img_abs.replace(os.sep, "/")}" '
                        f'style="max-width:320px;border-radius:10px;"/>'
                        f'<div style="font-size:11px;color:{THEME["faint"]};margin-top:4px;">'
                        f'{html_mod.escape(result)}</div></div>'
                    )
                    return self._wrap_msg(bubble, idx)
            card = tools_mod.card_html(name, m.get("args", ""), result)
            bubble = f'<div class="tool-wrap">{card}</div>'
            return self._wrap_msg(bubble, idx)
        else:
            raw = m.get("content", "")
            if isinstance(raw, list):
                # v4.104：多模态（文本+图片）——flex 标记，文本部分同旧逻辑转义+换行
                mm_parts = []
                for part in raw:
                    if isinstance(part, dict) and part.get("type") == "text":
                        t = part.get("text", "")
                        if t:
                            mm_parts.append(html_mod.escape(t).replace("\n", "<br>"))
                    elif isinstance(part, dict) and part.get("type") == "image_url":
                        url = part.get("image_url", {}).get("url", "")
                        if url:
                            mm_parts.append(
                                '<br><img src="' + html_mod.escape(url) + '" '
                                'style="max-width:300px;border-radius:10px;margin:6px 0;"><br>'
                            )
                inner = "".join(mm_parts) if mm_parts else "[图片消息]"
                if m.get("role") == "user":
                    bubble = (
                        '<div class="msg-row user">'
                        '<div class="col"><div class="who">You</div>'
                        f'<div class="bubble">{inner}</div></div>'
                        f'{self._avatar_img_html("user")}'
                        '</div>'
                    )
                else:
                    bubble = (
                        '<div class="msg-row ai">'
                        f'{self._avatar_img_html("ai")}'
                        '<div class="col"><div class="who">Agent</div>'
                        f'<div class="bubble">{inner}</div></div>'
                        '</div>'
                    )
                return self._wrap_msg(bubble, idx)
            # v4.57：空正文幽灵跳过
            if (m.get("role") == "assistant"
                    and not raw
                    and not m.get("tool_calls")):
                return ""
            return self._wrap_msg(self._fmt_bubble(m.get("role"), raw, idx), idx)

    # ============ 交付物面板 (260px) ============
    def _build_deliverables(self):
        self.deliverables = QWidget()
        self.deliverables.setFixedWidth(260)
        self.deliverables.setStyleSheet(
            f"QWidget#dvWidget{{background-color:{THEME['sidebar']};border:none;}}")
        self.deliverables.setObjectName("dvWidget")

        dv = QVBoxLayout(self.deliverables)
        dv.setContentsMargins(20, 24, 20, 12)
        dv.setSpacing(10)

        # 标题行
        head = QHBoxLayout()
        title = QLabel("交付物")
        title.setStyleSheet(
            f"font-size:14px;font-weight:600;color:{THEME['text']};background:transparent;")
        head.addWidget(title)
        head.addStretch(1)
        self.dv_count_label = QLabel("0")
        ut = THEME["user_text"]
        ac = THEME["accent"]
        self.dv_count_label.setStyleSheet(
            f"font-size:12px;color:{ut};background:{ac};"
            f"border-radius:10px;padding:1px 8px;font-weight:500;")
        head.addWidget(self.dv_count_label)
        dv.addLayout(head)

        # 分隔
        sep = QFrame()
        sep.setFrameShape(QFrame.HLine)
        sep.setStyleSheet(f"QFrame{{color:{THEME['separator']};max-height:1px;}}")
        dv.addWidget(sep)

        # 列表
        self.dv_scroll = QScrollArea()
        self.dv_scroll.setWidgetResizable(True)
        self.dv_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.dv_scroll.setStyleSheet("QScrollArea{background:transparent;border:none;}")
        self.dv_list = QWidget()
        self.dv_list.setStyleSheet("background:transparent;")
        self.dv_list_layout = QVBoxLayout(self.dv_list)
        self.dv_list_layout.setContentsMargins(0, 0, 0, 0)
        self.dv_list_layout.setSpacing(6)
        self.dv_list_layout.addStretch(1)
        self.dv_scroll.setWidget(self.dv_list)
        dv.addWidget(self.dv_scroll, 1)

        self.dv_empty_hint = QLabel("工具生成的文件会出现在这里\n（生图 / 写文件 / 跑代码产出等）")
        self.dv_empty_hint.setStyleSheet(
            f"color:{THEME['faint']};font-size:11px;line-height:1.5;padding:6px 4px 20px 4px;"
            f"background:transparent;")
        self.dv_empty_hint.setWordWrap(True)
        dv.addWidget(self.dv_empty_hint)
        dv.addSpacing(8)

        # 底部按钮（参考设计稿 .right-actions）
        # v4.86 修复：打开「真正的产物文件夹」= config.PRODUCTS_DIR
        # （~/Documents/小臭玩AI/产物，图片/截图/视频/workspace 统一落点），
        # 而非小臭程序目录 APP_DIR。
        open_dir_btn = QPushButton("📂  打开产物文件夹")
        open_dir_btn.setFixedHeight(38)
        open_dir_btn.setStyleSheet(
            f"QPushButton{{"
            f"background:transparent;"
            f"border:1.5px solid {THEME['border']};"
            f"border-radius:10px;"
            f"color:{THEME['dim']};font-size:13px;font-weight:500;"
            f"padding:0 14px;text-align:left;"
            f"}}"
            f"QPushButton:hover{{background:{THEME['panel2']};color:{THEME['text']};"
            f"border-color:{THEME['border_highlight']};}}"
        )
        open_dir_btn.clicked.connect(self._on_open_products_dir)
        dv.addWidget(open_dir_btn)

        clear_btn = QPushButton("🗑  清空")
        clear_btn.setFixedHeight(38)
        clear_btn.setStyleSheet(
            f"QPushButton{{"
            f"background:transparent;"
            f"border:1.5px solid {THEME['border']};"
            f"border-radius:10px;"
            f"color:{THEME['danger']};font-size:13px;font-weight:500;"
            f"padding:0 14px;text-align:left;"
            f"}}"
            f"QPushButton:hover{{background:rgba(239,68,68,0.1);"
            f"border-color:{THEME['danger']};}}"
        )
        clear_btn.clicked.connect(self._on_deliverables_clear)
        dv.addWidget(clear_btn)

        # 收起按钮
        collapse_btn = QPushButton("收起面板")
        collapse_btn.setFixedHeight(28)
        collapse_btn.setStyleSheet(
            f"QPushButton{{color:{THEME['faint']};font-size:11px;"
            f"border:none;border-radius:4px;background:transparent;}}"
            f"QPushButton:hover{{background:{THEME['sidebar_hover']};color:{THEME['text']};}}")
        collapse_btn.clicked.connect(self._toggle_deliverables)
        dv.addWidget(collapse_btn)

    def _toggle_sidebar(self):
        if not hasattr(self, "sb_expand_btn"):
            return
        visible = self.sidebar.isVisible()
        self.sidebar.setVisible(not visible)
        self.sb_expand_btn.setVisible(visible)

    def _toggle_deliverables(self):
        visible = self.deliverables.isVisible()
        self.deliverables.setVisible(not visible)
        self.dv_expand_btn.setVisible(visible)

    def _refresh_deliverables(self):
        while self.dv_list_layout.count() > 1:
            item = self.dv_list_layout.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
        items = self.store.active().deliverables
        self.dv_count_label.setText(str(len(items)))
        self.dv_empty_hint.setVisible(len(items) == 0)
        for d in items:
            rel = d.get("rel", "")
            kind = d.get("kind", "file")
            name = d.get("name", os.path.basename(rel))
            desc = d.get("desc", rel)

            color_key, tag_color_key = DELIVERY_COLORS.get(kind, ("delivery_blue", "accent"))
            border_color = THEME.get(color_key, THEME["accent"])
            tag_bg = THEME.get(tag_color_key, THEME["accent"])

            # 交付物卡片（参考设计稿 .delivery-item）
            card_item = QWidget()
            card_item.setCursor(Qt.PointingHandCursor)
            card_item.setStyleSheet(
                f"QWidget#dvCard{{"
                f"background:{THEME['surface_raised']};"
                f"border-radius:10px;"
                f"border-left:3px solid {border_color};"
                f"}}")
            card_item.setObjectName("dvCard")

            ci_lay = QVBoxLayout(card_item)
            ci_lay.setContentsMargins(14, 14, 14, 12)
            ci_lay.setSpacing(4)

            name_lbl = QLabel(name if len(name) <= 24 else name[:23] + "…")
            name_lbl.setStyleSheet(
                f"font-size:13px;font-weight:600;color:{THEME['text']};background:transparent;")
            ci_lay.addWidget(name_lbl)

            if desc:
                desc_text = desc if len(desc) <= 40 else desc[:39] + "…"
                desc_lbl = QLabel(desc_text)
                desc_lbl.setWordWrap(True)
                desc_lbl.setStyleSheet(
                    f"font-size:12px;color:{THEME['dim']};line-height:1.4;background:transparent;")
                ci_lay.addWidget(desc_lbl)

            tag_lbl = QLabel(kind.upper() if len(kind) <= 8 else kind[:7].upper())
            ut2 = THEME["user_text"]
            tag_lbl.setStyleSheet(
                f"font-size:10px;color:{ut2};background:{tag_bg};"
                f"border-radius:4px;padding:1px 6px;")
            tag_lbl.setFixedHeight(16)
            ci_lay.addWidget(tag_lbl)

            card_item.mousePressEvent = lambda e, r=rel: self._on_deliverable_open(r)
            self.dv_list_layout.insertWidget(self.dv_list_layout.count() - 1, card_item)

    def _on_deliverable_added(self, rel, kind, name):
        sess = self.store.active()
        if any(d.get("rel") == rel for d in sess.deliverables):
            return
        sess.deliverables.append({
            "rel": rel, "kind": kind, "name": name,
            "time": datetime.now().strftime("%H:%M:%S"),
        })
        self.store.save()
        self._refresh_deliverables()

    def _on_deliverable_open(self, rel):
        abs_path = os.path.abspath(os.path.join(APP_DIR, rel))
        if os.path.exists(abs_path):
            QDesktopServices.openUrl(QUrl.fromLocalFile(abs_path))
        else:
            self.status_label.setText(f"文件不存在：{abs_path}")

    def _on_deliverables_clear(self):
        if self._busy:
            self.status_label.setText("正在处理，稍后再清空")
            return
        sess = self.store.active()
        if not sess.deliverables:
            return
        sess.deliverables = []
        self.store.save()
        self._refresh_deliverables()

    def _on_open_products_dir(self):
        """v4.86 修复：打开真正的产物文件夹 PRODUCTS_DIR
        （~/Documents/小臭玩AI/产物），而非 APP_DIR。
        生成类工具（生图/截图/视频/跑代码）统一落此目录。"""
        pdir = getattr(config, "PRODUCTS_DIR", None) or os.path.join(
            os.path.expanduser("~"), "Documents", "小臭玩AI", "产物")
        try:
            os.makedirs(pdir, exist_ok=True)
        except Exception:
            pass
        if os.path.exists(pdir):
            QDesktopServices.openUrl(QUrl.fromLocalFile(pdir))
        else:
            self.status_label.setText(f"产物目录不可用：{pdir}")

    # ============ 弹层构建 ============
    def _popup_base(self, width):
        popup = QWidget(self, Qt.Popup | Qt.FramelessWindowHint)
        popup.setFixedWidth(width)
        popup.setStyleSheet(
            f"QWidget{{background:{THEME['card']};"
            f"border:1px solid {THEME['border_highlight']};border-radius:10px;}}"
            f"QLabel{{color:{THEME['text']};background:transparent;}}"
            f"QPushButton{{border:none;background:transparent;color:{THEME['text']};}}"
            f"QPushButton:hover{{background:rgba(10,10,12,0.04);}}"
            f"QScrollArea{{background:transparent;border:none;}}"
            f"QComboBox{{background:{THEME['elev']};border:none;border-radius:6px;"
            f"padding:6px 8px;font-size:12px;color:{THEME['text']};}}"
            f"QComboBox QAbstractItemView{{background:{THEME['card']};border:none;"
            f"color:{THEME['text']};selection-background-color:{THEME['accent']};}}"
            f"QLineEdit{{background:{THEME['elev']};border:1px solid {THEME['border']};"
            f"border-radius:6px;padding:6px 8px;color:{THEME['text']};font-size:12px;}}"
            f"QLineEdit:focus{{border-color:{THEME['accent']};}}"
            f"QCheckBox{{font-size:12px;color:{THEME['dim']};}}"
            f"QCheckBox::indicator{{width:16px;height:16px;}}")
        layout = QVBoxLayout(popup)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(10)
        return popup, layout

    def _popup_title(self, layout, text):
        t = QLabel(text)
        t.setStyleSheet(
            f"font-size:14px;font-weight:700;color:{THEME['text']};padding-bottom:4px;")
        layout.addWidget(t)

    def _build_skill_popup(self):
        popup, layout = self._popup_base(260)
        self._popup_title(layout, "技能库")
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setFixedHeight(360)
        inner = QWidget()
        inner_layout = QVBoxLayout(inner)
        inner_layout.setContentsMargins(0, 0, 0, 0)
        inner_layout.setSpacing(4)
        scroll.setWidget(inner)
        layout.addWidget(scroll)
        return popup, inner_layout

    def _build_settings_popup(self):
        popup, layout = self._popup_base(300)
        self._popup_title(layout, "设置")

        # ---- 模型组 ----
        group_model = QFrame()
        group_model.setStyleSheet(
            f"QFrame{{background:{THEME['elev']};border:none;border-radius:10px;padding:12px;}}")
        gml = QVBoxLayout(group_model)
        gml.setContentsMargins(12, 12, 12, 12)
        gml.setSpacing(6)

        gmt = QLabel("模型选择")
        gmt.setStyleSheet(f"font-size:12px;font-weight:600;color:{THEME['text']};")
        gml.addWidget(gmt)
        gmd = QLabel("切换 API 后端与模型")
        gmd.setStyleSheet(f"font-size:11px;color:{THEME['faint']};")
        gml.addWidget(gmd)

        self.model_combo = QComboBox()
        self.model_combo.setMinimumWidth(260)
        gml.addWidget(self.model_combo)
        layout.addWidget(group_model)

        # ---- API Key 组 ----
        group_apikey = QFrame()
        group_apikey.setStyleSheet(
            f"QFrame{{background:{THEME['elev']};border:none;border-radius:10px;padding:12px;}}")
        gakl = QVBoxLayout(group_apikey)
        gakl.setContentsMargins(12, 12, 12, 12)
        gakl.setSpacing(6)

        gakt = QLabel("API Key")
        gakt.setStyleSheet(f"font-size:12px;font-weight:600;color:{THEME['text']};")
        gakl.addWidget(gakt)
        gakd = QLabel("用于认证的密钥，回车保存")
        gakd.setStyleSheet(f"font-size:11px;color:{THEME['faint']};")
        gakl.addWidget(gakd)

        api_row = QHBoxLayout()
        self.api_key_edit = QLineEdit()
        self.api_key_edit.setEchoMode(QLineEdit.Password)
        self.api_key_edit.setPlaceholderText("输入 API Key 后回车保存")
        self.api_key_edit.setText(self.cfg.get("api_key", ""))
        self.api_key_edit.returnPressed.connect(self._save_api_key)
        api_row.addWidget(self.api_key_edit, 1)

        self.api_key_toggle = QPushButton("◉")
        self.api_key_toggle.setFixedSize(30, 30)
        self.api_key_toggle.setToolTip("显示/隐藏")
        self.api_key_toggle.setStyleSheet(
            f"QPushButton{{background:transparent;border:none;color:{THEME['dim']};"
            f"font-size:16px;border-radius:6px;}}"
            f"QPushButton:hover{{background:{THEME['sidebar_hover']};color:{THEME['text']};}}")
        self.api_key_toggle.setCheckable(True)
        self.api_key_toggle.clicked.connect(self._toggle_api_key_visible)
        api_row.addWidget(self.api_key_toggle)
        gakl.addLayout(api_row)
        layout.addWidget(group_apikey)

        # ---- 记忆加密（v4.75） ----
        group_enc = QFrame()
        group_enc.setStyleSheet(
            f"QFrame{{background:{THEME['elev']};border:none;border-radius:10px;padding:12px;}}")
        encl = QVBoxLayout(group_enc)
        encl.setContentsMargins(12, 12, 12, 12)
        encl.setSpacing(6)
        enct = QLabel("记忆加密口令")
        enct.setStyleSheet(f"font-size:12px;font-weight:600;color:{THEME['text']};")
        encl.addWidget(enct)
        encd = QLabel("启用后聊天记录与长期记忆以 Fernet 加密落盘（.enc），明文不残留；"
                      "口令经 PBKDF2 派生，salt 存于本地。留空=关闭加密。")
        encd.setWordWrap(True)
        encd.setStyleSheet(f"font-size:11px;color:{THEME['faint']};")
        encl.addWidget(encd)
        enc_row = QHBoxLayout()
        self.enc_pw_edit = QLineEdit()
        self.enc_pw_edit.setEchoMode(QLineEdit.Password)
        self.enc_pw_edit.setPlaceholderText("设置口令后回车保存（留空则关闭）")
        self.enc_pw_edit.setText(self.cfg.get("memory_encryption_passphrase", ""))
        self.enc_pw_edit.returnPressed.connect(self._save_enc_passphrase)
        enc_row.addWidget(self.enc_pw_edit, 1)
        enc_save = QPushButton("保存")
        enc_save.setFixedHeight(30)
        enc_save.setStyleSheet(
            f"QPushButton{{background:{THEME['accent']};border:none;border-radius:8px;"
            f"padding:0 14px;font-size:12px;color:#fff;}}"
            f"QPushButton:hover{{background:{THEME['accent_hover']};}}")
        enc_save.clicked.connect(self._save_enc_passphrase)
        enc_row.addWidget(enc_save)
        encl.addLayout(enc_row)
        layout.addWidget(group_enc)

        # ---- 功能开关组 ----
        group_toggles = QFrame()
        group_toggles.setStyleSheet(
            f"QFrame{{background:{THEME['elev']};border:none;border-radius:10px;padding:12px;}}")
        gtl = QVBoxLayout(group_toggles)
        gtl.setContentsMargins(12, 12, 12, 12)
        gtl.setSpacing(8)

        gtt = QLabel("功能开关")
        gtt.setStyleSheet(f"font-size:12px;font-weight:600;color:{THEME['text']};")
        gtl.addWidget(gtt)

        def _toggle_row(title, desc, var_name, checked, color=None):
            row_w = QWidget()
            rl = QVBoxLayout(row_w)
            rl.setContentsMargins(0, 0, 0, 0)
            rl.setSpacing(2)

            cb = QCheckBox(title)
            cb.setChecked(checked)
            if color:
                cb.setStyleSheet(f"color:{color};font-size:12px;")
            else:
                cb.setStyleSheet(f"font-size:12px;color:{THEME['dim']};")
            rl.addWidget(cb)

            dl = QLabel(desc)
            dl.setStyleSheet(f"font-size:10px;color:{THEME['faint']};padding-left:20px;")
            rl.addWidget(dl)

            if var_name == "search":
                cb.stateChanged.connect(self._on_search_toggle)
                self.search_toggle = cb
            elif var_name == "agent":
                cb.stateChanged.connect(self._on_agent_toggle)
                self.agent_toggle = cb
            elif var_name == "skip_confirm":
                cb.stateChanged.connect(self._on_skip_confirm_toggle)
                self.skip_confirm_toggle = cb
            return row_w

        gtl.addWidget(_toggle_row("联网搜索", "搜索失败自动降级为纯模型回答",
                                   "search", self.cfg.get("search_enabled", True)))
        gtl.addWidget(_toggle_row("Agent 模式", "启用工具调用与多步推理",
                                   "agent", self.cfg.get("agent_mode", False),
                                   THEME["ok"]))

        # ---- 执行模式（v4.50，借鉴 openworker 的权限引擎；替换旧「手动级操作免确认」开关）----
        _mode_label = QLabel("执行模式")
        _mode_label.setStyleSheet(f"font-size:13px;color:{THEME['text']};font-weight:600;padding-top:6px;")
        gtl.addWidget(_mode_label)
        _mode_cb = QComboBox()
        _mode_cb.setStyleSheet(f"font-size:12px;color:{THEME['text']};padding:4px;")
        for _m, _t in MODES.items():
            _mode_cb.addItem(_t, _m)
        _idx = _mode_cb.findData(self.permission_engine.mode)
        if _idx >= 0:
            _mode_cb.setCurrentIndex(_idx)
        _mode_cb.currentIndexChanged.connect(self._on_mode_change)
        gtl.addWidget(_mode_cb)
        _mode_hint = QLabel("交互：危险操作逐个问 ｜ 规划：只做只读 ｜ 自动：全直接执行 ｜ 仅讨论：不执行 ｜ 自定义：仅白名单免确认")
        _mode_hint.setWordWrap(True)
        _mode_hint.setStyleSheet(f"font-size:10px;color:{THEME['faint']};padding-left:4px;padding-bottom:4px;")
        gtl.addWidget(_mode_hint)
        _trust_btn = QPushButton("✅ 本次会话全部信任（危险操作不再逐个问）")
        _trust_btn.setFixedHeight(34)
        _trust_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['elev']};border:none;border-radius:8px;"
            f"padding:6px;font-size:12px;color:{THEME['text']};text-align:left;}}"
            f"QPushButton:hover{{background:{THEME['border_highlight']};}}")
        _trust_btn.clicked.connect(self._on_trust_session)
        gtl.addWidget(_trust_btn)
        layout.addWidget(group_toggles)

        # ---- 自动备份（v4.76，OS 级任务计划程序） ----
        group_backup = QFrame()
        group_backup.setStyleSheet(
            f"QFrame{{background:{THEME['elev']};border:none;border-radius:10px;padding:12px;}}")
        bkl = QVBoxLayout(group_backup)
        bkl.setContentsMargins(12, 12, 12, 12)
        bkl.setSpacing(6)
        bkt = QLabel("自动备份（系统级）")
        bkt.setStyleSheet(f"font-size:12px;font-weight:600;color:{THEME['text']};")
        bkl.addWidget(bkt)
        bkd = QLabel("定时把 ~/Documents/小臭玩AI（记忆/配置/反馈）备份到 backups/。"
                     "由 Windows 任务计划程序执行，关程序也能跑。")
        bkd.setWordWrap(True)
        bkd.setStyleSheet(f"font-size:11px;color:{THEME['faint']};")
        bkl.addWidget(bkd)
        bk_row = QHBoxLayout()
        self.ab_freq_combo = QComboBox()
        self.ab_freq_combo.addItem("关闭", "")
        self.ab_freq_combo.addItem("每日", "daily")
        self.ab_freq_combo.addItem("每周", "weekly")
        _cur_freq = self.cfg.get("autobackup_freq", "")
        _fi = self.ab_freq_combo.findData(_cur_freq)
        if _fi >= 0:
            self.ab_freq_combo.setCurrentIndex(_fi)
        bk_row.addWidget(self.ab_freq_combo)
        self.ab_time_edit = QTimeEdit()
        self.ab_time_edit.setDisplayFormat("HH:mm")
        try:
            self.ab_time_edit.setTime(
                datetime.strptime(self.cfg.get("autobackup_time", "03:00"), "%H:%M").time())
        except Exception:
            self.ab_time_edit.setTime(datetime.strptime("03:00", "%H:%M").time())
        bk_row.addWidget(self.ab_time_edit)
        bkl.addLayout(bk_row)
        bk_btn_row = QHBoxLayout()
        ab_apply = QPushButton("应用计划")
        ab_apply.setFixedHeight(30)
        ab_apply.setStyleSheet(
            f"QPushButton{{background:{THEME['accent']};border:none;border-radius:8px;"
            f"padding:0 14px;font-size:12px;color:#fff;}}"
            f"QPushButton:hover{{background:{THEME['accent_hover']};}}")
        ab_apply.clicked.connect(self._save_autobackup_settings)
        bk_btn_row.addWidget(ab_apply)
        ab_now = QPushButton("立即备份")
        ab_now.setFixedHeight(30)
        ab_now.setStyleSheet(
            f"QPushButton{{background:{THEME['elev']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 14px;font-size:12px;color:{THEME['text']};}}"
            f"QPushButton:hover{{background:{THEME['border_highlight']};}}")
        ab_now.clicked.connect(self._run_backup_now)
        bk_btn_row.addWidget(ab_now)
        bk_btn_row.addStretch(1)
        bkl.addLayout(bk_btn_row)
        layout.addWidget(group_backup)

        # ---- 版本 / 更新（v4.76） ----
        group_ver = QFrame()
        group_ver.setStyleSheet(
            f"QFrame{{background:{THEME['elev']};border:none;border-radius:10px;padding:12px;}}")
        vl = QVBoxLayout(group_ver)
        vl.setContentsMargins(12, 12, 12, 12)
        vl.setSpacing(6)
        vt = QLabel("版本与更新")
        vt.setStyleSheet(f"font-size:12px;font-weight:600;color:{THEME['text']};")
        vl.addWidget(vt)
        from config import APP_VERSION, APP_BUILD_DATE, UPDATE_CHECK_URL
        vinfo = QLabel(f"当前版本：{APP_VERSION}　构建日期：{APP_BUILD_DATE}")
        vinfo.setStyleSheet(f"font-size:11px;color:{THEME['faint']};")
        vl.addWidget(vinfo)
        vrow = QHBoxLayout()
        chk_btn = QPushButton("检查更新")
        chk_btn.setFixedHeight(30)
        chk_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['elev']};border:1px solid {THEME['border']};"
            f"border-radius:8px;padding:0 14px;font-size:12px;color:{THEME['text']};}}"
            f"QPushButton:hover{{background:{THEME['border_highlight']};}}")
        chk_btn.clicked.connect(self._check_update)
        vrow.addWidget(chk_btn)
        if not (self.cfg.get("update_check_url", "") or UPDATE_CHECK_URL):
            note = QLabel("（本地构建，无在线更新通道）")
            note.setStyleSheet(f"font-size:10px;color:{THEME['faint']};")
            vrow.addWidget(note)
        vrow.addStretch(1)
        vl.addLayout(vrow)
        layout.addWidget(group_ver)

        # ---- 性能基线（v4.78） ----
        try:
            import perf_baseline as _pb
            group_perf = QFrame()
            group_perf.setStyleSheet(
                f"QFrame{{background:{THEME['elev']};border:none;border-radius:10px;padding:12px;}}")
            pl = QVBoxLayout(group_perf)
            pl.setContentsMargins(12, 12, 12, 12)
            pl.setSpacing(6)
            pt = QLabel("性能基线")
            pt.setStyleSheet(f"font-size:12px;font-weight:600;color:{THEME['text']};")
            pl.addWidget(pt)
            # 最近一次冷启动耗时分解
            su = _pb.last_startup()
            if su:
                parts = ["最近冷启动 %.2fs" % (su.get("total") or 0)]
                for s in su.get("stages", [])[-6:]:
                    parts.append("%s %.2fs" % (s.get("stage"), s.get("dt") or 0))
                sl = QLabel("  ·  ".join(parts))
            else:
                sl = QLabel("尚未采集到启动耗时（启动后才会记录）")
            sl.setWordWrap(True)
            sl.setStyleSheet(f"font-size:10px;color:{THEME['faint']};")
            pl.addWidget(sl)
            # 按钮行：跑基线 / 设为基线
            prow = QHBoxLayout()
            btn_style = (
                f"QPushButton{{background:{THEME['elev']};border:1px solid {THEME['border']};"
                f"border-radius:8px;padding:0 14px;font-size:12px;color:{THEME['text']};}}"
                f"QPushButton:hover{{background:{THEME['border_highlight']};}}")
            run_btn = QPushButton("跑性能基线")
            run_btn.setFixedHeight(30)
            run_btn.setStyleSheet(btn_style)
            run_btn.clicked.connect(lambda: self._run_perf(False))
            set_btn = QPushButton("设为当前基线")
            set_btn.setFixedHeight(30)
            set_btn.setStyleSheet(btn_style)
            set_btn.clicked.connect(lambda: self._run_perf(True))
            prow.addWidget(run_btn)
            prow.addWidget(set_btn)
            prow.addStretch(1)
            pl.addLayout(prow)
            layout.addWidget(group_perf)
        except Exception:
            pass

        # ---- 导出 ----
        export_btn = QPushButton("导出当前对话")
        export_btn.setFixedHeight(36)
        export_btn.setStyleSheet(
            f"QPushButton{{background:{THEME['elev']};border:none;border-radius:8px;"
            f"padding:6px 8px;font-size:13px;color:{THEME['text']};text-align:left;}}"
            f"QPushButton:hover{{background:{THEME['border_highlight']};}}")
        export_btn.clicked.connect(self.export_session)
        layout.addWidget(export_btn)

        return popup

    def _show_popup(self, popup):
        if popup.isVisible():
            popup.hide()
            return
        for p in (self._skill_popup, self._settings_popup):
            if p is not popup and p.isVisible():
                p.hide()
        pos = self.sidebar.mapToGlobal(QPoint(self.sidebar.width() + 4, 12))
        popup.move(pos)
        popup.show()

    # ============ 技能 ============
    def _current_skill(self):
        s = self.store.sessions.get(self.store.active_sid)
        skill_id = s.skill if s else None
        if not skill_id:
            return None
        for sk in self._skills:
            if sk.get("id") == skill_id:
                return sk
        return None

    # 🚦 权限分级执行规则（来自《小臭玩AI-权限分级方案.md》第三节，注入系统提示开头附近）
    _PERMISSION_RULES = (
        "## 🚦 执行规则\n\n"
        "每个技能/工具都有对应的权限等级标签：\n"
        "- [自主] → 直接执行，不需要问我，完成后告诉我结果\n"
        "- [半自主] → 执行完成后把结果发我看看，等我说\"发\"/\"确认\"/\"可以\"再正式对外操作\n"
        "- [手动] → 必须先问我，我同意了你再执行\n\n"
        "如果有多个技能组合调用，按最高权限等级执行：\n"
        "  手动 > 半自主 > 自主\n"
        "  例如：联网搜索[自主] + 发邮件[手动] → 整个流程按手动处理\n"
    )

    def _prompt_section_permission(self):
        """权限分级清单（按等级分组列出所有工具），注入系统提示末尾，让 LLM 自知等级。"""
        groups = tools_mod.grouped_tools()
        lines = ["## 🚦 权限分级清单（工具 → 等级）", ""]
        for _label, _key in (("🟢 自主（直接执行）", "auto"),
                             ("🟡 半自主（执行完展示，等确认再对外）", "semi"),
                             ("🔴 手动（先问后做）", "manual")):
            _names = groups.get(_key, [])
            if _names:
                lines.append(f"{_label}：{', '.join(sorted(_names))}")
        return "\n".join(lines)

    # ---------- v4.87：L0/L1/L2 分层注入（省 token，应对 DeepSeek 涨价） ----------
    _TOPIC_REQUEST_KW = ("选题", "方向", "盘点", "爆款", "写什么", "列几个",
                         "赛道", "建议", "做什么内容", "内容方向", "给我一些")

    def _recent_user_query(self):
        """取最近一条用户消息文本，用于按需注入判断（L2 触发）。"""
        try:
            msgs = self.store.active().messages if self.store else []
            for m in reversed(msgs):
                if isinstance(m, dict) and m.get("role") == "user":
                    c = m.get("content")
                    if isinstance(c, str):
                        return c
                    if isinstance(c, list):  # 多模态（含图片/文本）
                        for p in c:
                            if isinstance(p, dict) and p.get("type") == "text":
                                return p.get("text", "")
        except Exception:
            pass
        return ""

    def _is_topic_request(self, q):
        """选题/盘点/方向类需求才注入完整选题模板（其余轮次省下这段 token）。"""
        if not q:
            return False
        return any(k in q for k in self._TOPIC_REQUEST_KW)

    def _harness_hit(self, title, q):
        """title 取长度>=4 的连续中文字窗（步长2）做包含匹配，判断是否命中用户消息。"""
        for i in range(0, len(title) - 3, 2):
            if title[i:i + 4] in q:
                return True
        return False

    def _build_system_prompt(self):
        """构建系统提示：L0 常驻(persona+权限+模式) + L1 概览(技能名/经验标题/记忆)
        + L2 按需(选题模板/命中经验全文)，大幅降低每轮 token 消耗。

        分层原则（v4.87，借鉴 OpenViking L0/L1/L2，不引入框架）：
        - L0 常驻：角色、硬约束、模式说明、权限、本会话目标——每轮必有。
        - L1 概览：技能仅列「名字+首句摘要」、经验库仅列标题、长期记忆语义召回(已有)——
          让模型有全局视野但不平铺全文。
        - L2 按需：选题模板仅在聊选题时注入；经验完整铁律仅在用户消息命中相关标题时展开。
        """
        base = self.cfg["system_prompt"]
        base += "\n\n" + self._PERMISSION_RULES
        # 本会话最初目标——注入到顶部，让模型无论多轮都记得最初要干什么
        try:
            sess = self.store.active() if hasattr(self, "store") else None
            if sess and getattr(sess, "goal", ""):
                base += "\n\n【本会话目标（首条用户消息）】\n" + sess.goal
        except Exception:
            pass
        # 按需触发判断用的用户消息（L2）
        user_query = self._recent_user_query()

        sk_section = self._prompt_section_skill()
        if sk_section:
            base += "\n\n" + sk_section
        # 模式感知：普通对话模式没有 tools，必须明确告知模型
        agent = getattr(self, "_agent_active", False)
        if agent:
            base += AGENT_SYS_APPEND
            base += self._build_tool_overview()
        else:
            base += (
                "\n\n【模式说明：当前为「普通对话模式」，你无法直接调用任何工具】\n"
                "· 若下方消息中包含【联网搜索结果】，请直接基于它作答；不要声称自己会去搜索或调用 web_search。\n"
                "· 若没有提供搜索结果，请基于自身知识简明回答，并说明『未检索到联网资料』。\n"
                "· 需要写文件/执行代码/生成图片/实际操作等任务时，请明确提示用户：「请开启 Agent 模式让我执行」，不要假装已完成。\n"
                "· 禁止输出「我现在去搜索/马上执行/下一步」之类的承诺性文字却不真正行动。\n"
                "· 覆盖说明：上方系统提示里的「命中路由表关键词必须调工具」「事实问题必须 web_search」两条约束，"
                "在普通对话模式下不适用，请以本【模式说明】为准。"
            )
        # L1 概览：技能名 + 首句摘要（compact 省 token；完整 prompt 在 use_skill 时注入）
        base += load_dynamic_skills(compact=True)
        rules = self._prompt_section_rules()
        if rules:
            base += "\n\n" + rules
        mem = self._prompt_section_memory()
        if mem:
            base += "\n\n" + mem
        base += "\n\n" + self._prompt_section_permission()
        # L1 概览：经验库标题列表；L2 按需：命中相关操作时展开完整铁律
        base += "\n\n" + self._prompt_section_harness(user_query)
        base += "\n\n" + self._prompt_section_facts()
        # L2 按需：仅选题/盘点类需求才注入完整选题模板（其余轮次省下这段 token）
        if self._is_topic_request(user_query):
            try:
                from config import TOPIC_IDEA_TEMPLATE
                base += "\n\n" + TOPIC_IDEA_TEMPLATE
            except Exception:
                pass
        return base

    def _prompt_section_skill(self):
        """当前技能 prompt（普通/Agent 模式都注入）。"""
        sk = self._current_skill()
        if not sk:
            return ""
        return f'【当前技能：{sk.get("name", "")}】\n{sk.get("prompt", "")}'

    def _prompt_section_rules(self):
        """行为规范 agent_rules.md（根目录找不到回退 _internal/，文件缺失静默跳过）。"""
        _ar = os.path.join(APP_DIR, "agent_rules.md")
        if not os.path.exists(_ar):
            _ar2 = os.path.join(APP_DIR, "_internal", "agent_rules.md")
            if os.path.exists(_ar2):
                _ar = _ar2
        if not os.path.exists(_ar):
            return ""
        try:
            with open(_ar, encoding="utf-8") as _f:
                return "【行为规范 agent_rules.md】\n" + _f.read()[:3000]
        except Exception:
            return ""

    def _prompt_section_memory(self):
        """v4.73：跨对话长期记忆——语义召回（钉住核心画像 + 相关记忆 top-k）。

        用当前会话最近的用户消息作查询，召回相关长期记忆；核心画像无条件注入。
        替换旧版 load_recent 尾部平铺（无论聊什么都灌末尾 4KB，且早期重要记忆会被滚动淘汰冲掉）。
        """
        try:
            from memory_store import recall_memory
            msgs = []
            try:
                msgs = self.store.active().messages if self.store else []
            except Exception:
                msgs = []
            # 取最近 3 条用户消息作为召回查询
            qparts = []
            for m in msgs[-8:]:
                if isinstance(m, dict) and m.get("role") == "user":
                    c = m.get("content")
                    if isinstance(c, str):
                        qparts.append(c)
            query = "\n".join(qparts[-3:])
            mem = recall_memory(query, limit=8, max_chars=4000)
            if mem:
                return "【用户长期记忆】\n" + mem
        except Exception:
            pass
        return ""

    def _prompt_section_facts(self):
        """v4.60：系统事实锚点——注入运行时真实数据，防止自检时模型瞎猜。"""
        import os
        lines = ["【系统事实（自检时以此为准，不要猜测）】"]
        data_dir = os.path.expanduser("~/Documents/小臭玩AI")
        lines.append(f"- 数据目录：{data_dir}")
        lines.append("- 记忆库：memory.md + memory.db (SQLite FTS5)")
        ov = self.cfg.get("obsidian_vault_path", "")
        lines.append(f"- Obsidian Vault：{'未配置' if not ov else ov}")
        rd = self.cfg.get("rag_data_dir", "") or os.path.join(config.APP_DIR, "rag_data")
        lines.append(f"- RAG 目录：{rd}")
        try:
            import sqlite3
            for db, label in [("xiaochou.db", "应用库"), ("agent_log.db", "日志库")]:
                dp = os.path.join(data_dir, db)
                if os.path.exists(dp):
                    conn = sqlite3.connect(dp)
                    tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
                    conn.close()
                    tl = ", ".join(t[0] for t in tables)
                    lines.append(f"- {label} ({db})：{len(tables)} 表（{tl}）")
                else:
                    lines.append(f"- {label} ({db})：不存在")
        except Exception:
            lines.append("- 数据库状态：读取失败")
        try:
            dyn = config.load_dynamic_skills()
            dyn_count = dyn.count("name: ") if dyn else 0
            lines.append(f"- 技能：14 工具栏(SKILL.md) + ~{dyn_count} 动态加载")
        except Exception:
            lines.append("- 技能：14 工具栏(SKILL.md)")
        return "\n".join(lines)

    def _prompt_section_harness(self, user_query=""):
        """v4.80/87：操作经验库 harness_notes 补充节（不改动基础系统提示）。

        L1 概览：标题列表常驻（让模型知道有哪些经验）；
        L2 按需：user_query 命中某 note 标题时，展开该 note 完整铁律（不命中只列标题，省 token）。
        基础系统提示锁死，操作经验作为可检索、可版本化、可回滚的 supplemental 状态层。
        文件缺失/损坏时静默跳过，绝不拖垮主程序。
        """
        try:
            import harness
            base = harness.harness_section_text(self.cfg, compact=True)
            if user_query:
                data = harness.load_harness_notes(self.cfg)
                hits = []
                for e in data.get("entries", []):
                    title = e.get("title", "")
                    body = e.get("body", "")
                    if title and self._harness_hit(title, user_query):
                        hits.append(f"- {title}：{body}")
                if hits:
                    base += "\n\n【命中操作经验详情（按需展开）】\n" + "\n".join(hits)
            return base
        except Exception:
            return ""

    def _build_tool_overview(self):
        """构建 Agent 模式下的工具能力概览文本。

        从 config.get_all_tools(cfg) 提取每个工具的 name + description 首行，
        格式化为列表，注入到系统提示中，让 LLM 对自己有什么能力有完整认知。
        """
        try:
            tools = config.get_all_tools(self.cfg)
        except Exception as e:
            log.warning("工具概览生成失败: %s", e)
            return "\n\n【工具能力概览】生成失败，请重试。\n"
        if not tools:
            return "\n\n【工具能力概览】暂无可用工具。\n"
        lines = ["\n\n【工具能力概览】以下工具可在 Agent 模式下调用："]
        for t in tools:
            fn = t.get("function", {}) if isinstance(t, dict) else {}
            name = fn.get("name", "")
            desc = (fn.get("description", "") or "").strip().split("\n")[0][:120]
            if name:
                lines.append(f"- {name}: {desc}")
        return "\n".join(lines) + "\n"

    @staticmethod
    def _skill_btn_style(active):
        if active:
            return (f"QPushButton{{background:{THEME['sidebar_active']};border:none;"
                    f"border-radius:8px;padding:6px 10px;font-size:12px;"
                    f"color:{THEME['user_text']};font-weight:600;text-align:left;}}")
        return (f"QPushButton{{background:rgba(10,10,12,0.05);border:none;"
                f"border-radius:8px;padding:6px 10px;font-size:12px;"
                f"color:{THEME['dim']};text-align:left;}}"
                f"QPushButton:hover{{background:rgba(10,10,12,0.09);color:{THEME['text']};}}")

    def _on_skill_pick(self, skill_id):
        s = self.store.active()
        s.skill = None if s.skill == skill_id else skill_id
        self.store.save()
        self._refresh_skill_buttons()
        self._update_skill_bar()
        sk = self._current_skill()
        # v4.110 旁路埋点：只在「选中」时记一次，取消选中不算一次使用
        if sk:
            try:
                route_log.log_skill(sk.get("name") or skill_id, "manual")
            except Exception:
                pass
        self.status_label.setText(f"已切换技能：{sk['name']}" if sk else "已切换为通用模式")

    def _clear_skill(self):
        s = self.store.active()
        s.skill = None
        self.store.save()
        self._refresh_skill_buttons()
        self._update_skill_bar()
        self.status_label.setText("已切换为通用模式")

    def _refresh_skill_buttons(self):
        cur = self.store.active().skill
        for sid, btn in getattr(self, "skill_buttons", {}).items():
            btn.setStyleSheet(self._skill_btn_style(sid == cur))

    def _update_skill_bar(self):
        bar = getattr(self, "skill_bar", None)
        if bar is None:
            return
        sk = self._current_skill()
        if sk:
            self.skill_name_label.setText(f"当前技能：{sk.get('emoji', '')} {sk.get('name', '')}")
            bar.setVisible(True)
        else:
            bar.setVisible(False)

    def _populate_skill_lib(self, layout):
        cats = {}
        order = []
        for sk in self._skills:
            cat = sk.get("category", "其他")
            if cat not in cats:
                cats[cat] = []
                order.append(cat)
            cats[cat].append(sk)
        for cat in order:
            cat_lbl = QLabel(cat)
            cat_lbl.setStyleSheet(
                f"font-size:11px;font-weight:600;color:{THEME['faint']};"
                f"padding-left:4px;margin-top:2px;background:transparent;")
            layout.addWidget(cat_lbl)
            for sk in cats[cat]:
                btn = QPushButton(f'{sk.get("emoji", "")}  {sk.get("name", "")}')
                btn.setFixedHeight(30)
                btn.setStyleSheet(self._skill_btn_style(False))
                btn.setToolTip(sk.get("desc", ""))
                btn.clicked.connect(lambda _=False, sid=sk["id"]: self._on_skill_pick(sid))
                self.skill_buttons[sk["id"]] = btn
                layout.addWidget(btn)
        layout.addStretch(1)

    # ============ v4.84 技能审核队列（软自进化·人工闸门）============
    def _refresh_skill_review_label(self):
        """刷新托盘「技能审核」菜单项的待审数量徽标。"""
        try:
            import skill_review
            n = skill_review.count_pending(self.cfg)
            tray = getattr(self, "tray_app", None)
            action = getattr(tray, "_skill_review_action", None) if tray else None
            if action is not None:
                action.setText(f"📝 技能审核 ({n})" if n else "📝 技能审核")
        except Exception:
            pass

    def _scan_skill_review(self):
        """启动扫描：有待审核技能则弹系统托盘提示，引导用户去审核。"""
        try:
            import skill_review
            n = skill_review.count_pending(self.cfg)
        except Exception:
            return
        self._refresh_skill_review_label()
        if n > 0:
            try:
                self.tray_app.tray.showMessage(
                    "技能审核",
                    f"有 {n} 个模型自动创建的技能待你审核，右键托盘→「技能审核」查看。",
                    QSystemTrayIcon.Information, 8000)
            except Exception:
                pass

    def _reload_skill_lib(self):
        """审核通过/拒绝后热重载技能库（弹层 + 侧栏按钮）。"""
        try:
            layout = self.skill_inner_layout
            while layout.count():
                item = layout.takeAt(0)
                w = item.widget()
                if w is not None:
                    w.setParent(None)
                    w.deleteLater()
            self._skills = load_skills()
            self._populate_skill_lib(layout)
            self._refresh_skill_buttons()
        except Exception as e:
            log.warning("重载技能库失败: %s", e)

    def _open_skill_review_dialog(self):
        """弹出技能审核对话框：列出待审核技能，逐条「通过 / 拒绝」。"""
        try:
            import skill_review
        except Exception as e:
            QMessageBox.warning(self, "技能审核", f"模块加载失败：{e}")
            return

        pending = skill_review.list_pending(self.cfg)
        dlg = QDialog(self)
        dlg.setWindowTitle("技能审核（模型自动创建·待你通过）")
        dlg.setMinimumWidth(540)
        dlg.setMinimumHeight(380)
        root = QVBoxLayout(dlg)
        if not pending:
            root.addWidget(QLabel("暂无待审核技能。模型自动创建的技能会出现在这里，"
                                  "通过后才会正式生效。"))
            btn_close = QPushButton("关闭")
            btn_close.setStyleSheet(self._secondary_btn_style())
            btn_close.clicked.connect(dlg.accept)
            root.addWidget(btn_close, alignment=Qt.AlignRight)
            dlg.exec()
            self._refresh_skill_review_label()
            return

        hint = QLabel(f"共 {len(pending)} 个待审核技能。通过后移入正式技能目录并热重载；"
                      f"拒绝则删除。通过后即可在对话中调用。")
        hint.setStyleSheet(f"font-size:12px;color:{THEME['faint']};")
        root.addWidget(hint)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        inner = QWidget()
        vlay = QVBoxLayout(inner)
        vlay.setSpacing(10)
        for sk in pending:
            box = QFrame()
            box.setStyleSheet(
                f"QFrame{{background:{THEME['elev']};border:1px solid {THEME['border']};"
                f"border-radius:10px;padding:10px;}}")
            bl = QVBoxLayout(box)
            bl.setContentsMargins(10, 8, 10, 8)
            title = QLabel(f"{sk.get('emoji','⚡')} {sk.get('display_name', sk.get('name',''))}   "
                           f"[{sk.get('category','自动生成')}]   {sk.get('created','')}")
            title.setStyleSheet(f"font-size:13px;font-weight:600;color:{THEME['text']};")
            bl.addWidget(title)
            desc = QLabel(sk.get("description", "") or "（无描述）")
            desc.setWordWrap(True)
            desc.setStyleSheet(f"font-size:12px;color:{THEME['faint']};")
            bl.addWidget(desc)
            row = QHBoxLayout()
            row.addStretch(1)
            reject_btn = QPushButton("拒绝 ✗")
            reject_btn.setFixedHeight(30)
            reject_btn.setStyleSheet(self._secondary_btn_style())
            approve_btn = QPushButton("通过 ✓")
            approve_btn.setFixedHeight(30)
            approve_btn.setStyleSheet(self._primary_btn_style())
            name = sk.get("name", "")
            reject_btn.clicked.connect(
                lambda _=False, nm=name: self._on_skill_review_decision(nm, False, dlg))
            approve_btn.clicked.connect(
                lambda _=False, nm=name: self._on_skill_review_decision(nm, True, dlg))
            row.addWidget(reject_btn)
            row.addWidget(approve_btn)
            bl.addLayout(row)
            vlay.addWidget(box)
        vlay.addStretch(1)
        scroll.setWidget(inner)
        root.addWidget(scroll, 1)
        close_btn = QPushButton("关闭")
        close_btn.setStyleSheet(self._secondary_btn_style())
        close_btn.clicked.connect(dlg.accept)
        root.addWidget(close_btn, alignment=Qt.AlignRight)
        dlg.exec()
        self._refresh_skill_review_label()

    def _on_skill_review_decision(self, name, approve, dlg):
        """处理单条审核决定：通过→移入正式目录+重载；拒绝→删除。"""
        try:
            import skill_review
            if approve:
                msg = skill_review.approve_skill(self.cfg, name)
                self._reload_skill_lib()
            else:
                msg = skill_review.reject_skill(self.cfg, name)
            self.status_label.setText(msg)
            try:
                self.tray_app.tray.showMessage("技能审核", msg, QSystemTrayIcon.Information, 4000)
            except Exception:
                pass
        except Exception as e:
            msg = f"审核操作失败：{e}"
            self.status_label.setText(msg)
        dlg.accept()
        # 刷新列表：还有剩余则重开对话框
        try:
            import skill_review
            if skill_review.count_pending(self.cfg) > 0:
                self._open_skill_review_dialog()
        except Exception:
            pass

    # ============ 模型切换 / 导出 ============
    def _on_model_change(self, name):
        prof = self.cfg.get("model_profiles", {}).get(name)
        if not prof:
            return
        self.cfg["base_url"] = prof["base_url"]
        self.cfg["model"] = prof["model"]
        if prof.get("api_key"):
            self.cfg["api_key"] = prof["api_key"]
            self.api_key_edit.setText(self.cfg["api_key"])
        self._save_cfg()
        self.status_label.setText(f"已切换模型：{name}（{prof['model']}）")

    def _save_api_key(self):
        key = self.api_key_edit.text().strip()
        self.cfg["api_key"] = key
        self._save_cfg()
        if key:
            self.status_label.setText("API Key 已保存")
        else:
            self.status_label.setText("API Key 已清空，发送前请重新填写")

    def _toggle_api_key_visible(self, checked):
        self.api_key_edit.setEchoMode(QLineEdit.Normal if checked else QLineEdit.Password)
        self.api_key_toggle.setText("◎" if checked else "◉")

    def export_session(self):
        """v4.76：富格式导出当前会话——Markdown / HTML / PDF / Word(.docx)。"""
        session = self.store.active()
        if not session.messages:
            self.status_label.setText("当前会话没有内容可导出")
            return
        default_name = (session.title or "会话").replace("/", "_")
        filters = ("Markdown 文件 (*.md);;HTML 网页 (*.html);;"
                   "PDF 文档 (*.pdf);;Word 文档 (*.docx);;纯文本 (*.txt)")
        path, sel = QFileDialog.getSaveFileName(self, "导出对话", default_name, filters)
        if not path:
            return
        ext = os.path.splitext(path)[1].lower()
        if not ext:
            mapping = {"Markdown": ".md", "HTML": ".html", "PDF": ".pdf",
                       "Word": ".docx", "纯文本": ".txt"}
            ext = next((v for k, v in mapping.items() if k in (sel or "")), ".md")
            path += ext
        try:
            if ext in (".md", ".txt"):
                self._export_md(session, path)
            elif ext == ".html":
                self._export_html(session, path)
            elif ext == ".pdf":
                self._export_pdf(session, path)
            elif ext == ".docx":
                self._export_docx(session, path)
            self.status_label.setText(f"已导出：{os.path.basename(path)}（{ext[1:].upper()}）")
        except Exception as e:
            log.error("导出失败: %s", e)
            self.status_label.setText(f"导出失败：{e}")

    # ============ v4.76：富格式导出 ============
    def _export_css(self):
        """导出 HTML/PDF 共用的内联样式。"""
        return (
            "body{font-family:'Microsoft YaHei','PingFang SC',sans-serif;"
            "max-width:820px;margin:24px auto;padding:0 16px;color:#1f2328;line-height:1.7;}"
            "h1{font-size:22px;margin-bottom:4px;} .meta{color:#888;font-size:13px;margin:2px 0;}"
            "hr{border:none;border-top:1px solid #e5e7eb;margin:16px 0;}"
            ".msg{margin:14px 0;} .role{font-size:12px;font-weight:600;color:#888;margin-bottom:4px;}"
            ".bubble{padding:10px 14px;border-radius:10px;white-space:pre-wrap;word-break:break-word;}"
            ".user .bubble{background:#e8f0fe;} .asst .bubble{background:#f1f3f5;}"
            ".tool{border:1px solid #e5e7eb;border-radius:8px;padding:8px 12px;margin:10px 0;"
            "background:#fafafa;font-size:13px;}"
            ".tool .role{color:#2563eb;} .result{color:#666;margin-top:4px;}"
            "pre{background:#0f172a;color:#e2e8f0;padding:10px;border-radius:6px;"
            "overflow:auto;font-size:12px;white-space:pre-wrap;}"
            "code{background:#f1f1f1;padding:1px 4px;border-radius:4px;font-size:12px;}"
            "b{font-weight:600;}"
        )

    def _md_to_html(self, text):
        """极简 Markdown→HTML：转义 + 代码块 + 加粗 + 换行；用于 HTML/PDF 导出。"""
        import re as _re
        text = html_mod.escape(text or "")
        out = []
        segs = _re.split(r"```", text)
        for i, seg in enumerate(segs):
            if i % 2 == 1:
                out.append(f"<pre>{seg}</pre>")
            else:
                seg = _re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", seg)
                seg = seg.replace("\n", "<br>")
                out.append(seg)
        return "".join(out)

    def _session_html_body(self, session):
        """生成会话正文的 HTML 片段（不含 <html>/<style> 外壳）。"""
        sk = self._current_skill()
        parts = [f"<h1>{html_mod.escape(session.title or '会话')}</h1>",
                 f"<p class='meta'>导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}</p>"]
        if sk:
            parts.append(f"<p class='meta'>当前技能：{sk.get('emoji', '')} "
                         f"{html_mod.escape(str(sk.get('name', '')))}</p>")
        if getattr(session, "deliverables", None):
            parts.append("<p class='meta'><b>交付物：</b></p><ul>")
            for d in session.deliverables:
                parts.append(f"<li>{html_mod.escape(str(d.get('name', '')))} → "
                             f"<code>{html_mod.escape(str(d.get('rel', '')))}</code></li>")
            parts.append("</ul>")
        parts.append("<hr>")
        for m in session.messages:
            if m.get("role") == "tool_log":
                parts.append(
                    f"<div class='tool'><div class='role'>工具调用："
                    f"{html_mod.escape(str(m.get('name', '')))}</div>"
                    f"<pre>{html_mod.escape(str(m.get('args', '')))}</pre>"
                    f"<div class='result'>→ {html_mod.escape(str(m.get('result', '')))}</div></div>")
            else:
                who = "你" if m["role"] == "user" else "助手"
                cls = "user" if m["role"] == "user" else "asst"
                content = m.get("content", "")
                if isinstance(content, list):
                    content = self._extract_text(m)
                parts.append(
                    f"<div class='msg {cls}'><div class='role'>{who}</div>"
                    f"<div class='bubble'>{self._md_to_html(content)}</div></div>")
        return "\n".join(parts)

    def _export_md(self, session, path):
        lines = [f"# {session.title}", "",
                 f"> 导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", ""]
        sk = self._current_skill()
        if sk:
            lines.append(f"> 当前技能：{sk.get('emoji', '')} {sk.get('name', '')}\n")
        if getattr(session, "deliverables", None):
            lines.append("> **交付物：**")
            for d in session.deliverables:
                lines.append(f"> - {d.get('name', '')} → `{d.get('rel', '')}`")
            lines.append("")
        lines.append("---")
        for m in session.messages:
            if m.get("role") == "tool_log":
                lines.append(
                    f"**工具调用：{m.get('name', '')}**\n\n"
                    f"```\n{m.get('args', '')}\n```\n\n"
                    f"→ {m.get('result', '')}\n")
            else:
                who = "你" if m["role"] == "user" else "助手"
                content = m.get("content", "")
                if isinstance(content, list):
                    content = self._extract_text(m)
                lines.append(f"**{who}**：\n\n{content}\n")
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _export_html(self, session, path):
        body = self._session_html_body(session)
        html = (f"<!DOCTYPE html><html lang='zh-CN'><head>"
                f"<meta charset='utf-8'><title>"
                f"{html_mod.escape(session.title or '会话')}</title>"
                f"<style>{self._export_css()}</style></head>"
                f"<body>{body}</body></html>")
        with open(path, "w", encoding="utf-8") as f:
            f.write(html)

    def _export_pdf(self, session, path):
        from PySide6.QtPrintSupport import QPrinter
        from PySide6.QtGui import QTextDocument, QPageSize
        body = self._session_html_body(session)
        styled = (f"<html><head><meta charset='utf-8'>"
                  f"<style>{self._export_css()}</style></head><body>{body}</body></html>")
        printer = QPrinter()
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName(path)
        printer.setPageSize(QPageSize.A4)
        doc = QTextDocument()
        doc.setHtml(styled)
        doc.print_(printer)

    def _export_docx(self, session, path):
        from docx import Document
        doc = Document()
        doc.add_heading(session.title or "会话", level=1)
        doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        sk = self._current_skill()
        if sk:
            doc.add_paragraph(f"当前技能：{sk.get('emoji', '')} {sk.get('name', '')}")
        for m in session.messages:
            if m.get("role") == "tool_log":
                p = doc.add_paragraph()
                p.add_run(f"工具调用：{m.get('name', '')}").bold = True
                doc.add_paragraph(str(m.get('args', ''))).style = "No Spacing"
                doc.add_paragraph(f"→ {m.get('result', '')}")
            else:
                who = "你" if m["role"] == "user" else "助手"
                p = doc.add_paragraph()
                p.add_run(f"{who}：").bold = True
                content = m.get("content", "")
                if isinstance(content, list):
                    content = self._extract_text(m)
                doc.add_paragraph(str(content))
        doc.save(path)

    # ============ 发送流程 ============
    def send_user_prompt(self, text, force_agent=False):
        """工作流模板调用：把一段文本当作用户输入发到当前会话。

        force_agent=True 时临时启用 Agent 模式（多工具执行），发完即恢复，
        不影响全局 agent_mode 设置。复用 send() 的全部路由与渲染逻辑。
        """
        text = (text or "").strip()
        if not text:
            return
        if not self.cfg.get("api_key"):
            self.status_label.setText("还没填 API Key，请在设置中输入后回车")
            return
        if self._busy:
            self.status_label.setText("上一条还在处理，稍等…")
            return
        if force_agent:
            prev = self.agent_mode
            self.agent_mode = True
            try:
                self.input_box.setPlainText(text)
                self.send()
            finally:
                self.agent_mode = prev
        else:
            self.input_box.setPlainText(text)
            self.send()

    def _on_input_changed(self):
        if not self._busy:
            has_text = bool(self.input_box.toPlainText().strip())
            has_images = bool(self.input_box.get_images())
            self.send_btn.setEnabled(has_text or has_images)

    def send(self):
        # v4.102 fix7：_busy 检查提到最前——上一条任务未结束时，直接提示并保留输入，
        # 不做图片压缩/文件解析等无效工作（之前压缩日志写了消息却被丢弃，用户以为卡死）。
        if self._busy:
            self.status_label.setText("上一条还在处理，请稍候再发（本条已保留在输入框）")
            return
        text = self.input_box.toPlainText().strip()
        # Strip image placeholder markers from text
        text = re.sub(r"\[\u56fe\u7247\u5df2\u7c98\u8d34 \d+\]", "", text).strip()
        # v4.102 hotfix：把 [文件: incoming/xxx.png] / [file: ...] 这类文件标记也解析成
        # 真实的 image_url content parts，避免附件只以纯文本路径发给模型、导致「看图」失效。
        clean_text, file_image_parts = _extract_file_image_parts(text, APP_DIR)
        images = self.input_box.get_images()
        has_images = bool(images or file_image_parts)
        if not clean_text and not has_images:
            return
        # v4.107：导演指令拦截——主对话框输入「改第3镜关键帧 / 主角换发型 / 合成成片 /
        # 导演台进度」等，不进主控 Agent、不写主会话（与主对话模块零交集）。
        # 改为在聊天区渲染一条引导，提示去导演台底部「导演对话」条下指令。
        if clean_text and self._is_director_command(clean_text):
            self.input_box.clear()
            self._render_director_redirect(clean_text)
            return
        if not self.cfg["api_key"]:
            self.status_label.setText("还没填 API Key，请在设置中输入后回车")
            return

        # v4.101：发出新消息即隐藏「继续上次任务」入口（若本消息开启新 Agent 任务，
        # _agent_run 会清理旧暂停检查点；普通消息则直接让入口消失）
        try:
            self.resume_agent_btn.setVisible(False)
        except Exception:
            pass

        session = self.store.active()
        edit_idx = getattr(self, "_edit_target_idx", None)
        if edit_idx is not None:
            # v4.75 改写问题模式：替换目标 user 消息内容并丢弃其后所有消息（重开分支）
            self._edit_target_idx = None
            if has_images:
                content = [{"type": "text", "text": clean_text}]
                total_orig = 0; total_final = 0
                for img_path in images:
                    comp = _compress_image_for_api(img_path)
                    if comp:
                        url, mime, orig_kb, final_kb, _ = comp
                        content.append({"type": "image_url", "image_url": {"url": url}})
                        total_orig += orig_kb; total_final += final_kb
                    else:
                        try:
                            import base64
                            with open(img_path, "rb") as f:
                                b64 = base64.b64encode(f.read()).decode()
                            content.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}})
                        except Exception as e:
                            content.append({"type": "text", "text": f"\n[图片加载失败: {e}]"})
                content.extend(file_image_parts)
                if total_orig > total_final > 0:
                    _vision_debug(f"send(edit): 图片压缩 {total_orig}KB→{total_final}KB ({len(images)}张)")
                session.messages[edit_idx]["content"] = content
                self.input_box.clear_images()
            else:
                session.messages[edit_idx]["content"] = clean_text
            del session.messages[edit_idx + 1:]
            self._track_context("user", session.messages[edit_idx]["content"])
            self.store.save()
            self.input_box.clear()
            self._rendered_msg_count = 0
            self._render_messages()
        elif has_images:
            content = [{"type": "text", "text": clean_text}]
            total_orig = 0; total_final = 0
            for img_path in images:
                # v4.102 fix4：先走 _compress_image_for_api（缩放+JPEG），
                # 避免真实截图多张发图时 payload 超 DeepSeek 视觉模型限制。
                comp = _compress_image_for_api(img_path)
                if comp:
                    url, mime, orig_kb, final_kb, _ = comp
                    content.append({"type": "image_url", "image_url": {"url": url}})
                    total_orig += orig_kb; total_final += final_kb
                else:
                    # 压缩失败回退直接读（保证不阻断发送）
                    try:
                        import base64
                        with open(img_path, "rb") as f:
                            b64 = base64.b64encode(f.read()).decode()
                        content.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}})
                    except Exception as e:
                        content.append({"type": "text", "text": f"\n[图片加载失败: {e}]"})
            content.extend(file_image_parts)
            if total_orig > total_final > 0:
                _vision_debug(f"send: 图片压缩 {total_orig}KB→{total_final}KB ({len(images)}张)")
            session.messages.append({"role": "user", "content": content})
            self._track_context("user", content)
            self.input_box.clear_images()
        else:
            session.messages.append({"role": "user", "content": clean_text})
            self._track_context("user", clean_text)

        if session.title in ("", "???", "新会话"):
            # 自动用首条消息命名（方便按项目名查找）
            seed = clean_text or ""
            if not seed and session.messages:
                last = session.messages[-1].get("content", "")
                if isinstance(last, list):
                    seed = "".join(p.get("text", "") for p in last if isinstance(p, dict))
                elif isinstance(last, str):
                    seed = last
            seed = seed.strip().replace("\n", " ")
            if seed:
                session.title = seed[:20] + ("…" if len(seed) > 20 else "")
                self._refresh_session_combo()

        # 记录本会话最初目标（首条 user 消息原话），用于长对话中防止模型「中途忘了要干什么」
        if not session.goal:
            seed = (clean_text or "").strip()
            if not seed and session.messages:
                last = session.messages[-1].get("content", "")
                if isinstance(last, list):
                    seed = "".join(p.get("text", "") for p in last if isinstance(p, dict))
                elif isinstance(last, str):
                    seed = last
                seed = seed.strip()
            session.goal = seed

        self.store.save()
        self.input_box.clear()
        self._render_messages()

        # 自动路由：默认是普通对话模式，但"搜索→写文件→数据分析"这类多步任务
        # 普通模式根本干不了（没工具、一次性回答）。检测到执行意图时自动走 Agent 真正执行。
        # v4.56 新增：纯"列方向/选题/盘点"型需求不联网不走 Agent——直接 LLM 裸出，避免
        # 「给我列几个方向」被误判为「去搜实时榜单」导致调一堆无意义搜索。
        # v4.102 hotfix：含图片附件/贴图的消息优先走普通对话 + 视觉模型，不进 Agent 空转调工具。
        use_agent = (self.agent_mode or self._message_needs_agent(clean_text)) and not has_images
        # v4.61：纯咨询 / 问意见类（如「给点自媒体平台方面的意见」）——即使开了 Agent 执行模式，
        # 也降级为普通对话直答，不进 Agent。否则 content-gap-analysis 等技能的「Stop and ask」
        # 会让模型反复追问、刷出十几个问问题气泡。
        advice_only = self._message_is_advice_only(clean_text)
        if use_agent and advice_only:
            use_agent = False
        topic_only = self._message_is_topic_only(clean_text)
        # v4.57：纯陈述 / 感慨（如「现在好多平台反应太严了」）即使 agent_mode 开着也走普通对话，
        # 不进 Agent 空转搜、不刷空气泡
        statement_only = self._message_is_statement_only(clean_text)
        if use_agent and not statement_only:
            self._busy = True
            self.send_btn.setEnabled(False)
            self._busy_timeout.start(120000)
            self.status_label.setText(
                "Agent 启动…" if self.agent_mode
                else "Agent 启动…（检测到需要执行操作，已自动切换）")
            try:
                self._agent_run()
            except Exception as e:
                log.error("Agent 启动失败: %s", e)
                self._reset_busy()
                self.status_label.setText(f"Agent 出错: {e}")
                self._on_input_changed()
        elif statement_only:
            # v4.57：纯陈述 / 感慨 → 普通对话（LLM 直接回应），不进 Agent、不联网搜
            self._busy = True
            self.send_btn.setEnabled(False)
            self._busy_timeout.start(120000)
            self.status_label.setText("（日常闲聊）")
            self._start_stream(clean_text, None)
        elif advice_only:
            # v4.61：纯咨询 / 问意见 → 普通对话（LLM 直接给建议），不进 Agent、不联网搜、不刷屏
            self._busy = True
            self.send_btn.setEnabled(False)
            self._busy_timeout.start(120000)
            self.status_label.setText("（给建议中…）")
            self._start_stream(clean_text, None)
        elif has_images:
            self._start_stream(clean_text, None)
        elif topic_only:
            # v4.56: 纯选题/盘点型需求 → 跳过联网搜，直接 LLM 裸出
            self._busy = True
            self.send_btn.setEnabled(False)
            self._busy_timeout.start(120000)
            self.status_label.setText("选题生成中…（用 AI 常识列方向）")
            self._start_stream(clean_text, None)
        elif self.cfg.get("search_enabled", True):
            self._busy = True
            self.send_btn.setEnabled(False)
            self._busy_timeout.start(120000)
            self.status_label.setText("搜索中…")
            self._do_search(clean_text)
        else:
            self._start_stream(clean_text, None)
    def _do_search(self, text):
        chain = search_mod.provider_chain(self.cfg.get("search_provider", "auto"))
        self._search_text = text
        self._search_acc = []
        self._search_step(text, chain, 0)

    def _search_step(self, text, chain, idx):
        if idx >= len(chain):
            ctx = search_mod.format_context(self._search_acc) if self._search_acc else None
            self._start_stream(text, ctx)
            return
        provider = chain[idx]
        url = search_mod.search_url(provider, text, self.cfg.get("search_top_k", 5))
        req = QNetworkRequest(QUrl(url))
        req.setHeader(QNetworkRequest.UserAgentHeader,
                      "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36")
        reply = self.search_manager.get(req)
        reply.finished.connect(lambda r=reply, t=text, c=chain, i=idx:
                               self._on_search_done(r, t, c, i))

    def _on_search_done(self, reply, text, chain, idx):
        provider = chain[idx]
        raw = bytes(reply.readAll()).decode("utf-8", "ignore")
        reply.deleteLater()
        results = search_mod.parse_search(raw, provider, self.cfg.get("search_top_k", 5))
        if results:
            self._search_acc.extend(results)
        if self._search_acc or idx == len(chain) - 1:
            ctx = search_mod.format_context(self._search_acc) if self._search_acc else None
            self._start_stream(text, ctx)
        else:
            self.status_label.setText(f"搜索中…（{chain[idx+1]}）")
            self._search_step(text, chain, idx + 1)

    # ============ Agent 模式 ============
    def _agent_run(self, resume=False, resume_task_id=None):
        # v4.101：开启新任务前，清理本会话残留的「已暂停」Agent 检查点（用户放弃旧任务、开启新任务）
        if not resume:
            self._cleanup_paused_agent_checkpoints()
        self._agent_active = True  # 本次实际走 Agent 管线，提示词按 Agent 模式构建
        session = self.store.active()
        self._streaming = False
        self._streaming_text = ""
        sys_msg = {"role": "system",
                   "content": self._build_system_prompt()}
        # v4.102 图像输入链路：agent 任务若带图（含 image_url）或命中工具意图，路由会
        # 升级到视觉模型（deepseek-v4-flash-vision-exp）。预判路由结果是否视觉，是则
        # 保留图像 content，让 AgentWorker 真正"看图"干活；非视觉仍归一化纯文本。
        _msgs = [m for m in session.messages if isinstance(m, dict)]
        _ti = self._needs_tool_intent(_msgs)
        _has_img = any(
            isinstance(m.get("content"), list)
            and any(isinstance(p, dict) and p.get("type") == "image_url"
                    for p in m["content"])
            for m in _msgs
        )
        _, _m, _ = self._route_model(_msgs, force_complex=(_ti or _has_img),
                                     reason="agent_worker")
        _vision_ok = _model_supports_vision(_m)
        # 统一走 _sanitize_msg_for_api——滤掉 tool/tool_log/None；视觉模型保留多模态
        # list（含图像），非视觉归一化为纯文本（list 原样发接口会 400）。
        hist = []
        for m in session.messages:
            sm = _sanitize_msg_for_api(m, vision_ok=_vision_ok)
            if sm:
                hist.append(sm)
        if len(hist) > self.cfg["max_history"]:
            hist = hist[-self.cfg["max_history"]:]
        messages = [sys_msg] + hist

        all_tools = config.get_all_tools(self.cfg)
        w = AgentWorker(self, messages, all_tools, config.mcp_clients,
                        task_id=resume_task_id, resume=resume)
        self._agent_worker = w
        w.status.connect(self.status_label.setText)
        w.render.connect(self._render_throttled)
        w.tool_log.connect(self._on_tool_log)
        w.tool_started.connect(self._on_tool_started)
        w.tool_finished.connect(self._on_tool_finished)
        w.deliverable_added.connect(self._on_deliverable_added)
        w.confirm_action.connect(self._on_confirm_action)
        w.schedule_reminder.connect(self._on_schedule_reminder)
        w.stream_begin.connect(self._on_stream_begin)
        w.stream_chunk.connect(self._on_stream_chunk)
        w.stream_commit.connect(self._on_stream_commit)
        w.done.connect(self._on_agent_done)
        w.started.connect(lambda: self.stop_btn.setVisible(True))
        w.finished.connect(lambda: self.stop_btn.setVisible(False))
        w.start()

    def _on_tool_log(self, entry):
        # v4.108 M-22：tool_log 的 args/result 截断存储——web_search/read_file 等
        # 全文可达数 KB，全量入库导致 sessions.json 无限膨胀（实测 1.3MB+），
        # 每次落盘/加载都拖慢。保留前段 + 省略号即可回溯工具干了什么。
        def _clip(s, n=500):
            s = str(s or "")
            return s if len(s) <= n else s[:n] + "…(截断)"
        self.store.active().messages.append({
            "role": "tool_log",
            "name": entry.get("name", ""),
            "args": _clip(entry.get("args", ""), 300),
            "result": _clip(entry.get("result", ""), 500),
        })
        self._save_throttled()  # v4.58：批量保存，避免每个工具一次磁盘 IO
        self._render_throttled()  # v4.58：节流渲染，避免信号洪水

    _TOOL_NAME_MAP = {
        "web_search": "联网搜索", "read_file": "读取文件", "write_file": "写入文件",
        "run_command": "执行命令", "run_python": "执行Python", "image_gen": "生成图片",
        "delete_file": "删除文件", "list_dir": "列目录", "rag_index": "索引文档",
        "rag_search": "知识库搜索", "screenshot": "截图",
        "browser_open": "浏览器打开", "browser_click": "浏览器点击",
        "browser_fill": "浏览器填表", "browser_read": "浏览器读文本",
        "remember": "写入长期记忆",
    }

    def _tool_card_html(self, ui_seq, display_name, args_preview="", done=False,
                        result_preview=""):
        """v4.108 M-25：统一生成 live 工具卡 HTML（running/done 两态共用）。"""
        if done:
            return (
                f'<div class="tool-card" id="live-tool-{ui_seq}" '
                f'style="margin-left: 20px;">'
                f'<span class="tool-dot done"></span>'
                f'<span class="tool-name">{display_name} ✓</span>'
                f'<span class="tool-result">{result_preview}</span>'
                f'</div>'
            )
        return (
            f'<div class="tool-card" id="live-tool-{ui_seq}">'
            f'<span class="tool-dot running"></span>'
            f'<span class="tool-name">{display_name}</span>'
            f'<span class="tool-args">{args_preview}</span>'
            f'</div>'
        )

    def _on_tool_started(self, data):
        name = data.get("name", "")
        args = data.get("args", {})
        agent_index = data.get("index", 0)

        display_name = self._TOOL_NAME_MAP.get(name, name)
        args_preview = ""
        if isinstance(args, dict):
            vals = [str(v) for v in args.values() if v]
            if vals:
                args_preview = html_mod.escape(vals[0][:30])
        # v4.108 M-25：UI 自管唯一 DOM id（agent index 只做 started↔finished 配对键）
        self._live_seq += 1
        ui_seq = self._live_seq
        self._live_tools[agent_index] = {
            "ui_seq": ui_seq, "name": name, "display_name": display_name,
            "args_preview": args_preview, "status": "running", "result": "",
        }
        self.chat_view.append(
            self._tool_card_html(ui_seq, display_name, args_preview))

    def _on_tool_finished(self, data):
        name = data.get("name", "")
        display_name = self._TOOL_NAME_MAP.get(name, name)
        result_preview = html_mod.escape(str(data.get("result_preview", ""))[:100])
        agent_index = data.get("index", 0)

        rec = self._live_tools.get(agent_index)
        if rec is not None and rec.get("name") == name:
            ui_seq = rec["ui_seq"]
            rec["status"] = "done"
            rec["result"] = result_preview
        else:
            # 找不到 started（历史/直发 finished）：兜底追加一张 done 卡，不丢信息
            self._live_seq += 1
            ui_seq = self._live_seq
            self._live_tools[agent_index] = {
                "ui_seq": ui_seq, "name": name, "display_name": display_name,
                "args_preview": "", "status": "done", "result": result_preview,
            }
        # 原位替换 running 卡 → done 卡（同一 id，不追加新 DOM，零重复）
        self.chat_view.replace_live(
            f"live-tool-{ui_seq}",
            self._tool_card_html(ui_seq, display_name, done=True,
                                 result_preview=result_preview))

        # v4.84：模型刚提交了一个待审核技能，刷新徽标并提示
        if name == "create_skill":
            self._refresh_skill_review_label()
            try:
                import skill_review
                n = skill_review.count_pending(self.cfg)
                if n > 0 and getattr(self, "tray_app", None):
                    self.tray_app.tray.showMessage(
                        "技能审核", f"新技能已提交到审核队列，共 {n} 个待你通过。",
                        QSystemTrayIcon.Information, 6000)
            except Exception:
                pass

    def _on_confirm_action(self, title, detail):
        w = self._agent_worker
        if self.permission_engine.session_trusted:
            w._confirm_val = True
            w._confirm_event.set()
            return
        dlg = ConfirmDialog(self)
        dlg.set_text(title, f"{title}：\n{detail}\n\n允许执行？")
        dlg.exec()
        w._confirm_val = dlg.result()
        if dlg.trusted():
            self.permission_engine.set_session_trusted()
            self.session_trusted = True
        w._confirm_event.set()
        self._clear_ghost()

    def _on_stream_begin(self):
        self._streaming = True
        self._streaming_text = ""
        self._render_throttled()

    def _on_stream_chunk(self, text):
        self._streaming_text = text
        self._render_throttled()  # v4.58：节流渲染，~200 chunk → ~20 次 setHtml

    def _track_context(self, role, content):
        """轻量同步消息到上下文管理器（模块4），不侵入 Session 本身。"""
        try:
            if not self.cfg.get("context_enabled", True):
                return
            if isinstance(content, list):
                content = "".join(
                    p.get("text", "") for p in content
                    if isinstance(p, dict) and p.get("type") == "text")
            if not isinstance(content, str):
                return
            from context_manager import get_context_manager
            sid = self.store.active_sid if self.store else None
            get_context_manager(sid).add_message(role, content)
        except Exception:
            pass

    def _on_stream_commit(self, text):
        self._streaming = False
        msgs = self.store.active().messages
        # v4.60o：末条防重——渲染竞态或循环重提同一句时，避免重复气泡
        if (msgs and msgs[-1].get("role") == "assistant"
                and msgs[-1].get("content") == text):
            self.chat_view.end_stream()  # v4.104 fix：先清流式残留再重渲染
            self._render_throttled()
            self._speak(text)
            return
        msgs.append({"role": "assistant", "content": text})
        self._track_context("assistant", text)
        self._save_throttled()  # v4.58：批量保存
        # v4.108 M-25：不再强制全量重建（旧 v4.96 方案）——jsRenderAll 会清空 DOM，
        # 把只活在 DOM 里的 live 工具卡（running/done）一并抹掉，随后 replace_live
        # 找不到目标静默丢弃。增量渲染计数不变 → 只 append 本句 + end_stream 清流式
        # 残留（v4.104 增量路径已覆盖清理），工具卡 DOM 原样保留。
        self._render_throttled()  # v4.58：节流渲染
        self._speak(text)

    def _request_agent_stop(self):
        if self._agent_worker and self._agent_worker.isRunning():
            self._agent_worker.request_stop()
            self.status_label.setText("⏹ 正在停止 Agent…（当前工具完成后生效）")
            self.stop_btn.setEnabled(False)

    def _on_agent_done(self):
        self._agent_active = False
        self._reset_busy()
        self.stop_btn.setVisible(False)
        self.stop_btn.setEnabled(True)
        self.resume_agent_btn.setVisible(False)
        w = self._agent_worker
        stopped = getattr(w, "stopped_by_user", False)
        if stopped:
            # v4.101：用户暂停 → 检查点已标记 paused，显示「继续上次任务」入口
            tid = getattr(w, "task_id", None)
            if tid and task_resume.load_checkpoint(self.cfg, tid):
                self._agent_resume_task_id = tid
                self.resume_agent_btn.setVisible(True)
                self.status_label.setText("⏸ 任务已暂停（可点「继续上次任务」接着干）")
            else:
                self.status_label.setText("Agent 完成")
        else:
            self.status_label.setText("Agent 完成")
        self._do_save()  # v4.58：agent 结束做一次最终落盘
        self._flush_render()  # v4.60o：冲刷待渲染定时器，避免末条消息重复插入
        self.input_box.setFocus()
        # v4.95：自动化任务完成 → 托盘+语音通知
        _pending = getattr(self, "_pending_done_notify", None)
        if _pending:
            self._pending_done_notify = None
            self._notify_task_done("任务完成", f"自动化任务「{_pending}」已完成")

    # ---- v4.101：普通 Agent 断点续传 ----
    def _resume_agent_task(self):
        """「继续上次任务」：复用同一检查点 task_id 重建 worker，从会话断点接着跑。"""
        if getattr(self, "_agent_active", False):
            return
        tid = getattr(self, "_agent_resume_task_id", None)
        if not tid:
            return
        self.resume_agent_btn.setVisible(False)
        self._agent_run(resume=True, resume_task_id=tid)

    def _cleanup_paused_agent_checkpoints(self, except_tid=None):
        """开启新 Agent 任务前，清理本会话残留的「已暂停」检查点（用户放弃旧任务）。
        except_tid：续传时跳过正在恢复的那个检查点。"""
        try:
            sess = self.store.active()
            sid = sess.sid if sess else None
        except Exception:
            sid = None
        try:
            for cp in task_resume.list_active(self.cfg):
                if cp.get("task_type") != "agent":
                    continue
                if cp.get("sid") != sid:
                    continue
                if cp.get("_task_id") == except_tid:
                    continue
                task_resume.mark_done(self.cfg, cp["_task_id"])
        except Exception:
            pass

    def _scan_agent_resume(self):
        """切换会话/启动时扫描本会话是否有已暂停的 Agent 任务，有则提示「继续」。"""
        try:
            sess = self.store.active()
            sid = sess.sid if sess else None
        except Exception:
            sid = None
        if getattr(self, "_agent_active", False):
            return
        btn = getattr(self, "resume_agent_btn", None)
        if btn is None:
            return
        try:
            for cp in task_resume.list_active(self.cfg):
                if cp.get("task_type") == "agent" and cp.get("sid") == sid:
                    self._agent_resume_task_id = cp.get("_task_id")
                    btn.setVisible(True)
                    self.status_label.setText("⏸ 检测到上次暂停的 Agent 任务，可点「继续上次任务」")
                    return
        except Exception:
            pass
        btn.setVisible(False)

    def _reset_busy(self):
        """重置忙碌状态（可由超时触发，防止永久卡住）。
        v4.58：超时时同时杀掉 agent worker + 中止流式 reply，防止后台僵尸线程。
        """
        self._busy_timeout.stop()
        if self._agent_worker and self._agent_worker.isRunning():
            self._agent_worker.request_stop()
        if self._reply is not None:
            try:
                self._reply.abort()
            except Exception:
                pass
        self._busy = False
        self.send_btn.setEnabled(True)
        self._on_input_changed()
    def _route_model(self, messages, force_complex=False, reason=""):
        """v4.94+v4.98 模型智能路由：默认主模型，复杂任务或工具意图升级到 complex_model。
        返回 (base_url, model, api_key)。complex_model 未配置 api_key 时回退主模型。
        force_complex=True 时直接强制走 complex_model（工具意图：杜绝弱模型退化成文字演工具）。

        v4.109 两处增量（均不改变 Auto 默认行为）：
        - reason：调用方可传入升舱归因（tool_intent / image / force），写入旁路日志。
        - 手动锁定 self._model_lock：下拉菜单选了 Auto 之外的档位时全程锁定，
          完全绕过智能路由；锁定档位失效（key 被删/配置不全）则回退主模型，
          **绝不静默跳到付费通道**。
        """
        cfg = self.cfg
        base_url = cfg.get("base_url", "")
        model = cfg.get("model", "")
        api_key = cfg.get("api_key", "")
        _upgraded = False
        _why = reason or ""

        # ---- v4.109 手动锁定优先：选了具体模型就不再自动切换 ----
        _lock = getattr(self, "_model_lock", "") or ""
        if _lock:
            if _lock == "__main__":
                _why = "manual_lock:main"
            else:
                _prof = (cfg.get("model_profiles") or {}).get(_lock) or {}
                if _prof.get("api_key") and _prof.get("model") and _prof.get("base_url"):
                    base_url = _prof["base_url"]
                    model = _prof["model"]
                    api_key = _prof["api_key"]
                    _why = "manual_lock:" + _lock
                else:
                    _why = "lock_invalid:fallback_main"
            self._route_reason = _why
            self._log_route(messages, model, base_url, False, _why)
            return base_url, model, api_key

        routing = cfg.get("model_routing") or {}
        if not routing.get("enabled", True):
            self._route_reason = "routing_disabled"
            self._log_route(messages, model, base_url, False, "routing_disabled")
            return base_url, model, api_key
        if not (force_complex or self._is_complex(messages, routing)):
            self._route_reason = "default"
            self._log_route(messages, model, base_url, False, "default")
            return base_url, model, api_key
        prof_name = routing.get("complex_model", "")
        prof = (cfg.get("model_profiles") or {}).get(prof_name) or {}
        if prof.get("api_key") and prof.get("model") and prof.get("base_url"):
            if not _why:
                _why = ("force_complex" if force_complex
                        else self._complexity_reason(messages, routing))
            base_url = prof["base_url"]
            model = prof["model"]
            api_key = prof["api_key"]
            _upgraded = True
        else:
            _why = "complex_profile_unset"
        self._route_reason = _why
        self._log_route(messages, model, base_url, _upgraded, _why)
        return base_url, model, api_key

    def _complexity_reason(self, messages, routing):
        """v4.109：复杂度判定的具体归因（命中哪个关键词 / 长度超阈值），供旁路日志复盘。"""
        hints = routing.get("complex_hint") or []
        threshold = routing.get("length_threshold", 1500)
        total = 0
        for m in messages or []:
            c = m.get("content", "")
            if not isinstance(c, str):
                continue
            total += len(c)
            for h in hints:
                if h in c:
                    return "kw:" + h
        if total > threshold:
            return "len:%d>%s" % (total, threshold)
        return "complex"

    def _log_route(self, messages, model, base_url, upgraded, reason):
        """v4.109 旁路路由日志：只记录不干预，任何异常一律吞掉。"""
        try:
            _n = 0
            for _m in messages or []:
                _c = _m.get("content")
                if isinstance(_c, str):
                    _n += len(_c)
            route_log.log_route(event="route", model=model, base_url=base_url,
                                upgraded=upgraded, reason=reason,
                                lock=getattr(self, "_model_lock", "") or "",
                                msgs_len=_n)
        except Exception:
            pass

    def _is_complex(self, messages, routing):
        """复杂任务判定：命中关键词 或 消息总长度超阈值。"""
        hints = routing.get("complex_hint") or []
        threshold = routing.get("length_threshold", 1500)
        total = 0
        for m in messages or []:
            c = m.get("content", "")
            if not isinstance(c, str):
                continue
            total += len(c)
            if any(h in c for h in hints):
                return True
        return total > threshold

    def _needs_tool_intent(self, messages):
        """v4.98 工具意图检测：用户明显要 Agent 用工具干活（改/写/编辑文件、跑 python、
        搜索、生图生视频、读取分析文件、调用工具/自动化等）。这类任务弱模型（Agnes）极易
        退化成"文字演工具"（伪造 [工具]/✅ 已保存/run_python(），直接强制走 complex_model
        并把 tool_choice 设为 required，从根上杜绝撒谎。命中最后一条用户原话即视为工具意图。"""
        KEYWORDS = (
            # 文件读写
            "写文件", "保存", "导出", "生成文件", "创建文件", "新建文件", "建个文件",
            "改文件", "编辑文件", "修改文件", "打开文件", "读取文件", "读文件", "写入",
            # 代码执行
            "运行", "跑", "执行", "python", "代码", "脚本", "py 脚本",
            # 搜索
            "搜索", "查一下", "上网查", "fetch", "爬虫", "爬取",
            # 生图生视频
            "生图", "画图", "画一张", "画图片", "生成图片", "生成一张", "一张图片",
            "一张图", "作图", "生视频", "生成视频", "做视频", "口播视频",
            "数字人视频", "数字人口播", "剪辑", "配音", "旁白", "做成视频", "做口播",
            "重新生成", "再画", "换一张", "重画",
            # 工具/自动化
            "调用工具", "用工具", "自动化", "定时", "提醒",
            # 数据分析
            "分析", "统计", "报表", "数据处理", "excel", "表格", "csv",
            # 自省/能力盘点（命中即强制走 DeepSeek 并 required，确保真调 sys_info 工具，
            # 避免弱模型退化成吐文本、被循环误判为『没调工具』）
            "你的能力", "你会什么", "你能做什么", "你会哪些", "能力清单", "功能清单",
            "有哪些功能", "有哪些工具", "会干啥", "能干什么", "能调用什么",
            "介绍一下你自己", "介绍你自己", "自我介绍", "列个清单", "清单给我",
            "你有什么功能", "你能调用",
        )
        last_user = ""
        for msg in reversed(messages or []):
            if msg.get("role") == "user":
                c = msg.get("content", "")
                if isinstance(c, str):
                    last_user = c
                break
        if not last_user:
            return False
        # v4.102 fix11：咨询问句豁免——「配音怎么学 / 剪辑怎么入门 / 数字人是什么」这类
        # 纯咨询不该被当成工具意图升舱，否则误走 DeepSeek + required 空转。
        # 命中这些非工具意图问句特征直接返回 False（交由普通对话直答）。
        if self._looks_like_learning_question(last_user):
            return False
        if any(kw in last_user.lower() for kw in KEYWORDS):
            return True
        # v4.102 fix11：媒体生成组合判定兜底——"生成口播视频 / 做个数字人"这类
        # 分开写的表述连写词表会漏判，用「对象词×动作词」同现识别。
        if self._is_media_gen_request(last_user):
            return True
        # v4.107：导演指令已从主路由摘除（统一在导演台底部对话条处理），不再升舱。
        return False

    def _user_refuses_tools(self, messages):
        """v4.100：检测用户是否明确要求『不要调用工具/纯聊天』。
        若用户说过此类约束，且其后没有下达新的明确工具指令，则本轮禁止调工具
        （由调用方设 tool_choice=none），尊重用户约束，避免闲聊被 remember 等
        工具自发调用打断。
        判断依据：扫描全部消息，取最后一次『拒绝调工具』与最后一次『明确工具动作』
        的位置——若拒绝在动作之后（或从未有动作），视为当前仍应尊重『不调工具』。"""
        REFUSE_KW = (
            "不要调用工具", "不要使用工具", "别调用工具", "别用工具", "不用工具",
            "不要调工具", "别调工具", "纯聊天", "只是聊", "只是聊天", "光聊天",
            "不要动工具", "先别用工具", "不要开工具", "不用开工具", "别开工具",
        )
        ACTION_KW = (
            "搜", "写", "生成", "运行", "执行", "python", "创建", "导出", "分析",
            "读文件", "打开文件", "读一下", "做图", "生图", "生视频", "截图", "下载",
            "安装", "提醒", "调用工具", "用工具", "查一下", "上网查", "爬", "整理",
            "发", "监控", "剪辑", "配音", "做视频", "画图",
        )
        last_refuse = -1
        last_action = -1
        for i, m in enumerate(messages or []):
            if m.get("role") != "user":
                continue
            c = m.get("content", "")
            if not isinstance(c, str):
                continue
            if any(k in c for k in REFUSE_KW):
                last_refuse = i
            if any(k in c for k in ACTION_KW):
                last_action = i
        return last_refuse >= 0 and last_refuse > last_action

    def _looks_like_learning_question(self, text):
        """v4.102 fix11：判断是否为「学习/科普/了解的纯咨询」问句。
        『配音怎么学』『剪辑怎么入门』『数字人是什么』『AI视频怎么做的』这类句子里
        虽含「配音/剪辑/视频」等词，但用户是在**询问知识**而非**让我执行制作**，
        不应视为工具意图（否则误升舱 DeepSeek + 强制 required 空转）。
        命中返回 True → 上层不做工具意图判定。"""
        if not text:
            return False
        t = text.lower()
        # 咨询/学习问句特征：怎么学 / 怎么入门 / 是什么 / 什么意思 / 如何 / 会不会 /
        # 了解下 / 介绍一下 / 原理 / 教程 / 怎么弄出来的
        if any(k in t for k in ("怎么学", "怎么入门", "如何学", "如何入门", "怎么开始",
                                "是什么", "什么意思", "啥意思", "了解一下", "了解下",
                                "介绍下", "介绍一下", "怎么弄的", "怎么做出来",
                                "原理", "教程", "怎么来的", "会不会", "能不能学",
                                "学习路径", "从哪学", "选哪个", "推荐学习")):
            return True
        return False

    def _is_reasoning_model(self, model, base_url=""):
        """v4.102 fix10：判断当前模型是否为「思考/推理模式」模型。
        DeepSeek 官方推理模型（例如 deepseek-v4-flash-vision-exp）不支持 tool_choice="required"
        ——一旦工具意图命中（如用户说「执行任务」「运行」）被 v4.98 强制设 required，
        API 直接返回 400：`Thinking mode does not support this tool_choice`，
        异常又被 _agent_call 的 except 吞掉 → 空 content → 界面「Agent 完成」但无任何输出。
        判断依据：DeepSeek 官方 base_url + 模型含推理/v4 特征；模型名含 think/reason/r1 等。
        命中时上层不再强制 required，改为让模型自由决定（默认/auto）。
        """
        m = (model or "").lower()
        b = (base_url or "").lower()
        # 明确推理/思考特征
        if any(k in m for k in ("think", "reason", "-r1", "reasoning", "thinking")):
            return True
        # DeepSeek 官方通道（api.deepseek.com）当前推理模型均为思考模式
        if "api.deepseek.com" in b:
            return True
        return False

    def _agent_call(self, messages, tools, on_delta=None, force_required=False, force_tool=None, force_complex=False):
        import urllib.request as urllib_req
        import logging as _logging
        _tool_intent = self._needs_tool_intent(messages)
        # v4.102：消息含图 → 强制走视觉模型通道（复杂/工具意图路径），让模型真正"看图"
        _has_img_call = any(
            isinstance(m.get("content"), list)
            and any(isinstance(p, dict) and p.get("type") == "image_url" for p in m.get("content", []))
            for m in messages
        )
        # 路由：工具意图或含图都升舱到 complex_model（视觉模型）
        _route_force = _tool_intent or _has_img_call or force_complex
        _refuse_tools = self._user_refuses_tools(messages)
        # v4.109：升舱归因——记录"是谁把这次调用送进了付费通道"，供旁路日志复盘
        _route_reason = ("tool_intent" if _tool_intent
                         else ("image" if _has_img_call
                               else ("force_complex" if force_complex else "")))
        _base_url, _model, _api_key = self._route_model(
            messages, force_complex=_route_force, reason=_route_reason)
        url = _base_url.rstrip("/") + "/chat/completions"
        body = {
            "model": _model,
            "messages": messages,
            "tools": tools,
            "stream": True,
            "temperature": 0.3,
        }
        _REDO_KEYWORDS = (
            "重新生成", "重新画", "重新搜索", "再生成", "再画", "再来一张",
            "再来一次", "再查", "再搜索", "再搜", "regenerate", "redo", "重画",
            "重新试", "重新来", "再试试", "重新", "换个", "换一张",
        )
        redo = False
        for msg in reversed(messages):
            if msg.get("role") == "user":
                last_user = msg.get("content", "")
                if isinstance(last_user, str) and any(kw in last_user for kw in _REDO_KEYWORDS):
                    redo = True
                break
        # v4.60：强制调用指定工具（如 sys_info / video_gen），优先级最高
        if force_tool:
            if self._is_reasoning_model(_model, _base_url):
                # v4.102 fix11：DeepSeek 思考模式不支持任何 tool_choice 自定义——
                # 连指定函数 {"function":{"name":force_tool}} 也返回 400
                # 「Thinking mode does not support this tool_choice」（fix10 只豁免了
                # "required"，这里连 force_tool 也要豁免，否则视频/生图首步直接 400 → 空响应）。
                # 改为在消息尾部注入强制指令，让思考模型自然决定调用指定工具，
                # 不设 tool_choice（思考模式默认行为）。
                _ft_instr = (
                    f"当前任务必须通过调用工具 {force_tool} 完成。"
                    f"请直接调用 {force_tool} 工具：把用户请求的全部必要信息整理为它的参数"
                    f"（如 prompt / duration / aspect / dialogue 等），一次调用它并生成结果。"
                    f"不要调用其他无关工具，也不要只描述计划，必须真实调用 {force_tool}。"
                )
                body["messages"] = list(messages) + [
                    {"role": "user", "content": _ft_instr, "_internal": True},
                ]
            else:
                body["tool_choice"] = {"type": "function", "function": {"name": force_tool}}
        elif _refuse_tools:
            # v4.100：用户明确要求"不调工具/纯聊天"且其后无新工具指令时，
            # 本轮彻底禁止调工具（即便命中工具意图也尊重用户约束），避免
            # 闲聊被 remember 等工具自发调用打断。
            body["tool_choice"] = "none"
        elif (force_required or redo or _tool_intent) and not self._is_reasoning_model(_model, _base_url):
            # v4.98：工具意图任务强制 required，杜绝弱模型退化成"文字演工具"
            # v4.102 hotfix：仅因含图进入视觉模型时，不强制 required，让模型自由描述图片。
            # v4.102 fix10：思考/推理模式模型（DeepSeek 官方）不支持 required——
            # 命中时降级为不设 tool_choice（默认/auto），让模型自然决定输出文本或调工具，
            # 否则 API 返回 400「Thinking mode does not support this tool_choice」→ 空 content。
            body["tool_choice"] = "required"
        # v4.102 fix12：请求 usage 统计——多数 OpenAI 兼容通道在末个 chunk 返回 usage。
        # 少数通道不认识该参数会返回 400，下方 _stream_once 会自动去掉参数重试一次，
        # 因此新增参数永远不会让原本能跑通的调用失败。
        body["stream_options"] = {"include_usage": True}
        payload = json.dumps(body, ensure_ascii=False).encode("utf-8")

        full_content = ""
        tool_acc = {}  # index -> {"id", "function": {"name", "arguments"}}
        _usage = {}

        def _stream_once(_body):
            """发一次流式请求，返回 (content, tool_acc, usage)；失败时向外抛异常。"""
            _payload = json.dumps(_body, ensure_ascii=False).encode("utf-8")
            _req = urllib_req.Request(url, data=_payload, method="POST")
            _req.add_header("Content-Type", "application/json")
            _req.add_header("Authorization", f"Bearer {_api_key}")
            _req.add_header("Accept", "text/event-stream")
            _content = ""
            _acc = {}
            _u = {}
            with urllib_req.urlopen(_req, timeout=90) as resp:
                buf = ""
                for raw in resp:
                    buf += raw.decode("utf-8", "ignore")
                    # v4.58：防 SSE 缓冲区无限膨胀（代理乱码/网络异常时）
                    if len(buf) > 1_000_000:
                        break
                    while "\n" in buf:
                        line, buf = buf.split("\n", 1)
                        line = line.strip()
                        if not line or not line.startswith("data:"):
                            continue
                        data = line[len("data:"):].strip()
                        if data == "[DONE]":
                            break
                        try:
                            evt = json.loads(data)
                        except Exception:
                            continue
                        # v4.102 fix12：末个 chunk 携带 usage（prompt/completion tokens）
                        if evt.get("usage"):
                            _u = evt["usage"]
                        choice = (evt.get("choices") or [{}])[0]
                        delta = choice.get("delta") or {}
                        if delta.get("content") is not None:
                            _content += delta["content"]
                            if on_delta:
                                on_delta(_content)
                        for tc in (delta.get("tool_calls") or []):
                            idx = tc.get("index", 0)
                            acc = _acc.setdefault(idx, {"id": "", "function": {"name": "", "arguments": ""}})
                            if tc.get("id"):
                                acc["id"] = tc["id"]
                            fn = tc.get("function") or {}
                            # v4.108 M-18：兼容网关会多 chunk 重复发 function.name，
                            # 只取首次，禁止 += 拼接（否则聚合出 "web_sweb_search"）。
                            if fn.get("name") and not acc["function"].get("name"):
                                acc["function"]["name"] = fn["name"]
                            if fn.get("arguments"):
                                acc["function"]["arguments"] += fn["arguments"]
            return _content, _acc, _u

        try:
            full_content, tool_acc, _usage = _stream_once(body)
        except Exception as e:
            # v4.102 fix12：通道不认识 stream_options 时（多为 HTTP 400），
            # 去掉该参数重试一次——保持 fix12 之前的行为，绝不因新参数导致调用失败。
            _code = getattr(e, "code", None)
            if _code in (400, 404) or "stream_options" in str(e):
                try:
                    body.pop("stream_options", None)
                    full_content, tool_acc, _usage = _stream_once(body)
                except Exception as e2:
                    # v4.108 H-04：失败必须上抛交给 agent.py 兜底弹错，不能吞掉装"成功"。
                    _logging.getLogger("dsdesktop").error("Agent 流式调用失败: %s", e2)
                    raise
            elif _code in (429, 500, 502, 503, 504):
                # v4.108 H-04：限流/网关抖动 → 退避重试 2 次（2s/4s），仍失败则上抛。
                _last = e
                for _attempt in range(2):
                    time.sleep(2 * (_attempt + 1))
                    try:
                        full_content, tool_acc, _usage = _stream_once(body)
                        _last = None
                        break
                    except Exception as e3:
                        _last = e3
                if _last is not None:
                    _logging.getLogger("dsdesktop").error("Agent 流式调用重试仍失败: %s", _last)
                    raise
            else:
                # v4.108 H-04：其余失败（超时/断流等）同样上抛，禁止静默返回空响应。
                _logging.getLogger("dsdesktop").error("Agent 流式调用失败: %s", e)
                raise

        tool_calls = []
        for idx in sorted(tool_acc):
            t = tool_acc[idx]
            try:
                args = json.loads(t["function"]["arguments"] or "{}")
            except Exception:
                args = {}
            tool_calls.append({
                "id": t.get("id") or f"call_{idx}",
                "type": "function",
                "function": {
                    "name": t["function"]["name"],
                    "arguments": json.dumps(args, ensure_ascii=False),
                },
            })
        # v4.102 fix12：token 用量回传，供 agent.py 做预算熔断统计。
        # API 未返回 usage 时用字符数**保守估算**（宁可高估，防烧钱）：
        # 中文约 1.5 字符/token、英文约 4 字符/token，统一按 2 字符/token 估。
        try:
            _pt = int(_usage.get("prompt_tokens") or 0)
            _ct = int(_usage.get("completion_tokens") or 0)
        except (TypeError, ValueError):
            _pt = _ct = 0
        if not (_pt or _ct):
            _est_in = 0
            for _m in messages:
                _c = _m.get("content")
                if isinstance(_c, str):
                    _est_in += len(_c)
                elif isinstance(_c, list):
                    _est_in += sum(len(str(p.get("text", "")))
                                   for p in _c if isinstance(p, dict))
            _est_out = len(full_content) + sum(
                len(str(tc.get("function", {}).get("arguments", "")))
                for tc in tool_calls)
            _pt = _est_in // 2
            _ct = max(_est_out // 2, 1)
        # v4.109：旁路成本日志（只记录，不干预；拿真实 token 才能算清付费通道花了多少）
        try:
            route_log.log_route(event="usage", model=_model, base_url=_base_url,
                                prompt_tokens=_pt, completion_tokens=_ct,
                                total_tokens=_pt + _ct,
                                reason=getattr(self, "_route_reason", ""),
                                lock=getattr(self, "_model_lock", "") or "")
        except Exception:
            pass
        return {
            "content": full_content,
            "tool_calls": tool_calls,
            "usage": {"prompt_tokens": _pt, "completion_tokens": _ct,
                      "total_tokens": _pt + _ct},
            "model": _model,
            "channel": "deepseek" if "deepseek" in (_base_url or "").lower() else "other",
        }

    def _on_schedule_reminder(self, delay_ms, message, repeat_secs=0):
        if repeat_secs <= 0:
            QTimer.singleShot(delay_ms, lambda: self._fire_reminder(message))
            return
        # 循环定时：首次 delay_ms 后触发，之后每 repeat_secs 秒触发
        def _fire_repeat():
            self._fire_reminder(message)
            QTimer.singleShot(repeat_secs * 1000, _fire_repeat)
        QTimer.singleShot(delay_ms, _fire_repeat)

    def _fire_reminder(self, message):
        dlg = InfoDialog(self)
        dlg.set_text("定时提醒", message)
        dlg.exec()
        self.status_label.setText(f"定时提醒已触发：{message}")
        self._clear_ghost()

    # ============ 流式生成 ============
    def _start_stream(self, text, search_context):
        self._agent_active = False  # 普通对话模式：提示词按普通模式构建
        self._busy = True
        self.send_btn.setEnabled(False)
        self._busy_timeout.start(120000)
        self.status_label.setText(
            "联网资料已就绪，生成中…" if search_context else "搜索无结果，使用模型知识回答")

        session = self.store.active()
        sys_msg = {"role": "system", "content": self._build_system_prompt()}

        # v4.102 图像输入链路：用户发图时，强制路由到视觉模型（complex_model profile，
        # 现指向 deepseek-v4-flash-vision-exp），并保留图像 content 让模型真正"看图"；
        # 非视觉模型（如 Agnes）仍把图归一化为纯文本，避免 list 原样发送导致 400。
        last_msg = session.messages[-1] if session.messages else None
        _has_img = bool(
            last_msg and isinstance(last_msg.get("content"), list)
            and any(isinstance(p, dict) and p.get("type") == "image_url"
                    for p in last_msg["content"])
        )
        if _has_img:
            # 带图 → 走视觉模型通道（复杂/工具意图路径，强制 complex）
            _bu, _m, _k = self._route_model(session.messages, force_complex=True,
                                            reason="image")
            # v4.102：视觉模型思考阶段会先吐 reasoning_content（content 为空，可能持续
            # 10-30秒），给用户明确反馈避免误以为「卡死无回复」。
            self.status_label.setText("看图中…（视觉模型思考中，请稍候 10-30 秒）")
        else:
            _bu, _m, _k = self.cfg["base_url"], self.cfg["model"], self.cfg["api_key"]
            # v4.109：手动锁定在这条路径上生效（Auto 时保持原样——普通对话仍只走主模型，
            # 不因为命中 complex_hint 关键词就升舱付费，行为与 v4.108 完全一致）。
            if getattr(self, "_model_lock", ""):
                _bu, _m, _k = self._route_model(session.messages, reason="stream_lock")
        _vision_ok = _model_supports_vision(_m)
        if _has_img and not _vision_ok and getattr(self, "_model_lock", ""):
            # 手动锁定了非视觉模型 + 用户发图：图会被归一化成纯文本，明确告知避免困惑
            try:
                self.status_label.setText(
                    "已锁定 %s（非视觉模型）：本轮图片按文字描述处理，看图请切回 Auto" % _m)
            except Exception:
                pass

        # v4.79 hotfix：历史必须经 _sanitize_msg_for_api 清洗——session 里混有
        # tool/tool_log/None/list 内容（UI 展示用），直接发接口会 400。
        # 视觉模型保留图像列表，非视觉归一化为纯文本。
        others = []
        for m in session.messages[:-1]:
            sm = _sanitize_msg_for_api(m, vision_ok=_vision_ok)
            if sm:
                others.append(sm)
        if len(others) > self.cfg["max_history"]:
            others = others[-self.cfg["max_history"]:]
        api_messages = [sys_msg] + others
        if search_context:
            api_messages.append({"role": "system", "content": search_context})

        # 最后一条用户消息：视觉模型保留图，非视觉归一化纯文本
        if last_msg and isinstance(last_msg.get("content"), list):
            lc = last_msg["content"]
            if _has_img and _vision_ok:
                # 真·多模态（含图片）：保留原样发给视觉模型，但图片统一重编码为 RGB JPEG
                # v4.102 fix6：兜底各种来源（贴图/附件/历史），避免特殊格式/超大图被拒
                norm_lc = []
                for _part in lc:
                    if isinstance(_part, dict) and _part.get("type") == "image_url":
                        _u = (_part.get("image_url") or {}).get("url", "")
                        norm_lc.append({"type": "image_url", "image_url": {"url": _normalize_image_dataurl(_u)}})
                    else:
                        norm_lc.append(_part)
                api_messages.append({"role": "user", "content": norm_lc})
            else:
                # 纯文本 list（语音/粘贴等）或非视觉模型 → 归一化为字符串
                api_messages.append({"role": "user", "content": _flatten_text_content(lc) or text})
        else:
            api_messages.append({"role": "user", "content": text})

        url = QUrl(_bu.rstrip("/") + "/chat/completions")
        req = QNetworkRequest(url)
        req.setHeader(QNetworkRequest.ContentTypeHeader, "application/json")
        req.setRawHeader(b"Authorization", f"Bearer {_k}".encode("utf-8"))
        req.setRawHeader(b"Accept", b"text/event-stream")
        req.setRawHeader(b"Cache-Control", b"no-cache")
        payload = json.dumps({
            "model": _m,
            "messages": api_messages,
            "stream": True,
            "temperature": 0.7,
        }, ensure_ascii=False).encode("utf-8")

        self._streaming = True
        self._streaming_text = ""
        self._sse_buf = ""
        self._streaming_error = ""
        _img_n = 0; _img_b = 0
        for _am in api_messages:  # v4.102 fix7：循环变量改名，勿覆盖模型名变量 _m
            _c = _am.get("content")
            if isinstance(_c, list):
                for _p in _c:
                    if isinstance(_p, dict) and _p.get("type") == "image_url":
                        _img_n += 1
                        _u = (_p.get("image_url") or {}).get("url", "")
                        if _u.startswith("data:"):
                            _img_b += len(_u)
        _vision_debug(f"_start_stream post: model={_m} has_img={_has_img} vision_ok={_vision_ok} msgs={len(api_messages)} imgs={_img_n} img_bytes={_img_b}B payload={len(payload)}B url={url.toString()}")
        try:
            reply = self.manager.post(req, QByteArray(payload))
            self._reply = reply  # v4.102: 存到实例，让 _reset_busy 超时能 abort 流式请求
            reply.readyRead.connect(lambda: self._on_stream_ready(reply))
            reply.finished.connect(lambda: self._on_stream_finished(reply))
        except Exception as e:
            _vision_debug(f"_start_stream post RAISED: {type(e).__name__}: {e}")
            self._streaming = False
            self._streaming_text = ""
            self._streaming_error = str(e)
            self._reset_busy()
            self.status_label.setText(f"发送失败：{e}")

    def _on_stream_ready(self, reply):
        chunk = bytes(reply.readAll()).decode("utf-8", "ignore")
        self._sse_buf += chunk
        # v4.58：防 SSE 缓冲区无限膨胀（与 _agent_call 同步流保持一致）
        if len(self._sse_buf) > 1_000_000:
            self._sse_buf = ""
        self._drain_sse()

    def _drain_sse(self):
        while "\n\n" in self._sse_buf:
            raw_event, self._sse_buf = self._sse_buf.split("\n\n", 1)
            for line in raw_event.split("\n"):
                line = line.strip()
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                if data == "[DONE]":
                    continue
                try:
                    obj = json.loads(data)
                except Exception:
                    continue
                # v4.102 hotfix：捕获 API 返回的错误事件（如图片不支持/超限/
                # 模型临时不可用），否则会被当成「成功但空内容」静默吞掉。
                if isinstance(obj, dict):
                    err = obj.get("error")
                    if err and isinstance(err, dict) and err.get("message"):
                        self._streaming_error = (self._streaming_error + " " + err["message"]).strip()
                        continue
                    if err and isinstance(err, str):
                        self._streaming_error = (self._streaming_error + " " + err).strip()
                        continue
                try:
                    delta_obj = obj["choices"][0]["delta"]
                except Exception:
                    continue
                # v4.102：视觉模型（deepseek-v4-flash-vision-exp）会先吐 reasoning_content
                # （思考阶段，content 为空/null），可能持续很多秒。期间给用户反馈，避免
                # 误以为「发图后无回复卡死」。
                rc = delta_obj.get("reasoning_content")
                if rc and not delta_obj.get("content"):
                    cur = self.status_label.text()
                    if "思考中" not in cur:
                        self.status_label.setText("思考中…（视觉模型分析图片，请稍候）")
                delta = delta_obj.get("content") or ""
                if delta:
                    self._streaming_text += delta
            self._render_messages()

    def _extract_api_error(self, buf):
        """v4.102 hotfix：从 SSE/原始缓冲里提取 API 返回的错误信息
        （data:{"error":{...}} 或 data:{"error":"..."}）。无则返回空串。"""
        if not buf:
            return ""
        for raw_event in buf.split("\n\n"):
            for line in raw_event.split("\n"):
                line = line.strip()
                if not line.startswith("data:"):
                    continue
                data = line[5:].strip()
                try:
                    obj = json.loads(data)
                except Exception:
                    continue
                if isinstance(obj, dict):
                    err = obj.get("error")
                    if isinstance(err, dict) and err.get("message"):
                        return err["message"]
                    if isinstance(err, str):
                        return err
        return ""

    def _on_stream_finished(self, reply):
        try:
            self._on_stream_finished_impl(reply)
        except Exception as e:
            _vision_debug(f"_on_stream_finished RAISED: {type(e).__name__}: {e}")
            log.error("_on_stream_finished 异常: %s", e)
            try:
                self._streaming = False
                self._streaming_text = ""
                self._reset_busy()
                self.status_label.setText(f"回复处理异常：{e}")
            except Exception:
                pass
            try:
                reply.deleteLater()
            except Exception:
                pass

    def _on_stream_finished_impl(self, reply):
        status_code = reply.attribute(QNetworkRequest.HttpStatusCodeAttribute)
        tail = bytes(reply.readAll()).decode("utf-8", "ignore")
        if tail:
            self._sse_buf += tail
            self._drain_sse()

        # v4.102 hotfix：Qt 对 HTTP 4xx/5xx 默认把 reply.error() 报成 NoError，
        # 单靠 error() 判断会把这些错误当成「成功但空内容」→ 用户看到「什么都不显示」。
        # 这里同时用状态码 + SSE 错误事件 + 尾部 JSON 判断，任何失败都明确展示。
        api_err = getattr(self, "_streaming_error", "")
        net_err = reply.error()
        err_str = reply.errorString() or ""
        _vision_debug(f"_on_stream_finished: http={status_code} net_err={net_err} err='{err_str}' api_err={api_err[:80]!r} text_len={len(self._streaming_text)} tail_len={len(tail)} tail={tail[:240]!r}")
        if net_err != QNetworkReply.NoError or (status_code and status_code >= 400) or api_err:
            err_str = reply.errorString() or ""
            # 优先 SSE 错误事件；否则整段 _sse_buf 当 JSON 解析（v4.102 fix4：HTTP 错误
            # 响应可能是 plain JSON 非 SSE，扫 data: 行拿不到），最后回退 Qt 错误串
            msg = api_err or self._extract_api_error(self._sse_buf) or self._extract_api_error(tail) or err_str
            if not msg or msg == err_str:
                for src in (self._sse_buf, tail):
                    if not src: continue
                    try:
                        ej = json.loads(src)
                        if isinstance(ej, dict):
                            err = ej.get("error") or {}
                            m = err.get("message") if isinstance(err, dict) else (err if isinstance(err, str) else "")
                            if m:
                                msg = m
                                break
                    except Exception:
                        continue
            log.error("流式请求失败 HTTP=%s err=%s msg=%s", status_code, err_str, msg)
            if not msg:
                # 服务器无错误正文（如连接被重置 / 协议层错误）时，给出可读提示并指向调试日志
                msg = err_str or f"HTTP {status_code}（无错误详情，详见 vision_debug.log）"
            self.status_label.setText(f"接口错误（HTTP {status_code}）：{msg[:280]}")
        else:
            text = self._streaming_text
            if text:
                session = self.store.active()
                session.messages.append({"role": "assistant", "content": text})
                self._track_context("assistant", text)
                self.store.save()
                self._speak(text)
                self.status_label.clear()
            elif not api_err:
                # 成功状态但模型返回空（极少见）：给个提示，避免「发图后毫无反应」
                self.status_label.setText("（模型返回为空，请重试）")

        self._streaming = False
        self._streaming_text = ""
        self._streaming_error = ""
        self._reply = None  # v4.102: reply 已结束，清掉引用
        self._reset_busy()
        # v4.96：强制全量重建，清除流式气泡残留
        self._rendered_msg_count = 0
        self._flush_render()  # v4.60o：冲刷待渲染定时器，避免末条消息重复插入
        self.input_box.setFocus()
        reply.deleteLater()

    # ============ 配置 ============
    def _on_search_toggle(self, state):
        self.cfg["search_enabled"] = bool(state)
        self._save_cfg()

    def _on_agent_toggle(self, state):
        self.agent_mode = bool(state)
        self.cfg["agent_mode"] = self.agent_mode
        self._save_cfg()
        self.status_label.setText("Agent 模式：开" if self.agent_mode else "Agent 模式：关")

    # 触发自动路由到 Agent 的执行意图关键词（普通模式下命中即按 Agent 执行）
    _ACTION_HINTS = (
        "写文件", "写个文件", "写一份", "保存", "创建文件", "生成文件", "写入文件", "导出文件",
        "运行", "执行", "命令", "脚本", "python", "代码", "计算", "数据处理", "爬取",
        "数据分析", "数据报告", "分析报告",
        "生成图", "画图", "配图", "海报", "插画", "生图", "做图", "生成图片", "生成一张",
        "生成视频", "做视频", "剪视频", "口播视频", "数字人视频", "做口播",
        "定时", "提醒", "闹钟", "倒计时",
        "截图", "截屏",
        "图表", "柱状图", "折线图", "饼图", "散点图", "可视化",
        "浏览器", "清空回收站", "锁屏", "关机", "打开文件", "打开应用", "控制软件", "输入文字", "点击",
        "记笔记", "待办", "备忘",
        "装技能", "搜索技能", "安装技能",
        "公众号", "写文章", "续写", "写稿", "写文案",
        "做ppt", "做 ppt", "做报告", "生成报告", "生成ppt", "做一张",
        "下载", "导出", "批量", "整理文件",
    )

    # v4.57：纯陈述 / 感慨拦截词。命中且不含执行意图时，不当成"需要研究/执行"的任务，
    # 走普通对话（不进 Agent、不联网搜），避免「现在好多平台反应太严了」这类感慨被误判成
    # 研究任务 → agent 连环调 web_search 空转、UI 刷出一堆空「…」气泡。
    _STATEMENT_HINTS = (
        "太严", "太严了", "不让发", "发不了", "发不出", "发不出去", "越来越难", "越来越严",
        "没法发", "没法做", "管得真宽", "限制多", "限制太多", "审核严", "审核太严",
        "烦", "好烦", "无语", "唉", "哎", "哎哟", "算了", "无奈", "难受", "头大",
        "太难了", "真难", "卷", "太卷了", "搞不动", "干不动", "心累", "累",
    )

    # 选题/盘点/列方向 类关键词（v4.56 新增）：命中时不走 Agent 也不走联网搜索，
    # 直接让 LLM 用训练知识出文本——避免把「给我列几个选题」误判为「去搜实时榜单」。
    # 若同时含"写文件/导出/搜最新"等执行词则走 Agent（写文件型选题清单）。
    _TOPIC_HINTS = (
        "列方向", "列选题", "想几个", "盘点", "选题", "爆款方向", "做什么内容",
        "给我想", "给我建议", "推荐方向", "哪些方向", "哪些选题", "哪些赛道",
        "有什么选题", "给我几个", "列几个", "出主意", "给我列", "写什么",
    )

    # v4.56 补：内容平台 + 方向词（隐含"列方向"意图但用户没明说"列"字）
    # 命中时不走联网搜（搜狗/百度/Bing 拿不到真实榜单数据，徒增延迟和频控风险），
    # 直接让 LLM 用训练知识出文本。
    _TOPIC_PLATFORMS = ("小红书", "抖音", "视频号", "公众号", "知乎", "微博", "b站", "bilibili", "快手")
    _TOPIC_DIR_KEYWORDS = (
        "爆款", "趋势", "风向", "方向", "赛道", "品类", "选题", "增长", "画像",
        "做什么", "写什么", "发什么", "内容", "玩法", "风格", "推荐", "建议",
        "榜单", "最新", "热门", "爆火", "火",
    )
    # v4.61：纯咨询 / 问意见类关键词——命中且无执行动作词时，即便 Agent 模式也降级为普通对话直答
    _ADVICE_HINTS = (
        "意见", "建议", "给点", "给些", "说说看", "说说", "怎么看", "怎么想",
        "怎么破", "破局", "看法", "你觉得", "支招", "出主意", "点评", "评价",
        "帮我看看", "怎么弄", "怎么办", "咋办", "有啥建议", "有啥想法",
        "给个主意", "参谋", "把脉", "指点",
    )

    def _message_is_statement_only(self, text):
        """v4.57：判断消息是否仅为"陈述 / 感慨"而非请求执行的任务。
        命中陈述词且不含任何执行意图词时返回 True → 走普通对话，避免误进 Agent / 联网搜。
        """
        if not text:
            return False
        t = text.lower()
        if not any(h.lower() in t for h in self._STATEMENT_HINTS):
            return False
        # 虽含陈述词，但明显也在交代执行任务（如「帮我整理文件，太乱了」）→ 仍走 Agent
        if any(h.lower() in t for h in self._ACTION_HINTS):
            return False
        return True

    def _message_needs_agent(self, text):
        """普通模式下，若用户消息含执行/写文件/分析/生成等意图，自动走 Agent 真正执行。"""
        if not text:
            return False
        t = text.lower()
        # v4.57：纯陈述 / 感慨不当成执行任务，回到普通对话
        if self._message_is_statement_only(text):
            return False
        # v4.102 fix11：媒体生成组合判定——用户说"生成口播视频 / 做个数字人 / 生成视频"
        # 时，"生成/做"与"视频/口播/数字人"常分开写，_ACTION_HINTS 的连写词（生成视频/做视频）
        # 会漏判 → 普通模式不自动进 Agent。组合判定兜底命中这类表述。
        if self._is_media_gen_request(text):
            return True
        # v4.107：导演指令已从主路由摘除——统一在导演台底部「导演对话」条处理，
        # 主对话框不再进 Agent。故此处不再识别 _is_director_command。
        return any(h.lower() in t for h in self._ACTION_HINTS)

    # v4.102 fix11：媒体生成组合词表——"生成/做/制作/创建 视频/口播/数字人/图"常被用户
    # 分开写（如「生成口播视频」「做个数字人」「生成一段视频」），连写词表
    # （生成视频/做视频/生视频）会漏判。这里用「媒体对象词 × 媒体动作词」同现判定。
    # 注意：动作词收敛为明确的「生成/制作/做一段/做一个/来一段/出一段/合成/剪/拍」等，
    # 刻意不放单字「做/出/配」——否则「视频号好难做」「我写了口播文案」这类纯陈述/分享
    # 会被误判为媒体生成。媒体对象词也排除宽泛的「配音」单列（单用「配音」由连写词表管）。
    _MEDIA_OBJ_KW = ("视频", "短视频", "口播", "数字人", "数字分身", "旁白",
                     "图片", "海报", "插画", "头像", "混剪", "微电影", "短片",
                     "配音", "字幕")
    _MEDIA_VERB_KW = ("生成", "制作", "做一个", "做一段", "做个", "做一条", "做几",
                      "来一段", "来一个", "出一段", "出一个", "合成", "剪一个", "剪一段",
                      "拍一段", "录一段", "创作一段", "帮我做", "帮我生成", "给我做",
                      "给我生成", "配一个", "配一段")

    def _is_media_gen_request(self, text):
        """是否「媒体生成/制作类」请求：命中一个媒体对象词 与 一个媒体动作词 同现即 True。
        覆盖：生成口播视频 / 帮我按文案做视频 / 做个数字人 / 生成一段视频 / 出一段短片。
        排除：单字「做/出/配」的宽泛使用（视频号好难做 / 我写了口播文案都不是生成）。"""
        if not text:
            return False
        t = text.lower()
        has_obj = any(k in t for k in self._MEDIA_OBJ_KW)
        has_verb = any(v in t for v in self._MEDIA_VERB_KW)
        if not (has_obj and has_verb):
            return False
        # 排除纯陈述/感慨：如「视频号最近好难做」——命中陈述词则不当媒体生成
        if self._message_is_statement_only(text):
            return False
        return True

    # v4.106：导演台对话指令词表——对象词收窄到导演台特有名词（分镜/关键帧/三视图/
    # 成片/导演台/主角/角色），避免「换个头像」「改改文章」这类普通编辑误进 Agent。
    # 单字「镜」只覆盖「第3镜/这镜/下一镜」写法；误伤面小（含"镜"且带明确改动的句子
    # 基本都在聊镜头）。
    _DIRECTOR_OBJ_KW = ("导演台", "分镜", "关键帧", "三视图", "成片",
                        "主角", "角色", "这镜", "这一镜", "下一镜", "末镜", "镜")
    _DIRECTOR_VERB_KW = ("重生成", "重新生成", "改成", "换成", "换回", "修改", "改一下",
                         "改改", "重做", "合成", "再来一版", "补一镜", "加一镜", "删掉",
                         "去掉", "调亮", "调暗")
    _DIRECTOR_STATUS_KW = ("导演台", "分镜", "关键帧", "三视图", "成片",
                           "第几镜", "哪几镜", "几镜")
    _DIRECTOR_STATUS_Q = ("进度", "状态", "怎么样", "好了吗", "好了没", "跑到哪",
                          "生成到哪", "到哪一步", "卡在哪", "到哪了", "停在哪",
                          "还没", "没生成", "没好", "没出", "没跑")

    def _is_director_command(self, text):
        """v4.106：是否导演台对话指令（查进度 / 改分镜 / 重生成关键帧·三视图 / 合成）。
        命中 → 普通模式自动进 Agent，且升舱 DeepSeek（director_* 工具意图）。"""
        if not text:
            return False
        t = text.lower()
        if self._message_is_statement_only(text):
            return False
        # 进度查询：「导演台进度怎么样 / 分镜跑到哪了」
        if any(o in t for o in self._DIRECTOR_STATUS_KW) and \
                any(q in t for q in self._DIRECTOR_STATUS_Q):
            return True
        # 修改指令：「把第3镜的关键帧改成夜晚 / 主角换成短发 / 合成成片」
        return any(o in t for o in self._DIRECTOR_OBJ_KW) and \
            any(v in t for v in self._DIRECTOR_VERB_KW)

    def _message_is_topic_only(self, text):
        """v4.56：用户消息是否属于"列方向/选题/盘点"型需求，且**不**含执行意图。
        命中则普通模式下直接走 LLM 裸出，不联网搜——避免误判为搜索任务。
        命中但同时含"写文件/导出"等执行词则返回 False（让 _message_needs_agent 接管）。
        扩展：内容平台名 + 方向词组合（如「小红书爆款」）也算"列方向"型。
        """
        if not text:
            return False
        t = text.lower()
        if any(h.lower() in t for h in self._TOPIC_HINTS):
            # 含执行意图词则不算"纯盘点"，交给 _message_needs_agent 走 Agent
            if any(h.lower() in t for h in self._ACTION_HINTS):
                return False
            return True
        # 隐式匹配：平台 + 方向词（不带"搜"字）
        has_platform = any(p in text for p in self._TOPIC_PLATFORMS)
        has_dir = any(k in text for k in self._TOPIC_DIR_KEYWORDS)
        if has_platform and has_dir and "搜" not in t and "查" not in t:
            # 若含执行意图（写文件/导出/抓取）则让 _message_needs_agent 接管
            if any(h.lower() in t for h in self._ACTION_HINTS):
                return False
            return True
        return False

    def _message_is_advice_only(self, text):
        """v4.61：判断消息是否为「纯咨询 / 问意见 / 求建议」类，且不含任何执行动作词。
        命中则即便 Agent 执行模式开启，也应降级为普通对话直答，避免进 Agent 后
        被技能（如 content-gap-analysis 的 Stop and ask）带偏，反复追问刷屏。
        """
        if not text:
            return False
        t = text.lower()
        if not any(h in t for h in self._ADVICE_HINTS):
            return False
        # 含执行动作词（写文件 / 生成 / 搜索 / 下载等）→ 不是纯咨询，交给 Agent 执行
        if any(h.lower() in t for h in self._ACTION_HINTS):
            return False
        return True

    def _on_skip_confirm_toggle(self, state):
        self.agent_skip_confirm = bool(state)
        self.cfg["agent_skip_confirm"] = self.agent_skip_confirm
        self._save_cfg()
        self.status_label.setText(
            "危险操作免确认：开" if self.agent_skip_confirm else "危险操作免确认：关")

    def _on_mode_change(self, index):
        """执行模式下拉变更：写回引擎 + 持久化 + 状态栏反馈。"""
        mode = self.sender().itemData(index)
        if not mode:
            return
        self.permission_engine.set_mode(mode)
        self.permission_mode = mode
        # 兼容旧开关：auto 等价于旧的「危险操作免确认」全开
        self.agent_skip_confirm = (mode == "auto")
        self.cfg["agent_skip_confirm"] = self.agent_skip_confirm
        self.cfg["permission_mode"] = mode
        self._save_cfg()
        self.status_label.setText(f"执行模式：{MODES[mode].split('（')[0]}")

    def _on_trust_session(self):
        """本次会话全部信任：危险操作不再逐个问（重启后失效）。"""
        self.permission_engine.set_session_trusted()
        self.session_trusted = True
        self.status_label.setText("执行模式：本次会话已信任")

    def _save_cfg(self):
        """v4.58：原子写入 config.json，防崩溃损坏配置。"""
        try:
            tmp_path = CONFIG_PATH + ".tmp"
            with open(tmp_path, "w", encoding="utf-8") as f:
                json.dump(self.cfg, f, ensure_ascii=False, indent=2)
            os.replace(tmp_path, CONFIG_PATH)
        except Exception as e:
            log.error("写 config.json 失败: %s", e)

    def _on_clipboard_event(self, result):
        """剪贴板分类结果 → 托盘通知（轻量，不打断）。"""
        tray = getattr(self, "tray_app", None)
        ti = getattr(tray, "tray", None) if tray else None
        if result["type"] == "url":
            titles = result.get("titles") or result.get("urls")
            title = "剪贴板 · 链接"
            msg = "🔗 " + "  |  ".join(titles)[:200]
        elif result["type"] == "code":
            title = "剪贴板 · 代码"
            msg = f"💻 检测为 {result['language']} 代码，{result['lines']} 行"
        elif result["type"] == "image_path":
            title = "剪贴板 · 图片"
            msg = "🖼️ 检测到图片路径，可让小臭识别"
        elif result["type"] == "text":
            title = "剪贴板 · 文本"
            msg = f"📝 长文本 {result['length']} 字，可翻译/摘要"
        else:
            return
        if ti:
            ti.showMessage(title, msg, QSystemTrayIcon.Information, 8000)

    def _clear_ghost(self):
        try:
            QApplication.processEvents()
            self.repaint()
            if sys.platform.startswith("win"):
                user32 = ctypes.windll.user32
                RDW_ERASE = 0x0001
                RDW_FRAME = 0x0400
                RDW_INVALIDATE = 0x0002
                RDW_ALLCHILDREN = 0x0080
                user32.RedrawWindow(
                    None, None, None,
                    RDW_ERASE | RDW_FRAME | RDW_INVALIDATE | RDW_ALLCHILDREN,
                )
        except Exception:
            pass

    def closeEvent(self, event):
        # v4.47: 点×彻底退出（之前是 event.ignore()+hide 缩到托盘，导致多实例/僵尸累积）
        event.accept()
        try:
            # v4.108 M-20：关窗前终止 Agent worker——request_stop 会唤醒阻塞中的
            # 危险操作确认（_confirm_event.set），wait(3s) 给子线程收尾，避免
            # 「QThread: Destroyed while thread is still running」崩溃与确认弹窗挂死。
            w = self._agent_worker
            if w is not None and w.isRunning():
                w.request_stop()
                w.wait(3000)
        except Exception:
            pass
        try:
            self._clear_ghost()
        except Exception:
            pass
        app = QApplication.instance()
        if app is not None:
            app.quit()


class TrayApp:
    def __init__(self, app, window, cfg):
        self.app = app
        self.window = window
        self.cfg = cfg
        self.tray = QSystemTrayIcon(get_app_icon(), app)

        menu = QMenu()
        show_action = QAction("显示 / 隐藏", app)
        show_action.triggered.connect(self.toggle)
        skill_action = QAction("🧩 技能管理器", app)
        skill_action.triggered.connect(self._open_skill_manager)
        market_action = QAction("🛍 技能市场", app)
        market_action.triggered.connect(self._open_skill_market)
        wf_action = QAction("⚙️ 工作流模板", app)
        wf_action.triggered.connect(self._open_workflow_manager)
        review_action = QAction("📝 技能审核", app)
        review_action.triggered.connect(lambda: self.window._open_skill_review_dialog())
        self._skill_review_action = review_action  # v4.84：动态更新待审数量
        diag_action = QAction("📦 诊断包", app)
        diag_action.triggered.connect(self.export_diagnostic)
        quit_action = QAction("退出", app)
        quit_action.triggered.connect(self.quit)
        menu.addAction(show_action)
        menu.addAction(skill_action)
        menu.addAction(market_action)
        menu.addAction(wf_action)
        menu.addAction(review_action)
        menu.addAction(diag_action)
        menu.addAction(quit_action)
        self.tray.setContextMenu(menu)
        self.tray.setToolTip("Agent 桌面助手")
        self.tray.activated.connect(self._on_activated)
        self.tray.show()
        self._setup_hotkey()

    def _on_activated(self, reason):
        if reason == QSystemTrayIcon.DoubleClick:
            self.toggle()

    def toggle(self):
        if self.window.isVisible():
            self.window.hide()
        else:
            self.window.show()
            self.window.raise_()
            self.window.activateWindow()

    def _open_skill_manager(self):
        try:
            from skill_manager_ui import open_skill_manager
            open_skill_manager(self.cfg)
        except Exception as e:
            try:
                self.tray.showMessage("技能管理器", f"打开失败：{e}",
                                      QSystemTrayIcon.Warning, 5000)
            except Exception:
                pass

    def _open_skill_market(self):
        try:
            from skill_market_ui import open_skill_market
            open_skill_market(self.cfg)
        except Exception as e:
            try:
                self.tray.showMessage("技能市场", f"打开失败：{e}",
                                      QSystemTrayIcon.Warning, 5000)
            except Exception:
                pass

    def _open_workflow_manager(self):
        try:
            from workflow_manager_ui import open_workflow_manager
            open_workflow_manager(self.cfg, self.window)
        except Exception as e:
            try:
                self.tray.showMessage("工作流模板", f"打开失败：{e}",
                                      QSystemTrayIcon.Warning, 5000)
            except Exception:
                pass

    def export_diagnostic(self):
        try:
            from diagnostic_export import export_diagnostic_package
            export_diagnostic_package(self.window)
        except Exception as e:
            try:
                self.tray.showMessage("诊断包", f"导出失败：{e}",
                                      QSystemTrayIcon.Warning, 5000)
            except Exception:
                pass

    def _setup_hotkey(self):
        try:
            import keyboard
            keyboard.add_hotkey(
                self.cfg["hotkey"],
                lambda: QTimer.singleShot(0, self.toggle),
            )
        except Exception:
            pass

    def quit(self):
        self.tray.hide()
        self.app.quit()




